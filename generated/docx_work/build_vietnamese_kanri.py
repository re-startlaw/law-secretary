from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
SRC = Path((ROOT / ".codex" / "shell_path_utf8.txt").read_text(encoding="utf-8").splitlines()[0])


def output_path(src: Path) -> Path:
    return src.with_name(f"{src.stem}_ベトナム語版{src.suffix}")


def set_latin_font(run, size_pt=None, bold=None):
    run.font.name = "Times New Roman"
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")


def replace_paragraph(paragraph, text, *, title=False, heading=False):
    if not paragraph.runs:
        run = paragraph.add_run()
    else:
        run = paragraph.runs[0]
    run.text = text
    for extra in paragraph.runs[1:]:
        extra.text = ""
    if title:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_latin_font(run, size_pt=16, bold=True)
    elif heading:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.keep_with_next = True
        set_latin_font(run, size_pt=12, bold=True)
    else:
        set_latin_font(run, size_pt=11, bold=False)


translations = [
    ("監理措置決定を希望する理由", "Lý do đề nghị được áp dụng biện pháp giám sát", "title"),
    (
        "私は、このたび監理措置決定を希望し、その理由を以下のとおり申し述べます。",
        "Tôi xin đề nghị được áp dụng biện pháp giám sát và trình bày lý do như sau.",
        "body",
    ),
    (
        "私は現在、妻及び0歳の子と同居して生活しております（資料１　住民票）。また、0歳の子を養育しております（資料３　子の出生届受理証明書）。",
        "Hiện nay, tôi đang sống cùng vợ và con 0 tuổi (Tài liệu 1: Phiếu cư trú). Ngoài ra, tôi đang nuôi dưỡng con 0 tuổi (Tài liệu 3: Giấy chứng nhận thụ lý khai sinh của con).",
        "body",
    ),
    (
        "妻は現在育児休業中です（資料４　育児休業給付金支給決定通知書）。子がまだ0歳であり、授乳、通院、夜間対応を含めて日常的に手が掛かる時期であるため、私が不在となることの影響は非常に大きいです。特に、0歳の子については日常的な見守りと世話が欠かせず、妻一人に育児及び生活上の負担が集中することになります。私が家族と同居し、生活を支えることは、子の健全な養育環境を維持する上でも必要です。",
        "Vợ tôi hiện đang nghỉ chăm con (Tài liệu 4: Thông báo quyết định chi trả trợ cấp nghỉ chăm con). Vì con tôi vẫn còn 0 tuổi, đang trong giai đoạn cần được chăm sóc hằng ngày, bao gồm bú sữa, đi khám bệnh và chăm sóc ban đêm, nên việc tôi vắng mặt sẽ gây ảnh hưởng rất lớn. Đặc biệt, đối với con 0 tuổi, việc trông nom và chăm sóc hằng ngày là không thể thiếu; nếu tôi bị vắng mặt, toàn bộ gánh nặng chăm con và sinh hoạt sẽ dồn lên một mình vợ tôi. Việc tôi sống cùng gia đình và hỗ trợ đời sống gia đình cũng là cần thiết để duy trì môi trường nuôi dưỡng lành mạnh cho con.",
        "body",
    ),
    (
        "また、私は現在の住居において安定して生活しております（資料２　賃貸借契約書）。今後も同住所で家族と生活する予定であり、所在が不安定になる事情はありません。",
        "Ngoài ra, tôi đang sinh sống ổn định tại nơi ở hiện nay (Tài liệu 2: Hợp đồng thuê nhà). Trong thời gian tới, tôi vẫn dự định tiếp tục sống cùng gia đình tại địa chỉ này, và không có hoàn cảnh nào khiến nơi ở của tôi trở nên không ổn định.",
        "body",
    ),
    (
        "監理人候補は妻です。妻は私と同居しており、私の日常生活の状況、生活時間、交友関係、家庭内の様子を継続的に把握することができます。妻の身分証及び監理人承諾書兼誓約書を提出します（資料５　妻の身分証、資料６　監理人承諾書兼誓約書）。したがって、妻は私に対する日常的な指導・監督を行うことができ、監理人として適切に監理を行うことが可能です。私自身も、妻の監理の下で生活し、必要な指示に従う意思があります。",
        "Người dự kiến làm người giám sát là vợ tôi. Vợ tôi sống cùng tôi và có thể liên tục nắm bắt tình hình sinh hoạt hằng ngày, thời gian sinh hoạt, quan hệ bạn bè cũng như tình hình trong gia đình của tôi. Tôi sẽ nộp giấy tờ tùy thân của vợ và Giấy chấp thuận kiêm cam kết của người giám sát (Tài liệu 5: Giấy tờ tùy thân của vợ; Tài liệu 6: Giấy chấp thuận kiêm cam kết của người giám sát). Vì vậy, vợ tôi có thể hướng dẫn và giám sát tôi trong sinh hoạt hằng ngày, và có thể thực hiện việc giám sát một cách phù hợp với tư cách là người giám sát. Bản thân tôi cũng có ý định sống dưới sự giám sát của vợ và tuân theo các chỉ thị cần thiết.",
        "body",
    ),
    (
        "私は、逃亡したり、証拠を隠滅したりする意思は全くありません。地方出入国在留管理官署その他関係機関から出頭や報告を求められた場合には、必ずこれに応じ、指定された条件を遵守します。今後、指定された日時・場所への出頭、連絡、届出等については、すべて誠実に対応することを誓います。今後も現在の住居で生活を継続する予定です。",
        "Tôi hoàn toàn không có ý định bỏ trốn hoặc che giấu, tiêu hủy chứng cứ. Nếu Cục Quản lý xuất nhập cảnh địa phương hoặc các cơ quan liên quan khác yêu cầu tôi trình diện hoặc báo cáo, tôi nhất định sẽ thực hiện và tuân thủ các điều kiện được chỉ định. Tôi xin cam kết sẽ thành thật thực hiện tất cả các yêu cầu về trình diện tại ngày giờ, địa điểm được chỉ định, liên lạc, khai báo và các thủ tục khác. Trong thời gian tới, tôi dự định tiếp tục sinh sống tại nơi ở hiện nay.",
        "body",
    ),
    (
        "私は、今回違反調査を受けるに至ったことについて深く反省しております。今後は家族の監督の下で生活を整え、交友関係や生活環境を見直し、再発防止に努めます。社会内で生活することができれば、家族との生活を維持しながら、関係機関の指示に従って適切に手続を受けることができます。",
        "Tôi đang suy nghĩ sâu sắc và hối hận về việc đã dẫn đến việc bị điều tra vi phạm lần này. Trong thời gian tới, dưới sự giám sát của gia đình, tôi sẽ ổn định lại sinh hoạt, xem xét lại các mối quan hệ bạn bè và môi trường sống, đồng thời nỗ lực phòng ngừa tái phạm. Nếu được sinh sống trong cộng đồng, tôi có thể duy trì đời sống cùng gia đình, đồng thời tuân theo chỉ thị của các cơ quan liên quan và tham gia thủ tục một cách thích hợp.",
        "body",
    ),
    (
        "以上のとおり、私には安定した住居があり、監理人候補である妻による継続的な監理を受けることができます。また、私には逃亡や証拠隠滅をする動機も理由もありません。他方で、収容によって妻及び0歳の子に及ぶ不利益は極めて大きいと考えております。したがって、私については、監理措置の下で社会内において生活しながら手続を進められるよう、監理措置決定をしていただきたく、何卒お願いいたします。",
        "Như đã trình bày ở trên, tôi có nơi ở ổn định và có thể được vợ tôi, người dự kiến làm người giám sát, giám sát liên tục. Ngoài ra, tôi không có động cơ hay lý do để bỏ trốn hoặc che giấu, tiêu hủy chứng cứ. Mặt khác, nếu tôi bị đưa vào cơ sở lưu trú, bất lợi đối với vợ tôi và con 0 tuổi sẽ là vô cùng lớn. Vì vậy, tôi kính đề nghị cho tôi được áp dụng biện pháp giám sát để có thể tiến hành thủ tục trong khi sinh sống trong cộng đồng dưới biện pháp giám sát.",
        "body",
    ),
    ("添付資料一覧", "Danh mục tài liệu đính kèm", "heading"),
    (
        "資料1　住民票（妻・子との同居が分かるもの）",
        "Tài liệu 1: Phiếu cư trú (thể hiện việc sống cùng vợ và con)",
        "body",
    ),
    ("資料2　賃貸借契約書", "Tài liệu 2: Hợp đồng thuê nhà", "body"),
    (
        "資料3　子の出生届受理証明書",
        "Tài liệu 3: Giấy chứng nhận thụ lý khai sinh của con",
        "body",
    ),
    (
        "資料4　育児休業給付金支給決定通知書",
        "Tài liệu 4: Thông báo quyết định chi trả trợ cấp nghỉ chăm con",
        "body",
    ),
    ("資料5　妻の身分証", "Tài liệu 5: Giấy tờ tùy thân của vợ", "body"),
    (
        "資料6　監理人承諾書兼誓約書",
        "Tài liệu 6: Giấy chấp thuận kiêm cam kết của người giám sát",
        "body",
    ),
]

OUT = output_path(SRC)
shutil.copy2(SRC, OUT)

doc = Document(OUT)
mapping = {jp: (vi, kind) for jp, vi, kind in translations}
remaining = set(mapping)

for paragraph in doc.paragraphs:
    original = paragraph.text
    if original in mapping:
        vi_text, kind = mapping[original]
        replace_paragraph(
            paragraph,
            vi_text,
            title=(kind == "title"),
            heading=(kind == "heading"),
        )
        remaining.remove(original)

if remaining:
    raise RuntimeError(f"Untranslated paragraphs: {len(remaining)}")

doc.save(OUT)
print(OUT)
