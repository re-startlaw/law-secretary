from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "監理措置決定を希望する理由.docx"


def set_run_font(run, size_pt=11, bold=False):
    run.font.name = "Yu Mincho"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "游明朝")
    r_fonts.set(qn("w:ascii"), "Yu Mincho")
    r_fonts.set(qn("w:hAnsi"), "Yu Mincho")


def set_paragraph_font(paragraph, size_pt=11, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size_pt=size_pt, bold=bold)


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.35, first_line=True):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(11)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    set_paragraph_font(p)
    set_paragraph_spacing(p)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    set_paragraph_font(p, size_pt=12, bold=True)
    set_paragraph_spacing(p, before=12, after=6, line=1.25, first_line=False)
    return p


def add_material_item(doc, index, text):
    p = doc.add_paragraph()
    p.add_run(f"資料{index}　{text}")
    set_paragraph_font(p)
    set_paragraph_spacing(p, after=3, line=1.2, first_line=False)
    p.paragraph_format.left_indent = Pt(22)
    p.paragraph_format.first_line_indent = Pt(-22)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(25)
section.right_margin = Mm(25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Yu Mincho"
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
normal._element.rPr.rFonts.set(qn("w:ascii"), "Yu Mincho")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Yu Mincho")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("監理措置決定を希望する理由")
set_run_font(title_run, size_pt=16, bold=True)
title.paragraph_format.space_after = Pt(18)

paragraphs = [
    "私は、このたび監理措置決定を希望し、その理由を以下のとおり申し述べます。",
    "私は現在、妻及び0歳の子と同居して生活しております。この同居関係は、住民票により明らかにすることができます（資料１　住民票）。また、子が0歳であることは、出生届受理証明書によって確認することができます（資料３　子の出生届受理証明書）。",
    "妻は現在育児休業中であり、そのことは育児休業給付金支給決定通知書により確認することができます（資料４　育児休業給付金支給決定通知書）。子がまだ0歳であり、授乳、通院、夜間対応を含めて日常的に手が掛かる時期であるため、私が不在となることの影響は非常に大きいです。特に、0歳の子については日常的な見守りと世話が欠かせず、妻一人に育児及び生活上の負担が集中することになります。私が家族と同居し、生活を支えることは、子の健全な養育環境を維持する上でも必要です。",
    "また、私は現在の住居において安定して生活しており、住居については賃貸借契約書により明らかにすることができます（資料２　賃貸借契約書）。今後も同住所で家族と生活する予定であり、所在が不安定になる事情はありません。",
    "監理人候補は妻です。妻は私と同居しており、私の日常生活の状況、生活時間、交友関係、家庭内の様子を継続的に把握することができます。妻の本人確認資料として身分証を提出し（資料５　妻の身分証）、また、妻は監理人としての承諾及び誓約を行っています（資料６　監理人承諾書兼誓約書）。したがって、妻は私に対する日常的な指導・監督を行うことができ、監理人として適切に監理を行うことが可能です。私自身も、妻の監理の下で生活し、必要な指示に従う意思があります。",
    "私は、逃亡したり、証拠を隠滅したりする意思は全くありません。地方出入国在留管理官署その他関係機関から出頭や報告を求められた場合には、必ずこれに応じ、指定された条件を遵守します。今後、指定された日時・場所への出頭、連絡、届出等については、すべて誠実に対応することを誓います。これまでの所在も安定しており、今後も現在の住居で生活を継続する予定です。",
    "私は、今回違反調査を受けるに至ったことについて深く反省しております。今後は家族の監督の下で生活を整え、交友関係や生活環境を見直し、再発防止に努めます。社会内で生活することができれば、家族との生活を維持しながら、関係機関の指示に従って適切に手続を受けることができます。",
    "以上のとおり、私には安定した住居があり、監理人候補である妻による継続的な監理が可能であり、逃亡又は証拠隠滅のおそれは低いと考えております。他方で、収容によって妻及び0歳の子に及ぶ不利益は極めて大きいと考えております。したがって、私については、監理措置の下で社会内において生活しながら手続を進めることが相当であると考えますので、監理措置決定をしていただきたく、お願いいたします。",
]

for text in paragraphs:
    add_body_paragraph(doc, text)

add_heading(doc, "添付資料一覧")
materials = [
    "住民票（妻・子との同居が分かるもの）",
    "賃貸借契約書",
    "子の出生届受理証明書",
    "育児休業給付金支給決定通知書",
    "妻の身分証",
    "監理人承諾書兼誓約書",
]
for i, text in enumerate(materials, 1):
    add_material_item(doc, i, text)

doc.save(OUT)
print(OUT)
