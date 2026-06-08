"""馬さん案件・受任通知_3 を作成（修正履歴ナシ・番号繰上げ済み）。

- ベース: 元 260512電子内容証明案.docx を _3 にコピー済み
- 修正履歴は付けず、最終文面で直接書き換える
- 指示6（G10学費の事実関係加筆）は **適用しない**
- 第４損害項目で旧(5)削除 → 旧(6)(7)を(5)(6)に繰上げ、新規(7)(8)(9)を末尾追加
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
NSMAP = {"w": W_NS}

INPUT = Path(
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/"
    "共有用/01_事件記録/ま_馬さん/Wordファイル/260512電子内容証明案_3.docx"
)


def extract_first_rpr(p_elem: etree._Element) -> etree._Element | None:
    return p_elem.find(f".//{W}r/{W}rPr", namespaces=NSMAP)


def set_paragraph_text(p_elem: etree._Element, new_text: str) -> None:
    """段落本文を単一 run で置換（pPr は保持、書式は先頭 run の rPr を継承）。"""
    base_rpr = extract_first_rpr(p_elem)
    pPr = p_elem.find(f"{W}pPr", namespaces=NSMAP)
    # 既存の本文子要素（pPr 以外）を全削除
    for child in list(p_elem):
        if child.tag == f"{W}pPr":
            continue
        p_elem.remove(child)
    # 新 run
    r = etree.SubElement(p_elem, f"{W}r")
    if base_rpr is not None:
        r.insert(0, copy.deepcopy(base_rpr))
    t = etree.SubElement(r, f"{W}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = new_text


def delete_paragraph_element(p_elem: etree._Element) -> None:
    p_elem.getparent().remove(p_elem)


def insert_paragraph_after(after_p_elem: etree._Element, text: str) -> etree._Element:
    """after の直後に同書式の新段落を挿入して返す。"""
    base_pPr = after_p_elem.find(f"{W}pPr", namespaces=NSMAP)
    base_rpr = extract_first_rpr(after_p_elem)
    new_p = etree.Element(f"{W}p")
    if base_pPr is not None:
        new_p.append(copy.deepcopy(base_pPr))
    r = etree.SubElement(new_p, f"{W}r")
    if base_rpr is not None:
        r.insert(0, copy.deepcopy(base_rpr))
    t = etree.SubElement(r, f"{W}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    after_p_elem.addnext(new_p)
    return new_p


def main() -> None:
    doc = Document(str(INPUT))
    paragraphs = doc.paragraphs

    # ---- 修正1+2: Para 24 見出し ----
    p24 = paragraphs[24]
    assert p24.text.strip() == "（１）　小牧氏がAngelina氏の襟元を掴んで引き移動させた行為"
    set_paragraph_text(
        p24._element,
        "（１）　小牧氏がAngelina氏の襟元付近に身体的接触を加えて移動させ、トラック脇で停止させた行為",
    )

    # ---- 修正1+2: Para 25 本文 ----
    p25 = paragraphs[25]
    assert p25.text.startswith("　令和８年４月１３日午前９時１８分頃")
    set_paragraph_text(
        p25._element,
        (
            "　令和８年４月１３日午前９時２０分頃から同９時３０分頃までの間、貴校体育館において、"
            "小牧氏は、Angelina氏に対し、後方からその襟元付近に身体的接触を加えてこれを掴み、"
            "Angelina氏の身体を伴って移動させた上、トラック脇においてこれを停止させるとともに、"
            "大声で叱責しました。当該事実は、Angelina氏の同日付書面陳述、当方が保有する事案"
            "当日の映像等から、優に認められるものです。"
        ),
    )

    # ---- 修正3: Para 34 数学授業 ----
    p34 = paragraphs[34]
    assert p34.text.startswith("　Angelina氏は、令和８年４月１３日の１時間目の算数")
    set_paragraph_text(
        p34._element,
        (
            "　Angelina氏は、令和８年４月１３日の１時間目の数学の時間（午前８時４０分から"
            "同９時３５分まで）中、約４分間体育館に立ち寄ったことを認めています。これは、"
            "Angelina氏が自身の課題を比較的早く終えた後、長男であるGeorge氏の約４分間の"
            "クロスカントリー競技を観覧するため、体育館に向かったものです。当該時間は、"
            "前半が教員による講義の時間、後半が各生徒が課題を行う自習の時間として運用"
            "されていたものであり、また、当該自習時間中に生徒が一時的に教室を離れること"
            "についても、現場の教員において従前から黙認されてきたものです。"
        ),
    )

    # ---- 修正4: Para 40 (2) 精神状態悪化 ----
    p40 = paragraphs[40]
    assert p40.text.startswith("（２）　貴校から不適切な対応を受けたこと")
    set_paragraph_text(
        p40._element,
        (
            "（２）　貴校から不適切な対応を受けたこと、登校できないこと、及び事実に反して"
            "Angelina氏が虚偽の主張をしている旨を周囲に言いふらされたことによりAngelina氏"
            "の精神状態が悪化し、投薬調整及び継続的な心理カウンセリングを要する状態となった"
            "ことに伴う精神科通院費用及び心理カウンセリング費用"
        ),
    )

    # ---- 修正5: Para 43 (5) 妻関連損害 → 削除 ----
    p43 = paragraphs[43]
    assert p43.text.startswith("（５）　George氏の看護者の校内立入りが何らの根拠なく拒絶")
    delete_paragraph_element(p43._element)

    # 以降の paragraphs インデックスはずれるので、再取得して旧(6)(7)を新(5)(6)に
    doc2_paragraphs = doc.paragraphs
    # 旧(6) 「（６）　既に購入済みのAngelina氏の通学定期券に関する損失」を (5) に
    p_kippu = next(
        p
        for p in doc2_paragraphs
        if (p.text or "").startswith("（６）　既に購入済みのAngelina氏の通学定期券")
    )
    set_paragraph_text(
        p_kippu._element,
        "（５）　既に購入済みのAngelina氏の通学定期券に関する損失",
    )
    # 旧(7) 「（７）　弁護士費用…」を (6) に
    p_fee = next(
        p
        for p in doc.paragraphs
        if (p.text or "").startswith("（７）　弁護士費用その他本件への法的対応")
    )
    set_paragraph_text(
        p_fee._element,
        "（６）　弁護士費用その他本件への法的対応のために通知人が要した費用",
    )

    # ---- 修正9: 新規 (7)(8)(9) を弁護士費用段落の直後に追加 ----
    anchor = p_fee._element
    new_items = [
        "（７）　Angelina氏の長期にわたる学籍中断により、Angelina氏の学業継続性が毀損されていることに関する損害",
        "（８）　Angelina氏の試験成績、GPA及びTranscriptへの不利益な影響に関する損害",
        "（９）　Angelina氏の大学進学経路への現実的影響並びに教育機会の喪失及び将来的進学リスクに関する損害",
    ]
    for txt in new_items:
        anchor = insert_paragraph_after(anchor, txt)

    # ---- 修正8: 第５(2) 「直ちに復学」→ 公平・安全な教育環境 ----
    p_fukugaku = next(
        p
        for p in doc.paragraphs
        if (p.text or "").strip() == "（２）　Angelina氏を直ちに復学させていただくこと"
    )
    set_paragraph_text(
        p_fukugaku._element,
        (
            "（２）　Angelina氏に対する公平かつ安全な教育環境を確保した上で、"
            "Angelina氏の復学を実現していただくこと（なお、当方として復学権自体を放棄"
            "するものではありません。）"
        ),
    )

    # ---- 修正7: 第６(3) 学費相当額 → 1,388,000円 ----
    p_gakuhi = next(
        p for p in doc.paragraphs if "学費相当額（金○○万円）" in (p.text or "")
    )
    set_paragraph_text(
        p_gakuhi._element,
        (
            "（３）　もっとも、Angelina氏及び通知人は、貴校との在学契約の継続を強く"
            "希望しております。当方は、学費相当額（金１,３８８,０００円。Angelina氏の"
            "Ｇ１０学年第一期学費に相当します。）を当職の弁護士預り金口座において保全し、"
            "貴校による教育役務の提供再開と引換えに、直ちに同額をお支払いする用意がございます。"
        ),
    )

    # ---- 指示6 は **適用しない**（Para 63 は元のまま） ----

    doc.save(str(INPUT))
    print(f"saved: {INPUT}")


if __name__ == "__main__":
    main()
