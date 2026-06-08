"""馬強事件 損害賠償請求書 docx v4 生成。

修正指示（260522_馬氏4 続き）：
- セクションタイトル（第１・第２・（１）・（２）等）は左マージンから開始
- 本文は left_indent=1字 + first_line_indent=1字
  → 1行目=2字下げ、2行目以降=1字下げ
- 本文先頭の全角スペースは不要（first_line_indent で字下げを実現）
- 表は left_indent=1字（タイトルより右）
- 箇条書きア〜カは hanging indent（既存維持）

学習点（v3 ユーザー画像から）：
- v3 では全段落 left_indent=1字 で揃えてしまい、タイトルと本文2行目が
  同じ位置になっていた。
- 正解：タイトル left_indent=0、本文 left_indent=1字 + fi=1字、
  本文の見た目は「1行目2字下げ・2行目以降1字下げ」。
"""

import os
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260522_損害賠償請求書（馬様）_3.docx"
)
DST = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260522_損害賠償請求書（馬様）_4.docx"
)

CHAR_EMU = 133350  # ≒1文字（10.5pt）


HEADING_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十０-９]"),
    re.compile(r"^（[０-９0-9]+）"),
    re.compile(r"^【.*】"),
]


def is_heading(text):
    t = text.strip()
    if not t:
        return False
    return any(p.search(t) for p in HEADING_PATTERNS)


def is_bullet_a_kou(text):
    t = text.strip()
    if not t:
        return False
    if t[0] in "アイウエオカキクケコ" and len(t) >= 2 and t[1] in "　 ":
        return True
    return False


def rep(para, text):
    if para.runs:
        first = para.runs[0]
        first.text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    # ヘッダーブロック index：
    # 0: 日付
    # 1-5: 宛先
    # 6-12: 差出人
    # 13: 空行
    # 14: タイトル「損害賠償請求書」
    # 15: 空行
    paras = list(doc.paragraphs)
    HEADER_END = 16  # 16 以降が本文

    for i, p in enumerate(paras):
        pf = p.paragraph_format
        text = p.text

        if i < HEADER_END:
            # ヘッダー要素はそのまま（v3で正しく設定済）
            continue

        if text.strip() == "以上":
            # 「以上」は RIGHT 揃え、字下げなし
            pf.left_indent = None
            pf.first_line_indent = None
            continue

        if is_heading(text):
            # セクションタイトル：左マージンから開始
            pf.left_indent = None
            pf.first_line_indent = None
            # 先頭の全角スペースがあれば削除
            if text.startswith("　") or text.startswith(" "):
                rep(p, text.lstrip("　 "))
            continue

        if is_bullet_a_kou(text):
            # 箇条書き：hanging indent 維持（v3 設定）
            pf.left_indent = Emu(CHAR_EMU * 3)
            pf.first_line_indent = Emu(-CHAR_EMU * 2)
            continue

        # 本文段落：left_indent=1字、first_line_indent=1字
        # 結果：1行目=2字下げ、2行目以降=1字下げ
        # 先頭の全角スペース1個（v3で字下げに使ってたもの）を削除
        if text.startswith("　"):
            new_text = text[1:]
            rep(p, new_text)
        pf.left_indent = Emu(CHAR_EMU)
        pf.first_line_indent = Emu(CHAR_EMU)

    # ---- 表の left_indent は v3 から維持（210 twips ≒ 1字） ----
    # 念のため再設定
    for t in doc.tables:
        tblPr = t._element.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            t._element.insert(0, tblPr)
        existing = tblPr.find(qn("w:tblInd"))
        if existing is not None:
            tblPr.remove(existing)
        tblInd = OxmlElement("w:tblInd")
        tblInd.set(qn("w:w"), "210")
        tblInd.set(qn("w:type"), "dxa")
        tblPr.append(tblInd)

    doc.save(DST)
    print("saved:", DST)


if __name__ == "__main__":
    main()
