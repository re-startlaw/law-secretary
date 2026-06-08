"""馬強さん委任契約書（日中併記）を作成するスクリプト。
ファムさん契約書をテンプレートとしてコピーし、表内テキストを馬さん案件用に置換する。
"""
import shutil
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn

SRC = Path(
    "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com"
    "/マイドライブ/共有用/01_事件記録/ふ_ファム・ティ・フォン/委任関係/260417_委任契約書（ファム様）.docx"
)
DST_DIR = Path(
    "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com"
    "/マイドライブ/共有用/01_事件記録/ま_馬さん/委任関係"
)
DST = DST_DIR / "260507_委任契約書（馬様）.docx"

# --- コンテンツ定義 -------------------------------------------------

TITLE_JP = "委 任 契 約 書"
TITLE_CN = "委 任 合 同 书"

PREAMBLE_JP = "依頼者甲と受任者弁護士乙は、以下のとおり委任契約を締結する。"
PREAMBLE_CN = "委托人甲与受托人律师乙、按以下内容缔结委任合同。"

ART1_JP_TITLE = "第1条（事件等の表示と受任の範囲）"
ART1_JP_BODY = [
    "甲は乙に対し、甲の長女であるアンジェリーナ・マ氏（以下「対象者」という。）に関する、K International School Tokyoとの間の下記事件または法律事務（以下「本件事件」という）の処理を委任し、乙はこれを受任した。",
    "・対象者の復学、損害賠償、謝罪請求等に関する交渉",
]
ART1_CN_TITLE = "第1条（事件等表示及受任范围）"
ART1_CN_BODY = [
    "甲就甲之长女Angelina Ma女士（以下称「对象者」）与K International School Tokyo之间的下列事件或法律事务（以下称「本案件」）委任乙方处理，乙方并就此受任。",
    "・对象者复学、损害赔偿、谢罪请求等事项之交涉",
]

ART2_JP_TITLE = "第2条（報酬）"
ART2_JP_BODY = [
    "１　本件事件に関する着手金は以下のとおりとし、甲は乙の業務着手時にこれを支払う。（金額はいずれも税別）",
    "　着手金：金300,000円",
    "",
    "２　本件事件に関する報酬金は以下のとおりとし、甲は乙による本件事件の終了時に、乙に対して報酬金を支払う。（金額はいずれも税別）",
    "　報酬金：経済的利益の１５％",
    "　※経済的利益とは、相手方から受領する金銭、その他甲が本件事件の解決により金銭的に評価し得る利益の額をいう。",
    "",
    "３　本件事件が訴訟手続その他法的手続に移行する場合の弁護士費用については、甲乙協議のうえ、別途定めるものとする。",
    "",
    "４　甲は、乙が本件事件を処理するに際して必要とする印紙、郵券、交通費、文書取寄費、通信費、複写費、その他の実費を負担し、別途甲乙が協議して定める時期に乙に支払う。",
]
ART2_CN_TITLE = "第2条（报酬）"
ART2_CN_BODY = [
    "1.　本案件之委托费用如下，甲方应于乙方业务着手时支付。（金额均不含税）",
    "　委托费用：300,000日元",
    "",
    "2.　本案件之报酬金如下，本案件终结时，甲方应支付乙方下列报酬金。（金额均不含税）",
    "　报酬金：经济利益之15%",
    "　※「经济利益」系指甲方自相对人处取得之金钱，及甲方因本案件解决而可金钱评价之利益总额。",
    "",
    "3.　本案件移行至诉讼程序或其他法律程序时，律师费用由甲乙双方协议另行约定。",
    "",
    "4.　甲方应负担乙方处理本案件所需之印花费、邮票费、交通费、文书寄送费、通讯费、复印费及其他实际支出费用，并于甲乙协议另定之日期向乙方支付。",
]

ART3_JP_TITLE = "第3条（秘密保持）"
ART3_JP_BODY = [
    "乙は、本契約期間内及び本契約終了後においても、本件事件、甲から相談を受けた事項、その他甲の業務に関し職務上知り得た事項について、甲の秘密を厳守する。",
]
ART3_CN_TITLE = "第3条（保密义务）"
ART3_CN_BODY = [
    "乙方于本合同期间及本合同终止后，就本案件、甲方咨询事项及其他于职务上知悉之关于甲方事务之情报，应严格保守甲方之秘密。",
]

ART4_JP_TITLE = "第4条（中途精算）"
ART4_JP_BODY = [
    "本契約にもとづく本件事件の処理が、解任、辞任または継続不能により終了したときは、乙の処理の程度に応じて精算をおこなうこととし、処理の程度についての甲及び乙の協議結果にもとづき、報酬の全部もしくは一部の返還または支払いを行うものとする。",
]
ART4_CN_TITLE = "第4条（中途结算）"
ART4_CN_BODY = [
    "基于本合同之本案件处理，因解任、辞任或不能继续之事由而终止时，依乙方处理之程度进行结算，并基于甲乙双方就处理程度之协议结果，进行报酬全部或一部分之返还或支付。",
]

ART5_JP_TITLE = "第5条（準拠法）"
ART5_JP_BODY = ["乙による本件事件の処理は、日本法に基づき行うものとする。"]
ART5_CN_TITLE = "第5条（准据法）"
ART5_CN_BODY = ["乙方对本案件之处理，依日本法律为之。"]

ART6_JP_TITLE = "第6条（正文）"
ART6_JP_BODY = [
    "本契約書の正文は日本語とし、中国語訳は参考として付したものである。日本語の条項と中国語訳との間に差異がある場合には、日本語の条項に従って解釈するものとする。",
]
ART6_CN_TITLE = "第6条（正本）"
ART6_CN_BODY = [
    "本合同书之正本以日语为准，中文翻译仅供参考。日语条款与中文翻译之间存在差异时，应依日语条款解释。",
]

CLOSING_JP = "甲及び乙は、甲が乙の説明に基づき本契約の合意内容を十分理解したことを相互に確認し、その成立を証するため本契約書を2通作成し、相互に保管するものとする。"
CLOSING_CN = "甲乙双方互相确认，甲方已基于乙方之说明充分理解本合同之合意内容，为证明本合同之成立，作成本合同书一式两份，由甲乙双方各持一份保管。"

DATE_JP = "2026年5月7日"
DATE_CN = "2026年5月7日"

SIG_JP = [
    "甲（依頼者）",
    "住所：東京都文京区向丘1-3-2",
    "氏名：馬　強　　　　　　　　　　㊞",
    "",
    "",
    "乙（受任者）",
    "〒170-6012",
    "東京都豊島区東池袋3丁目1−1 サンシャイン60 12階",
    "弁護士法人Re-Start法律事務所",
    "弁護士　米谷　尚起　　　　　　　　㊞",
]
SIG_CN = [
    "甲（委托人）",
    "地址：东京都文京区向丘1-3-2",
    "姓名：马　强　　　　　　　　　　（盖章）",
    "",
    "",
    "乙（受托人 律师）",
    "〒170-6012",
    "东京都豊岛区东池袋3丁目1-1 阳光60 12楼",
    "律师法人Re-Start法律事务所",
    "律师：米谷　尚起　　　　　　　　（盖章）",
]


# --- 操作ヘルパー -----------------------------------------------------

def clear_cell(cell):
    """セル内の段落をすべて削除（最低1段落は残す）。"""
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)
    # 空状態だとWordがセルを壊すので新しい段落を1つ追加
    cell.add_paragraph("")


def set_cell_content(cell, lines, title=None, font_name="ＭＳ 明朝", font_size_pt=None):
    """セルにタイトル（任意・太字）＋本文行を入れる。

    lines: list[str] — 各要素が1段落。空文字列は空段落。
    """
    from docx.shared import Pt

    clear_cell(cell)
    # 最初の段落を再利用
    paragraphs = cell.paragraphs
    target = paragraphs[0]
    target.text = ""

    blocks = []
    if title is not None:
        blocks.append((title, True))
    for line in lines:
        blocks.append((line, False))

    for i, (text, bold) in enumerate(blocks):
        if i == 0:
            p = target
        else:
            p = cell.add_paragraph()
        run = p.add_run(text)
        run.bold = bold if bold else None
        if font_size_pt is not None:
            run.font.size = Pt(font_size_pt)
        # 日中混在対応：東アジア用フォント指定
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)


# --- メイン処理 -------------------------------------------------------

def main():
    DST_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, DST)
    print(f"copied → {DST}")

    doc = Document(str(DST))

    # Table 0: タイトル（1行2列）
    t0 = doc.tables[0]
    set_cell_content(t0.rows[0].cells[0], [TITLE_JP])
    set_cell_content(t0.rows[0].cells[1], [TITLE_CN])

    # Table 1: 本文（8行2列）
    t1 = doc.tables[1]
    rows = t1.rows
    # R0: 前文
    set_cell_content(rows[0].cells[0], [PREAMBLE_JP])
    set_cell_content(rows[0].cells[1], [PREAMBLE_CN])
    # R1: 第1条
    set_cell_content(rows[1].cells[0], ART1_JP_BODY, title=ART1_JP_TITLE)
    set_cell_content(rows[1].cells[1], ART1_CN_BODY, title=ART1_CN_TITLE)
    # R2: 第2条
    set_cell_content(rows[2].cells[0], ART2_JP_BODY, title=ART2_JP_TITLE)
    set_cell_content(rows[2].cells[1], ART2_CN_BODY, title=ART2_CN_TITLE)
    # R3: 第3条
    set_cell_content(rows[3].cells[0], ART3_JP_BODY, title=ART3_JP_TITLE)
    set_cell_content(rows[3].cells[1], ART3_CN_BODY, title=ART3_CN_TITLE)
    # R4: 第4条
    set_cell_content(rows[4].cells[0], ART4_JP_BODY, title=ART4_JP_TITLE)
    set_cell_content(rows[4].cells[1], ART4_CN_BODY, title=ART4_CN_TITLE)
    # R5: 第5条
    set_cell_content(rows[5].cells[0], ART5_JP_BODY, title=ART5_JP_TITLE)
    set_cell_content(rows[5].cells[1], ART5_CN_BODY, title=ART5_CN_TITLE)
    # R6: 第6条
    set_cell_content(rows[6].cells[0], ART6_JP_BODY, title=ART6_JP_TITLE)
    set_cell_content(rows[6].cells[1], ART6_CN_BODY, title=ART6_CN_TITLE)
    # R7: 末尾
    set_cell_content(rows[7].cells[0], [CLOSING_JP])
    set_cell_content(rows[7].cells[1], [CLOSING_CN])

    # Table 2: 日付＋署名
    t2 = doc.tables[2]
    set_cell_content(t2.rows[0].cells[0], [DATE_JP])
    set_cell_content(t2.rows[0].cells[1], [DATE_CN])
    set_cell_content(t2.rows[1].cells[0], SIG_JP)
    set_cell_content(t2.rows[1].cells[1], SIG_CN)

    doc.save(str(DST))
    print(f"saved → {DST}")


if __name__ == "__main__":
    main()
