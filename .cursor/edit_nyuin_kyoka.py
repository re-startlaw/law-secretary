# -*- coding: utf-8 -*-
"""
260511入院許可申請.docx をコピーして _2 を作成し、
第２　申請の理由 ブロックと附属書類欄のみを書き換える。
レイアウト（既存パラグラフのスタイル）は維持する。
"""
import shutil
from copy import deepcopy
from pathlib import Path
from docx import Document

SRC = Path(
    "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/"
    ".shortcut-targets-by-id/1pEJgbHuaj4L4ovFEISOuq-0LeHGInxYl/"
    "大中 忠生 : 詐欺未遂/ワードファイル/260511入院許可申請.docx"
)
DST = SRC.with_name("260511入院許可申請_2.docx")

# 新しい第２本文（段落ごとにリスト化）
NEW_REASON_PARAS = [
    "１　大中氏の本件における保釈条件においては、原則として住居を２日以上離れてはならない旨が定められている。",
    "２　大中氏は、令和８年４月頃から、目で物を見ることが困難になってきたことから、京都市内の眼科である「にしじま眼科」（資料１）を受診したところ、手術を要するとの所見を受け、同院から紹介を受けて、大阪医科薬科大学病院眼科（資料２）を受診することとなった。",
    "３　大中氏は、大阪医科薬科大学病院眼科において、令和８年４月２８日に初診を受け、矯正視力検査、精密眼底検査、眼底三次元画像解析、角膜曲率半径計測、細隙灯顕微鏡検査等の精密な検査を受けた（資料３）。続いて、同年５月１日にも再診を受け、精密眼底検査、眼底三次元画像解析等の追加検査が行われた（資料４）。これらの診療を受けたことは外来診療費領収書からも明らかである（資料５）。",
    "４　その後、同月８日には視野検査が、同月１２日にはＩＯＬ検査、手術説明及び入院申込みの手続が予定されている（資料６）。",
    "５　以上の検査の結果、大中氏は、左眼について手術を受ける必要があると診断され、同病院から、入院日を令和８年５月２５日（月）、手術日を同月２６日（火）、退院日を同月３１日（日）とする旨の入院案内を受けた（資料７）。",
    "６　もっとも、医師からは、術式の性質上、手術後は、眼内にガスを留置して網膜を視神経側に押し付ける必要があり、患者は１週間程度の間うつ伏せの姿勢を保ったまま安静療養しなければならない上、術後の経過によっては入院期間が延長され、おおむね２週間程度の入院療養を要する可能性もある旨の説明を受けている。",
    "７　大中氏が住居を離れるのは、専ら上記手術及びこれに伴う入院療養のためであり、罪証を隠滅し、又は逃亡することを意図するものでは全くない。手術直後から、大中氏は、上記のとおり安静療養に専念せざるを得ない状態となるのであるから、関係者と接触する余地もない。",
    "８　また、視覚は、大中氏の日常生活の基盤であるのみならず、ランゲート社及び関連会社の代表取締役としての職務遂行、並びに被告人として裁判記録の確認及び弁護人との打合せに臨むなど防御権の実質的行使のためにも、これを維持することが必要不可欠な身体機能である。本件手術及び入院療養を実施することは、視覚機能を維持し、本件裁判を継続して受けるためにも必要なものといえる。",
    "９　よって、職権により、上記第１の趣旨のとおり許可されたく、本申請に及ぶ次第である。",
]

# 新しい附属書類欄
NEW_APPENDIX_PARAS = [
    "資料１　にしじま眼科 医院案内（同院ホームページ）",
    "資料２　大阪医科薬科大学病院 診察券",
    "資料３　大阪医科薬科大学病院 外来診療明細書（令和８年４月２８日付）",
    "資料４　大阪医科薬科大学病院 外来診療明細書（令和８年５月１日付）",
    "資料５　大阪医科薬科大学病院 外来診療費領収書（令和８年４月２８日付・同年５月１日付）",
    "資料６　大阪医科薬科大学病院 予約票",
    "資料７　眼科手術（入院）を受ける患者様へ（大阪医科薬科大学病院作成）",
]


def find_index(paragraphs, predicate):
    for i, p in enumerate(paragraphs):
        if predicate(p.text):
            return i
    raise RuntimeError("anchor paragraph not found")


def main():
    if DST.exists():
        # _3, _4... に退避
        n = 3
        while DST.with_name(f"260511入院許可申請_{n}.docx").exists():
            n += 1
        new_dst = DST.with_name(f"260511入院許可申請_{n}.docx")
        print(f"既存の _2 があるため、新規コピー先を {new_dst.name} とする")
        target = new_dst
    else:
        target = DST
    shutil.copy2(SRC, target)
    print(f"コピー作成: {target}")

    doc = Document(target)
    paragraphs = doc.paragraphs

    idx_dai2 = find_index(paragraphs, lambda t: t.strip().startswith("第２"))
    idx_appendix = find_index(paragraphs, lambda t: t.strip() == "附属書類")
    idx_ijou = find_index(paragraphs, lambda t: t.strip() == "以上")

    print(f"第２ index={idx_dai2}, 附属書類 index={idx_appendix}, 以上 index={idx_ijou}")

    # --- 第２　申請の理由 ブロックの差し替え -----------------------------
    # 第２ 行と 附属書類 行の間の段落を全削除し、新規段落を挿入する。
    # スタイル保持のため、第２ 段落の deepcopy をテンプレートに用いる。
    anchor_dai2 = paragraphs[idx_dai2]
    style_template = anchor_dai2._p  # XMLエレメント
    parent = style_template.getparent()

    # 削除対象: idx_dai2+1 から idx_appendix-1 まで
    to_delete = paragraphs[idx_dai2 + 1: idx_appendix]
    appendix_p = paragraphs[idx_appendix]._p

    for p in to_delete:
        p._p.getparent().remove(p._p)

    # 新パラグラフを 附属書類 の直前に挿入
    # 先頭に空行 1 つを残す
    def insert_blank_before(ref_p):
        new_p = deepcopy(style_template)
        # テキスト消去
        for r in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            new_p.remove(r)
        ref_p.addprevious(new_p)
        return new_p

    def insert_text_before(ref_p, text):
        new_p = deepcopy(style_template)
        # 既存のすべての run を削除
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        for r in list(new_p.findall(f'{ns}r')):
            new_p.remove(r)
        # 新規 run を作成
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        run = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        run.append(t)
        new_p.append(run)
        ref_p.addprevious(new_p)
        return new_p

    # 1) 先頭の空行
    insert_blank_before(appendix_p)
    # 2) 本文
    for text in NEW_REASON_PARAS:
        insert_text_before(appendix_p, text)
        # 各項目の後にひと呼吸の空行
        insert_blank_before(appendix_p)

    # --- 附属書類 欄の差し替え -------------------------------------------
    # 附属書類 と 以上 の間を全削除して新規に。
    # 再取得（インデックス変化）
    paragraphs = doc.paragraphs
    idx_appendix = find_index(paragraphs, lambda t: t.strip() == "附属書類")
    idx_ijou = find_index(paragraphs, lambda t: t.strip() == "以上")
    print(f"差替後の 附属書類 index={idx_appendix}, 以上 index={idx_ijou}")

    appendix_anchor = paragraphs[idx_appendix]._p
    ijou_anchor = paragraphs[idx_ijou]._p

    # 削除
    to_delete = paragraphs[idx_appendix + 1: idx_ijou]
    for p in to_delete:
        p._p.getparent().remove(p._p)

    # 先頭の空行
    insert_blank_before(ijou_anchor)
    for text in NEW_APPENDIX_PARAS:
        insert_text_before(ijou_anchor, text)
    insert_blank_before(ijou_anchor)

    doc.save(target)
    print(f"保存完了: {target}")
    return target


if __name__ == "__main__":
    p = main()
    print(f"FINAL_PATH={p}")
