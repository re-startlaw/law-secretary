"""馬さん案件・受任通知 中国語版_2 を作成。

仕様:
- ベース: 既存 260512電子内容証明案_中国語版.docx を _2 にコピー済み
- _4 の最終形（修正履歴受諾後）に合わせて中国語本文を反映
- 修正履歴は付けない（クリーン中国語版）
- 指示6（G10学費請求事実関係の加筆）は反映しない（_4 でも不適用）
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
NSMAP = {"w": W_NS}

INPUT = Path(
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/"
    "共有用/01_事件記録/ま_馬さん/Wordファイル/260512電子内容証明案_中国語版_2.docx"
)


def extract_first_rpr(p_elem: etree._Element) -> etree._Element | None:
    return p_elem.find(f".//{W}r/{W}rPr", namespaces=NSMAP)


def set_paragraph_text(p_elem: etree._Element, new_text: str) -> None:
    base_rpr = extract_first_rpr(p_elem)
    for child in list(p_elem):
        if child.tag == f"{W}pPr":
            continue
        p_elem.remove(child)
    r = etree.SubElement(p_elem, f"{W}r")
    if base_rpr is not None:
        r.insert(0, copy.deepcopy(base_rpr))
    t = etree.SubElement(r, f"{W}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = new_text


def delete_paragraph_element(p_elem: etree._Element) -> None:
    p_elem.getparent().remove(p_elem)


def insert_paragraph_after(after_p_elem: etree._Element, text: str) -> etree._Element:
    base_pPr = after_p_elem.find(f"{W}pPr", namespaces=NSMAP)
    base_rpr = extract_first_rpr(after_p_elem)
    new_p = etree.Element(f"{W}p")
    if base_pPr is not None:
        new_p.append(copy.deepcopy(base_pPr))
    r = etree.SubElement(new_p, f"{W}r")
    if base_rpr is not None:
        r.insert(0, copy.deepcopy(base_rpr))
    t = etree.SubElement(r, f"{W}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    after_p_elem.addnext(new_p)
    return new_p


def main() -> None:
    doc = Document(str(INPUT))
    paragraphs = doc.paragraphs

    # ---- 第２(1) 見出し ----
    p24 = paragraphs[24]
    assert "抓拽Angelina女士衣领将其拖移之行为" in (p24.text or "")
    set_paragraph_text(
        p24._element,
        "（１）　小牧女士对Angelina女士衣领附近施以身体接触并使其移动、于跑道旁将其停止之行为",
    )

    # ---- 第２(1) 本文：時間・暴力性 ----
    p25 = paragraphs[25]
    assert p25.text.startswith("　2026年（令和8年）4月13日上午9时18分左右")
    set_paragraph_text(
        p25._element,
        (
            "　2026年（令和8年）4月13日上午9时20分左右至9时30分左右之间，于贵校体育馆内，"
            "小牧女士从背后对Angelina女士衣领附近施以身体接触并将其抓住，使Angelina女士之"
            "身体被带动移动，于跑道旁将其停止，同时大声叱责。该事实，根据Angelina女士同日"
            "书面陈述、本方所持之事案当日影像等，足以认定。"
        ),
    )

    # ---- 第３ 数学授業 ----
    p34 = paragraphs[34]
    assert p34.text.startswith("　Angelina女士承认于2026年4月13日第一节算数课时间内")
    set_paragraph_text(
        p34._element,
        (
            "　Angelina女士承认于2026年4月13日第一节数学课时间（上午8时40分至9时35分）内，"
            "曾前往体育馆。此乃Angelina女士在较早完成自身课题之后，为观看长男George先生"
            "约4分钟之越野跑（cross country）比赛，而前往体育馆。该时间，前半为教师讲课时间，"
            "后半为各学生进行课题之自习时间；且于该自习时间内学生暂时离开教室，亦系现场教师"
            "所默认。"
        ),
    )

    # ---- 第４(2) 精神状態悪化 ----
    p40 = paragraphs[40]
    assert "Angelina女士抑郁症恶化" in (p40.text or "")
    set_paragraph_text(
        p40._element,
        (
            "（２）　因受贵校不当应对、无法到校、及违背事实地散布「Angelina女士在进行虚假"
            "陈述」之言论，导致Angelina女士精神状态恶化、需进行药物调整及继续接受心理咨询"
            "之状态，由此产生之精神科诊疗费及心理咨询费"
        ),
    )

    # ---- 第４(5) 妻関連損害 → 削除 ----
    p43 = paragraphs[43]
    assert "通知人之妻每日为George先生前往贵校" in (p43.text or "")
    delete_paragraph_element(p43._element)

    # ---- 番号繰上げ・並べ替え ----
    # 削除後の paragraphs 取得：旧(6) 通学定期券 → 新(5)
    p_kippu = next(
        p
        for p in doc.paragraphs
        if "通学定期券之损失" in (p.text or "")
    )
    set_paragraph_text(
        p_kippu._element,
        "（５）　已购入之Angelina女士通学定期券之损失",
    )

    # 旧(7) 弁護士費用 → 末尾(9) に移動するため、まず段落取得＆書換
    p_fee = next(
        p
        for p in doc.paragraphs
        if "通知人为本件法律应对所需之律师费" in (p.text or "")
    )
    # 弁護士費用パラグラフは現状(7)→新(9) として残す（並べ替えは挿入で実現）
    set_paragraph_text(
        p_fee._element,
        "（９）　通知人为本件法律应对所需之律师费及其他费用",
    )

    # 弁護士費用の前に新規 (6)(7)(8) を挿入する。
    # 挿入位置：(5) 通学定期券 の直後 → そのあと弁護士費用が(9)で続く構造
    new_items = [
        "（６）　因Angelina女士长期处于学籍中断状态，导致Angelina女士学业连续性受到毁损所产生之损害",
        "（７）　对Angelina女士考试成绩、GPA及Transcript（成绩单）造成不利影响所产生之损害",
        "（８）　对Angelina女士升学路径造成现实性影响，以及教育机会丧失及未来升学风险所产生之损害",
    ]
    anchor = p_kippu._element
    for txt in new_items:
        anchor = insert_paragraph_after(anchor, txt)

    # ---- 第５(2) 直ちに復学 → 公平・安全な教育環境 ----
    p_fukugaku = next(
        p
        for p in doc.paragraphs
        if (p.text or "").strip() == "（２）　立即允许Angelina女士复学"
    )
    set_paragraph_text(
        p_fukugaku._element,
        "（２）　在确保对Angelina女士公平且安全之教育环境之基础上，实现Angelina女士之复学",
    )

    # ---- 第６(3) 学費相当額 → 1,388,000日元 ----
    p_gakuhi = next(
        p for p in doc.paragraphs if "学费相当额（金○○万日元）" in (p.text or "")
    )
    set_paragraph_text(
        p_gakuhi._element,
        (
            "（３）　然而，Angelina女士及通知人强烈希望继续与贵校之在学合同。本方将学费"
            "相当额（金1,388,000日元。相当于Angelina女士G10年级第一期学费。）保全于"
            "本律师之律师预收金账户，待贵校恢复提供教育役务之同时，立即支付该金额。"
        ),
    )

    # 指示6 は不適用 → G10学費の段落（既存）は元のまま

    doc.save(str(INPUT))
    print(f"saved: {INPUT}")


if __name__ == "__main__":
    main()
