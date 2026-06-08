"""馬強事件 — 寺井勇人弁護士宛回答書を連絡書面テンプレから作成"""

import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── パス ──────────────────────────────────
DRIVE = Path(
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/"
    "マイドライブ/共有用"
)
TEMPLATE = DRIVE / "02_ひな形/00_一般/連絡書面.docx"
OUT_DIR = DRIVE / "01_事件記録/ま_馬強/Wordファイル"
OUT = OUT_DIR / "260526_回答書.docx"

# 既存チェック（連番）
stem, suffix = OUT.stem, OUT.suffix
counter = 2
target = OUT
while target.exists():
    target = OUT_DIR / f"{stem}_{counter}{suffix}"
    counter += 1

shutil.copy2(TEMPLATE, target)

# ── ドキュメント読み込み・クリア ─────────────────────
doc = Document(str(target))
body = doc.element.body
sect_pr = body.find(qn("w:sectPr"))
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

# ── マージン ───────────────────────────────
sec = doc.sections[0]
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(2.54)
sec.right_margin = Cm(2.54)

# ── 定数 ───────────────────────────────
CHAR1 = Emu(133350)
SENDER_INDENT = Emu(3600450)
FONT_NAME = "游明朝"
FONT_SIZE = Pt(10.5)
TITLE_SIZE = Pt(16)


def _set_font(run, size=FONT_SIZE, bold=False):
    run.font.size = size
    run.font.name = FONT_NAME
    run.font.bold = bold
    r = run._element
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        rpr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rpr)
    ea = rpr.find(qn("w:rFonts"))
    if ea is None:
        ea = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(ea)
    ea.set(qn("w:eastAsia"), FONT_NAME)


def add_para(text, align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=None,
             first_indent=None, bold=False, size=FONT_SIZE, space_after=Pt(0),
             space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_after = space_after
    fmt.space_before = space_before
    if left_indent is not None:
        fmt.left_indent = left_indent
    if first_indent is not None:
        fmt.first_line_indent = first_indent
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    return p


def add_body(text, space_after=Pt(0)):
    """本文段落（1字下げ + first_line 1字 = 1行目2字、2行目1字）"""
    return add_para(text, left_indent=CHAR1, first_indent=CHAR1,
                    space_after=space_after)


def add_section_title(text, space_before=Pt(6)):
    """第○ 見出し（インデントなし）"""
    return add_para(text, space_before=space_before)


def add_subsection_title(text, space_before=Pt(3)):
    """（○）小見出し（インデントなし）"""
    return add_para(text, space_before=space_before)


def add_demand_item(text):
    """要求事項の番号付き箇条書き（hanging indent）"""
    return add_para(text, left_indent=Emu(400050), first_indent=Emu(-266700))


# ── ヘッダー部 ──────────────────────────────
add_para("令和８年５月２６日",
         align=WD_ALIGN_PARAGRAPH.RIGHT)

# 宛先
for line in [
    "〒108-0074",
    "東京都港区高輪四丁目23番6号",
    "ハイホーム高輪708",
    "寺井勇人法律事務所",
    "弁護士　寺井　勇人　先生",
]:
    add_para(line)

# 空行
add_para("")

# 差出人（RIGHT 揃え）
for line in [
    "（差出人）",
    "〒170-6012",
    "東京都豊島区東池袋３丁目１番１号",
    "サンシャイン60　12階",
    "弁護士法人Ｒｅ－Ｓｔａｒｔ法律事務所",
    "弁護士　米谷　尚起",
    "TEL：03-6820-3815",
]:
    add_para(line, align=WD_ALIGN_PARAGRAPH.RIGHT)

# 空行
add_para("")

# タイトル
add_para("回答書", align=WD_ALIGN_PARAGRAPH.CENTER,
         bold=True, size=TITLE_SIZE, space_after=Pt(6))

# ── 冒頭 ───────────────────────────────
add_body(
    "貴職が令和８年５月２２日付で"
    "送付された回答書（以下「貴職"
    "回答書」といいます。）を拝受"
    "いたしました。当職は、引き続"
    "き馬強氏（以下「通知人」とい"
    "います。）の代理人として、以"
    "下のとおり回答いたします。"
)

# ── 第1 ───────────────────────────────
add_section_title(
    "第１　回答の基本的立場"
)
add_body(
    "貴職回答書は、事実関係に関"
    "する主張を詳細に展開されて"
    "おりますが、当方が受任通知"
    "において求めた本質的な問題"
    "、すなわち「Angelina氏の登校を拒"
    "絶する法的根拠」及び「George氏"
    "の看護者の校内立入りを禁止"
    "する法的根拠」に対する回答"
    "が一切含まれておりません。"
)
add_body(
    "本書面においては、事実関係"
    "の細部には立ち入りません。"
    "以下では、万が一、貴職の主"
    "張する事実関係を全て前提と"
    "したとしても、登校拒絶及び"
    "看護者の排除が法的に正当化"
    "されないことを述べます。"
)

# ── 第2 ───────────────────────────────
add_section_title(
    "第２　登校拒絶に法的根拠"
    "がないこと"
)

add_subsection_title(
    "（１）在学契約上の教育役"
    "務提供義務"
)
add_body(
    "在学契約は有償双務契約で"
    "あり（最判平成18年11月27日）"
    "、学費納付義務と教育役務"
    "提供義務は対価関係に立ち"
    "ます。貴校は通知人から学"
    "費を受領しながら、令和８"
    "年４月13日以降、約１か月半"
    "にわたり教育役務の提供を"
    "拒絶しており、これは債務"
    "不履行にほかなりません。"
)

add_subsection_title(
    "（２）「虚偽申告」は履行"
    "拒絶の根拠たり得ないこと"
)
add_body(
    "貴職は、Angelina氏が理事長の行"
    "為について「虚偽の申告」を"
    "し、その撤回を拒んでいるこ"
    "とをもって登校拒絶の理由と"
    "されます。"
)
add_body(
    "しかしながら、万が一、Angelina"
    "氏の申告内容が客観的事実と"
    "異なるとしても、生徒が学校"
    "側の事実認定に同意しないこ"
    "とは、教育役務の提供を拒絶"
    "できる契約上の事由に該当し"
    "ません。在学契約において、"
    "生徒側が学校の見解に無条件"
    "に従う義務は存在せず、事実"
    "関係に争いがあること自体は"
    "、教育を受ける権利とは全く"
    "無関係です。"
)

add_subsection_title(
    "（３）「教育的観点」「安全"
    "確保の観点」は法的根拠に代"
    "わり得ないこと"
)
add_body(
    "貴職は「教育的観点」「学校"
    "の安全確保の観点」を繰り返"
    "し述べられますが、これらは"
    "抽象的な理念にすぎず、登校"
    "拒絶を法的に正当化する具体"
    "的根拠とはなり得ません。"
)
add_body(
    "貴職回答書第４の２（１）に"
    "おいて、「虚偽主張をするこ"
    "とにより物事を解決できると"
    "いう勘違いは絶対に糺さなけ"
    "ればならない」、「このよう"
    "な勘違いを抱えながら登校と"
    "して社会に送り出すことは、"
    "当校の教育方針からできない"
    "」と述べられておりますが、"
    "これは教育上の価値判断を理"
    "由に在学契約上の債務の履行"
    "を拒絶するものであり、法的"
    "に正当化される余地はありま"
    "せん。"
)

# ── 第3 ───────────────────────────────
add_section_title(
    "第３　George氏の看護者の立入"
    "り禁止に法的根拠がないこと"
)
add_body(
    "George氏は年少の児童であり、看"
    "護者の付添いなくして通常の"
    "学校生活を送ることは困難で"
    "す。看護者の校内立入りを禁"
    "止することは、実質的にGeorge氏"
    "の就学環境そのものを害する"
    "措置です。"
)
add_body(
    "貴職は、看護者が「虚偽の加"
    "担証言」をしたことをその理"
    "由とされますが、万が一その"
    "ような事実があったとしても"
    "、Angelina氏に関する事実調査にお"
    "ける証言内容を理由として、"
    "別の生徒であるGeorge氏の看護者"
    "を排除することに合理的関連"
    "性はありません。"
)
add_body(
    "施設管理権は無制限の権限で"
    "はなく、在学契約に基づく教"
    "育環境提供義務の範囲内で行"
    "使されるべきものです。本件"
    "措置はその逸脱にほかなりま"
    "せん。"
)

# ── 第4 ───────────────────────────────
add_section_title(
    "第４　貴職の要求について"
)
add_body(
    "貴職回答書第６において求め"
    "られた以下の各事項には、い"
    "ずれも応じません。"
)

add_subsection_title(
    "（１）映像データの提出につ"
    "いて"
)
add_body(
    "通知人が保有する映像データ"
    "は、通知人側の証拠です。貴"
    "職が一方的に設定した私的調"
    "査手続において提出すべき義"
    "務はなく、当該証拠は司法手"
    "続において適切に提出いたし"
    "ます。"
)

add_subsection_title(
    "（２）貴校アドバイザー立会"
    "いの現場検証について"
)
add_body(
    "貴校が独自に選任した人物の"
    "立会いによる検証は、中立性"
    "・客観性を欠く私的調査であ"
    "り、これに応じる法的義務は"
    "存在しません。"
)

add_subsection_title(
    "（３）K2保護者及び職員への事"
    "実確認の同意について"
)
add_body(
    "Angelina氏に関する本件紛争の内容"
    "を約50名の保護者に開示するこ"
    "とは、Angelina氏のプライバシーを"
    "著しく侵害するものであり、"
    "同意しません。"
)

# ── 第5 ───────────────────────────────
add_section_title(
    "第５　偶計業務妨害罪に関す"
    "る貴職の主張について"
)
add_body(
    "貴職は、偶計業務妨害罪（刑"
    "法第233条）による刑事告訴を示"
    "唆されております。同罪は「"
    "虚偽の風説を流布し、又は偶"
    "計を用いて、人の業務を妨害"
    "」した場合に成立するもので"
    "す。"
)
add_body(
    "通知人は、自身の子が学校内"
    "で受けた扱いについて、学校"
    "に対して事実を主張し、弁護"
    "士を通じて書面により交渉し"
    "ているにすぎません。これは"
    "正当な権利行使であり、「虚"
    "偽の風説の流布」にも「偶計"
    "」にも該当しません。また、"
    "当方の主張は貴校に対しての"
    "み行われており、第三者に流"
    "布された事実もないことから"
    "、「流布」の要件も充たしま"
    "せん。"
)

# ── 第6 ───────────────────────────────
add_section_title(
    "第６　通知人の要求"
)
add_body(
    "取り急ぎ、以下の各事項につ"
    "いて、本書面到達後７日以内"
    "に書面にてご回答ください。"
)

add_demand_item(
    "１　Angelina氏に対する登校拒絶を"
    "直ちに解除し、安全な教育環"
    "境を確保した上で復学を認め"
    "ること"
)
add_demand_item(
    "２　George氏の看護者による校内"
    "への立入りを直ちに認めるこ"
    "と"
)
add_demand_item(
    "３　上記各措置の法的根拠を"
    "、具体的な契約条項又は法令"
    "の条文を示して明らかにする"
    "こと"
)

add_body(
    "なお、上記期限内に貴校が対"
    "応されない場合、通知人は、"
    "生徒たる地位保全の仮処分（"
    "民事保全法第23条第２項）の申"
    "立て、教育役務の提供を求め"
    "る本案訴訟、及び貴校の債務"
    "不履行並びに不法行為に基づ"
    "く損害賠償請求訴訟の提起を"
    "含む法的手続に着手いたしま"
    "す。"
)
add_body(
    "また、受任通知において通知"
    "したとおり、学費相当額は当"
    "職の預り金口座において保全"
    "しており、貴校による教育役"
    "務提供の再開と引換えに、直"
    "ちに同額をお支払いする用意"
    "があることを改めて申し添え"
    "ます。"
)

# 以上
add_para("以上", align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(12))

# ── 保存 ───────────────────────────────
doc.save(str(target))
print(f"saved: {target}")
