"""馬強事件 続編内容証明（損害賠償請求書）docx 生成 v3。

naiyou-shoumei スキルに従い、雛形 02_ひな形/00_一般/内容証明.docx を
shutil.copy2 でコピーし、python-docx で run を温存しつつ段落を書き換え
る方式に変更（v2の新規Document生成方式は雛形レイアウトを失うため廃止）。
"""

import os
import shutil
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/02_ひな形/"
    "00_一般/内容証明.docx"
)
DST = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260522電子内容証明案_損害賠償請求.docx"
)


def rep(para, text):
    """段落のrunを温存しつつ本文だけ差し替える。"""
    if para.runs:
        first = para.runs[0]
        first.text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def insert_para_after(anchor_para, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """anchor_para の直後に新規段落を挿入し、その段落を返す。"""
    new_p = deepcopy(anchor_para._element)
    # 既存のrunを掃除
    for child in list(new_p):
        if child.tag.endswith("}r"):
            new_p.remove(child)
        elif child.tag.endswith("}pPr"):
            # alignment用にppr残す
            pass
    anchor_para._element.addnext(new_p)
    # rebuild Paragraph from element via doc.paragraphs scan
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, anchor_para._parent)
    p.alignment = alignment
    p.add_run(text)
    return p


def insert_before(anchor_para, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """anchor_para の直前に新規段落を挿入。"""
    new_p_el = deepcopy(anchor_para._element)
    for child in list(new_p_el):
        if child.tag.endswith("}r"):
            new_p_el.remove(child)
    anchor_para._element.addprevious(new_p_el)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p_el, anchor_para._parent)
    p.alignment = alignment
    p.add_run(text)
    return p


def remove_para(para):
    para._element.getparent().remove(para._element)


def main():
    os.makedirs(os.path.dirname(DST), exist_ok=True)
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    paras = list(doc.paragraphs)

    # ---- ヘッダー部分の書き換え ----
    rep(paras[0], "２０２６年５月◯日")
    rep(paras[1], "〒１０８－００７４")
    rep(paras[2], "東京都港区高輪四丁目２３番６号")
    rep(paras[3], "ハイホーム高輪７０８")
    rep(paras[4], "寺井勇人法律事務所")
    # 雛形にない「弁護士 寺井 勇人 先生」段落を Para 4 の後に追加
    insert_para_after(paras[4], "弁護士　寺井　勇人　先生")

    # Para 5 (差出人) はそのまま
    # Para 6-10 (差出人住所・事務所・米谷尚起) はそのまま
    # Para 11 空行 → TEL を入れる
    rep(paras[11], "TEL：０３－６８２０－３８１５")

    # Para 12-16 (共同受任弁護士欄) は削除
    for p in [paras[12], paras[13], paras[14], paras[15], paras[16]]:
        remove_para(p)

    # Para 17 (空行) はそのまま残す
    # Para 18 タイトル (CENTER)
    rep(paras[18], "損害賠償請求及び要求事項に関する御連絡")

    # Para 19 (空行) はそのまま
    # Para 20以降は本文。雛形の本文段落（20-34）を全削除し、
    # Para 35 (草々, RIGHT) を「以上」に書き換えて温存。
    for p in paras[20:35]:
        remove_para(p)

    # 末尾「草々」段落を「以上」に書き換え（alignment は雛形のRIGHTを温存）
    rep(paras[35], "以上")

    # ---- 本文を草々(以上)段落の前に挿入 ----
    anchor = paras[35]  # 「以上」段落

    body = [
        ("　当職は、貴職が代理人として対応されているK. International School "
         "Tokyo（以下「貴校」といいます。）にG9学年として在籍するAngelina "
         "Feili Ma氏（以下「Angelina氏」といいます。）の父であり、貴校との"
         "間で在学契約を締結している馬強氏（以下「通知人」といいます。）の"
         "代理人として、貴職に対し、令和８年５月１２日付「受任通知及び御連絡」"
         "（以下「前回通知」といいます。）に引き続き、以下のとおり通知人の"
         "貴校及び貴校理事会議長小牧孝子氏（以下「小牧氏」といいます。）に"
         "対する損害賠償請求並びに要求事項について、御連絡申し上げます。",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),

        ("第１　前回通知の整理", WD_ALIGN_PARAGRAPH.LEFT),
        ("　前回通知においては、貴校及び小牧氏による不法行為の概要、現に通知人"
         "及びその家族に生じている損害の項目並びに通知人の貴校に対する要求事項"
         "を御連絡したほか、貴校による教育役務提供義務の履行拒絶に対し、通知人"
         "が同時履行の抗弁権を行使し、Angelina氏のG10学年第一期学費相当額"
         "（金１,３８８,０００円）を当職弁護士預り金口座において保全している"
         "旨を御連絡いたしました。", WD_ALIGN_PARAGRAPH.LEFT),
        ("　本書面は、前回通知においては個別の金額を留保していた損害賠償請求の"
         "具体的内容並びに要求事項に対する御回答期限の設定について、改めて"
         "御連絡するものです。", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),

        ("第２　損害賠償請求", WD_ALIGN_PARAGRAPH.LEFT),
        ("　前記前回通知第２記載の貴校及び小牧氏による不法行為（民法第７０９条、"
         "同法第７１５条又は第７１９条）並びに貴校による在学契約上の教育役務"
         "提供義務の履行拒絶（同法第４１５条）により、通知人及びその家族には、"
         "以下のとおり、金銭に換算可能な損害が現に生じております。通知人は、"
         "貴校及び小牧氏に対し、連帯して、下記合計額の支払を請求いたします。",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),

        ("（１）　実費損害", WD_ALIGN_PARAGRAPH.LEFT),
        ("ア　Angelina氏の代替教育（Crimson Global Academy G9 Term 2）学費",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　金１,２６０,０００円", WD_ALIGN_PARAGRAPH.LEFT),
        ("イ　Angelina氏のG10学年オンラインスクール預託金（１０％ Deposit）",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　金２４４,８００円", WD_ALIGN_PARAGRAPH.LEFT),
        ("ウ　既に貴校に対し支払済みであるにもかかわらず、Angelina氏が登校を"
         "拒絶されている期間に対応するKIST G9学年学費相当額",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　貴校G9学年の年間学費総額金２,８９７,０００円のうち、令和８年"
         "４月１４日から正常復学日までの登校拒絶期間に対応する按分額。最終的な"
         "金額は復学日確定後に確定いたします。", WD_ALIGN_PARAGRAPH.LEFT),
        ("エ　既に購入済みの通学定期券に関する損失", WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　令和８年１月１７日購入の６か月通学定期券金３５,９８０円のうち、"
         "令和８年４月１４日以降の未利用期間相当額", WD_ALIGN_PARAGRAPH.LEFT),
        ("オ　精神科診療費及び心理カウンセリング費用",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　令和８年４月１６日付精神科診療費　金２１,８５０円",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　令和８年４月３０日付精神科診療費　金５７,０５０円",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　令和８年５月１日付心理カウンセリング費用　金１７,０００円",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　令和８年５月８日付心理カウンセリング費用　金１７,０００円",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　以後の精神科及び心理カウンセリング治療費用については、現に継続"
         "発生しており、最終的な金額は今後の領収書により確定いたします。",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("カ　Angelina氏の精神科通院及び心理カウンセリング通院に伴う交通費",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　Angelina氏の通院には保護者１名の同行が必要であり、精神科通院"
         "（１回あたり往復金８４０円、概ね２週から３週に１回）及び心理カウンセ"
         "リング通院（１回あたり往復金８４０円、概ね週１回）に係る交通費を"
         "実費損害として請求いたします。", WD_ALIGN_PARAGRAPH.LEFT),
        ("キ　弁護士費用その他本件への法的対応のために通知人が要した費用",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　　本件は、貴校及び小牧氏の不法行為と相当因果関係を有する弁護士"
         "費用として、上記金額並びに後記精神的損害の合計額の１０％相当額を"
         "請求いたします。", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),

        ("（２）　精神的損害（慰謝料）", WD_ALIGN_PARAGRAPH.LEFT),
        ("　貴校及び小牧氏の不法行為により、Angelina氏及び通知人らは、以下に"
         "列挙する精神的苦痛を受けております。", WD_ALIGN_PARAGRAPH.LEFT),
        ("ア　令和８年４月１３日における小牧氏の有形力行使（襟元付近の身体的"
         "接触を伴う移動及び大声による叱責）により、Angelina氏が受けた精神的"
         "苦痛", WD_ALIGN_PARAGRAPH.LEFT),
        ("イ　約１か月以上にわたる登校拒絶措置の継続により、Angelina氏が"
         "正常な学習機会及び学級内の人間関係を喪失したことに伴う精神的苦痛"
         "並びに不安、心理的圧迫", WD_ALIGN_PARAGRAPH.LEFT),
        ("ウ　貴校及び小牧氏が、Angelina氏の書面陳述記載の事実関係について、"
         "事実に反する主張を行っているかのような取扱いを継続したことにより、"
         "Angelina氏が受けた名誉権、人格権及び校内評価に対する侵害並びに精神"
         "的苦痛", WD_ALIGN_PARAGRAPH.LEFT),
        ("エ　Angelina氏の精神状態の悪化、投薬量の増量、継続的心理カウンセ"
         "リングを要する状態に至った結果、Angelina氏及び通知人らが受けた精神"
         "的苦痛", WD_ALIGN_PARAGRAPH.LEFT),
        ("オ　Angelina氏の学業継続性、試験成績、GPA及びTranscriptへの不利益"
         "並びに将来の大学進学経路への現実的影響に伴う精神的苦痛",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("カ　George氏の看護者に対する校内立入禁止措置の継続により、George氏"
         "の就学環境が害されたこと並びにこれに対応するために通知人らが負担を"
         "強いられたことに伴う精神的苦痛", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),
        ("　以上の精神的苦痛に対する慰謝料として、通知人は、貴校及び小牧氏に"
         "対し、連帯して、金７,０００,０００円の支払を請求いたします。",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),

        ("（３）　合計", WD_ALIGN_PARAGRAPH.LEFT),
        ("　前記（１）アからカまでに掲記の各実費損害、前記（２）の慰謝料金"
         "７,０００,０００円及び前記（１）キの弁護士費用の合計額を、本件請求"
         "金額といたします。なお、前記（１）ウの登校拒絶期間に対応する按分額"
         "及び前記（１）オの精神科及び心理カウンセリング治療費用は、現に継続"
         "発生しており、最終的な金額は復学日及び治療終了日の確定をもって精算"
         "するものといたします。", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),

        ("第３　支払方法及び支払期限", WD_ALIGN_PARAGRAPH.LEFT),
        ("　貴校及び小牧氏は、本書面到達後◯日以内に、前記第２（３）記載の本件"
         "請求金額を、下記当職指定口座へお振込みの方法によりお支払いください。",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　【振込先】", WD_ALIGN_PARAGRAPH.LEFT),
        ("　　金融機関：（追って当職よりご指定申し上げます。）",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　　口座種別・口座番号：（同上）", WD_ALIGN_PARAGRAPH.LEFT),
        ("　　口座名義：弁護士法人Ｒｅ－Ｓｔａｒｔ法律事務所　預り口",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),

        ("第４　要求事項及び御回答期限", WD_ALIGN_PARAGRAPH.LEFT),
        ("　前回通知第５記載の各要求事項について、改めて以下のとおり整理の上、"
         "本書面到達後◯日以内に、貴校の対応方針について御回答くださいますよう"
         "求めます。", WD_ALIGN_PARAGRAPH.LEFT),
        ("（１）　Angelina氏に対する登校拒絶措置を速やかに解除し、Angelina氏"
         "に対する公平かつ安全な教育環境を確保した上で、Angelina氏の正常な復学"
         "を実現すること。", WD_ALIGN_PARAGRAPH.LEFT),
        ("（２）　George氏の看護者による貴校構内への正常な立入り権限を回復"
         "すること。", WD_ALIGN_PARAGRAPH.LEFT),
        ("（３）　本件における貴校の対応に不適切な点があったことを認め、"
         "Angelina氏及び通知人に対し、書面をもって謝罪すること。",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("（４）　Angelina氏の名誉及び人格的評価の回復のため、貴校がこれまで"
         "Angelina氏について事実に反する申告を行っているかのような認識を伝えた"
         "全ての人物に対し、その旨の訂正を行う声明を発出する等の措置を講じる"
         "こと。", WD_ALIGN_PARAGRAPH.LEFT),
        ("（５）　Angelina氏の登校拒絶期間中における学業の遅れを取り戻すため"
         "の補習その他必要な措置を講じること。", WD_ALIGN_PARAGRAPH.LEFT),
        ("（６）　Angelina氏のG10学年の学費の支払方法及び支払時期について、"
         "復学方案の確定と併せて協議すること。", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),

        ("第５　法的措置の予告", WD_ALIGN_PARAGRAPH.LEFT),
        ("　万一、前記第３記載の支払期限内に貴校及び小牧氏から本件請求金額の"
         "全部又は相当部分の支払がない場合、又は前記第４記載の御回答期限内に"
         "貴校から誠実な御回答が得られない場合、当方としては、Angelina氏の"
         "生徒たる地位保全の仮処分命令申立て（民事保全法第２３条第２項）並びに"
         "本案訴訟の提起その他の法的措置を直ちに講じる所存です。",
         WD_ALIGN_PARAGRAPH.LEFT),
        ("　また、前記第３記載の支払期限を経過した本件請求金額については、"
         "支払期限の翌日から支払済みまで、年３％の割合による遅延損害金を併せ"
         "請求いたします。", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),

        ("　本書面の内容を御確認いただき、御不明な点がございましたら、ご遠慮"
         "なくお尋ねくださいますようお願い申し上げます。",
         WD_ALIGN_PARAGRAPH.LEFT),
    ]

    # 「以上」段落の前に順次挿入
    for text, align in body:
        insert_before(anchor, text, align)

    doc.save(DST)
    print("saved:", DST)


if __name__ == "__main__":
    main()
