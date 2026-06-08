#!/usr/bin/env python3
"""Build refined version of 260512上申書 by incorporating 聴取報告書 content."""
import shutil
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from copy import deepcopy

BASE = Path("/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/た_田村正宣/09_ワードファイル")
SRC = BASE / "260512上申書.docx"
DST = BASE / "260512上申書_2.docx"

# Step 1: copy
shutil.copy2(SRC, DST)
print(f"Copied: {SRC.name} -> {DST.name}")

doc = Document(str(DST))

# Sample style/font from existing body paragraphs (line 009-onward)
sample_body_para = doc.paragraphs[9]
sample_run_font = None
if sample_body_para.runs:
    sample_run_font = sample_body_para.runs[0].font

def make_para_match_body(p):
    """Copy default font from sample body paragraph to runs in p."""
    if sample_run_font is None:
        return
    for r in p.runs:
        if sample_run_font.name:
            r.font.name = sample_run_font.name
            # East Asian font requires explicit rFonts override
            from docx.oxml.ns import qn
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), sample_run_font.name)
            rFonts.set(qn('w:ascii'), sample_run_font.name)
            rFonts.set(qn('w:hAnsi'), sample_run_font.name)
        if sample_run_font.size:
            r.font.size = sample_run_font.size

# Strategy: clear all paragraphs, then rebuild.
# We need to keep at least one paragraph in the body, then add via add_paragraph.

body = doc.element.body
# Remove all paragraph children of body, except sectPr at end
sectPr = body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
for child in list(body):
    if child is sectPr:
        continue
    body.remove(child)

def add(text="", *, style=None, align=None):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = doc.styles[style]
        except KeyError:
            pass
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
    make_para_match_body(p)
    return p

# === Build document ===

# Title
p = doc.add_paragraph()
try:
    p.style = doc.styles['タイトル']
except KeyError:
    pass
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("上申書")
make_para_match_body(p)

add()  # blank

# Date
add("令和８年５月２０日", align=WD_ALIGN_PARAGRAPH.RIGHT)

# Court
add("東京地方裁判所　刑事部　御中")

add()  # blank
add()  # blank

# Signer block (right-aligned)
add("上申人（被疑者内縁の妻）", align=WD_ALIGN_PARAGRAPH.RIGHT)
add("氏名　大宮　瑞希　　印", align=WD_ALIGN_PARAGRAPH.RIGHT)

add()  # blank

# Intro
add(
    "私は、本件被疑者である田村正宣（以下「主人」といいます。）の内縁の妻、大宮瑞希と申します。"
)
add(
    "主人とは令和元年４月に出会い、今年で交際７年目を迎えます。"
    "その間に三人の子を授かりました。長女・大宮向日凜（令和３年９月２９日生）、"
    "長男・大宮優篤（令和５年１月５日生）、次男・大宮護（令和８年１月２２日生）の三人です。"
    "婚姻届こそ提出しておりませんが、主人は三人とも認知し、責任をもって育ててくれており、"
    "家族五人で寄り添うように暮らしてまいりました。"
)
add(
    "今般、主人が逮捕されたことに加え、その後の予期せぬ事態に直面し、"
    "私自身の精神状態も限界に近く、また三人の子どもたちの生活を守るために一刻の猶予もない状況にあります。"
    "下記の事情をご賢察いただき、主人との接見、せめて手紙のやりとりだけでもお認めいただきたく、"
    "重ねてお願い申し上げます。"
)

add()  # blank

# Section 1
add("１．報道による自宅の特定と、避難先での限界について")
add(
    "主人の逮捕後、実名や顔写真、さらには自宅住所の町名までが詳細に報道されました。"
    "主人は子どもたちの幼稚園の送り迎えを毎日欠かさず行っており、"
    "地域の方々や園の保護者の方にも顔が知られていたため、一気に噂が広まり、"
    "志木市の自宅には身の危険を感じていられなくなりました。"
)
add(
    "急遽、茨城県の実家へ子ども三人を連れて避難いたしましたが、"
    "私の両親は花屋を営む自営業者で、共働きで朝から晩まで店に出ております。"
    "日中の育児サポートはほとんど得られず、私一人が慣れない実家という環境で、"
    "生後三か月の乳児を含む三人の子どもの世話に追われております。"
    "両親の生活リズムを壊している申し訳なさもあり、いつまでも身を寄せていられる状況ではありません。"
)

add()  # blank

# Section 2
add("２．精神的な苦痛と切迫感について")
add(
    "主人が突然いなくなったショックに加え、連日のように流れる報道やネット上の反応、"
    "そして住み慣れた自宅を追われるという異常な事態に、"
    "私は現在、夜も全く眠れないほど精神的に追い詰められております。"
)
add(
    "三人の子どもたちの前では何とか気を張っていますが、"
    "今後の生活や住居、お金のこと、そして主人の安否を考えると、"
    "不安で押しつぶされそうになり、涙が止まらない時間もあります。"
    "特に生後三か月の次男を抱えながら、自分一人では何も決められない、"
    "どこへ行けばいいのかもわからないという状態です。"
)

add()  # blank

# Section 3
add("３．主人との協議の必要性について")
add(
    "報道により元の自宅へは戻れず、早急に新しい引越し先を見つける必要があります。"
    "しかし、引越し先をどの地域にするか、それに伴う長女・長男の幼稚園の転園手続き、"
    "そして次男の保育園入園の相談など、いずれも主人の同意なしに進めることは不可能です。"
)
add(
    "また、私はこれまで三人の子どもの育児で働いておらず、生活費はすべて主人から受け取っておりました。"
    "手持ちの資金は既に底をつきかけており、引越し費用も必要です。"
    "一、二か月であれば何とかしのげるとしても、それ以上勾留が続けば、"
    "三人の子どもが路頭に迷ってしまいます。"
    "主人の周囲の方々にどのようにサポートをお願いすればよいのか、"
    "今後の生活をどう立て直すべきかについて、主人の口から直接指示を受け、"
    "子どもたちの将来について安心させてほしいと切に願っております。"
)

add()  # blank

# Section 4 (NEW from 聴取報告書 item 3)
add("４．共犯とされる方々との関係について")
add(
    "報道で共犯と取り沙汰されている方々とは、いずれも主人を介して知り合った間柄に過ぎず、"
    "子ども同士を家族ぐるみで遊ばせたことがある程度です。"
    "私自身が個別に連絡先を交換しているわけではなく、"
    "主人を介さない限り、私から連絡を取る手段はそもそもございません。"
)
add(
    "加えて、弁護人の先生から、こうした方々への接触は証拠隠滅を疑われる行為であることを"
    "丁寧にご説明いただき、決して行ってはならないことをよく理解しております。"
    "主人に対しても、面会や手紙の場で、私から強く言い聞かせる所存です。"
)

add()  # blank

# Section 5: 結び
add("５．結び")
add(
    "面会には警察官の方が立ち会われ、手紙も検閲を受けると伺っております。"
    "事件の話をするつもりはございません。"
    "これからの引越し先や、子どもたちの幼稚園のこと、当面の生活費をどう工面するのかという、"
    "家族が生きていくための切実な相談をさせてください。"
)
add(
    "どうか、私と主人との面会、手紙のやりとりをお許しくださいますよう、"
    "心よりお願い申し上げます。"
)

add()  # blank

# 以上
add("以上", align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.save(str(DST))
print(f"Saved: {DST}")

# Verify
print("\n=== Verification ===")
doc2 = Document(str(DST))
for i, p in enumerate(doc2.paragraphs):
    style = p.style.name if p.style else ""
    align = p.alignment
    txt = p.text[:60]
    print(f"[{i:03d}][{style}][align={align}] {txt}")
