"""v3 と v7_2 の条文内容を詳細に比較。"""
from pathlib import Path
from docx import Document
import re

v3_path = Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v3.docx")
v7_2_path = Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v7_2.docx")

def load_doc(path):
    return Document(str(path))

def extract_full_article_text(doc, article_num):
    """条文番号から、その条文全体の本文を取得（項目を含む）。"""
    start_idx = None
    end_idx = None
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        match = re.match(r'^第(\d+)条', text)
        
        if match:
            num = int(match.group(1))
            if num == article_num and start_idx is None:
                start_idx = i
            elif num > article_num and start_idx is not None:
                end_idx = i
                break
    
    if start_idx is None:
        return None
    
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    
    lines = []
    for i in range(start_idx, end_idx):
        lines.append(doc.paragraphs[i].text)
    
    return '\n'.join(lines)

def check_internal_references(doc):
    """「第○条」への内部参照をすべて抽出。"""
    references = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        # 「第○条第△項」「第○条」などのパターンを抽出
        pattern = r'第(\d+)条（第(\d+)項)?'
        matches = re.findall(pattern, text)
        if matches:
            references[i] = {
                'text': text,
                'refs': matches
            }
    return references

# ====== main ======
print("=" * 100)
print("【v3 と v7_2 の条文内容 詳細比較】")
print("=" * 100)

doc_v3 = load_doc(v3_path)
doc_v7_2 = load_doc(v7_2_path)

# v3 と v7_2 の主要条文を比較
compare_articles = [2, 3, 6, 10, 11, 13, 16, 19, 20]

print("\n【主要条文の内容差分】")
for article_num in compare_articles:
    v3_text = extract_full_article_text(doc_v3, article_num)
    v7_2_text = extract_full_article_text(doc_v7_2, article_num)
    
    print(f"\n{'='*100}")
    print(f"【第{article_num}条】")
    print(f"{'='*100}")
    
    if v3_text == v7_2_text:
        print("（v3 と v7_2 で同一）")
    else:
        print(f"\n--- v3 ---\n{v3_text[:500]}")
        print(f"\n--- v7_2 ---\n{v7_2_text[:500]}")
        if len(v7_2_text) > 500:
            print(f"\n[後略... 全長: v7_2={len(v7_2_text)}文字]")

print(f"\n\n{'='*100}")
print("【v7_2 内部参照（第○条）の一覧】")
print(f"{'='*100}")
refs = check_internal_references(doc_v7_2)
all_refs = set()
for para_idx, ref_data in refs.items():
    text = ref_data['text']
    for article, item in ref_data['refs']:
        all_refs.add(int(article))
        print(f"[段落{para_idx:03d}] → 第{article}条 (項={item}) | {text!r}")

print(f"\n合計: {len(all_refs)}条が参照対象")
print(f"参照されている条文: {sorted(all_refs)}")

