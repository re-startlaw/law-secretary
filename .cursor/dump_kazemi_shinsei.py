from docx import Document

path = "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/か_KAZEMI・HOSSEIN/提出書類/妻/申請理由書（サンプル）.docx"

doc = Document(path)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text}")

for ti, t in enumerate(doc.tables):
    print(f"=== TABLE {ti} ===")
    for r in t.rows:
        print(" | ".join(c.text for c in r.cells))
