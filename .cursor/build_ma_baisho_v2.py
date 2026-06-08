"""馬強事件 損害賠償請求書 docx v2 生成（簡素版）。

修正指示（260522_馬氏3.md）：
- Word デフォルト余白
- 1文字下げ（本文段落先頭に全角スペース1個のみ。first_line_indent はクリア）
- タイトル・見出しは左端で本文より左
- 精神的損害の表削除→文章記載（最低100万円、長期化で増額）
- 弁護士費用の表削除（説明段落に金額追加）

戦略：ユーザー編集版をベースに、最小限の修正で済ます。
- 全パラグラフの first_line_indent を None
- マージンを Word デフォルト 2.54cm に
- 精神的損害テーブル削除＋文章追加
- 弁護士費用テーブル削除＋説明段落の金額追記
"""

import os
import shutil
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260522_損害賠償請求書（馬様）.docx"
)
DST = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260522_損害賠償請求書（馬様）_2.docx"
)


def rep(para, text):
    if para.runs:
        first = para.runs[0]
        first.text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    # ---- 1. Word デフォルトマージン (2.54cm = 1 inch) ----
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    # ---- 2. 全段落の first_line_indent をクリア ----
    # 本文の1字下げは段落先頭の「全角スペース1個」で行う
    for p in doc.paragraphs:
        p.paragraph_format.first_line_indent = None

    # 表セル内段落も同様
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.first_line_indent = None

    # ---- 3. 差出人ブロックを右寄せ ----
    paras = list(doc.paragraphs)
    paras[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 日付
    for i in range(6, 13):  # 差出人住所〜TEL
        paras[i].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ---- 4. 本文段落の先頭に全角スペース1個があるかチェックして補正 ----
    # ユーザー編集版で既に全角spaceが付いている段落はそのまま。
    # アイウエ箇条書き行（ア　... イ　...）はそのまま（先頭文字+全角space）。
    # 見出し（第○、（○）、【○】、損害賠償請求書、以上）は字下げ無し。

    # ---- 5. 精神的損害テーブル削除＋文章追加 ----
    body = doc.element.body
    target_mental = None
    target_fee = None
    for el in list(body):
        if el.tag != qn("w:tbl"):
            continue
        text = "".join(el.itertext())
        if "Angelina氏及び通知人らに対する慰謝料" in text:
            target_mental = el
        elif "弁護士費用（実費損害＋精神的損害" in text:
            target_fee = el

    if target_mental is not None:
        # 表前に新規段落を挿入
        new_p_el = OxmlElement("w:p")
        target_mental.addprevious(new_p_el)
        new_para = Paragraph(new_p_el, target_mental.getparent())
        new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = new_para.add_run(
            "　以上の精神的苦痛に対する慰謝料の額は、少なくとも金"
            "１,０００,０００円を下ることはないため、通知人は、貴校及び"
            "小牧氏に対し、連帯して、同金額の支払を請求いたします。なお、"
            "Angelina氏が貴校に通学できない期間が長期化する場合には、"
            "本慰謝料請求額を増額する予定です。"
        )
        run.font.size = Pt(10.5)
        target_mental.getparent().remove(target_mental)

    # ---- 6. 弁護士費用テーブル削除 ----
    if target_fee is not None:
        target_fee.getparent().remove(target_fee)

    # 弁護士費用説明段落の金額を明示
    for p in doc.paragraphs:
        if p.text.startswith("　通知人は、本件法的対応のために当職に委任せざるを得ず"):
            rep(
                p,
                "　通知人は、本件法的対応のために当職に委任せざるを得ず、"
                "貴校及び小牧氏の不法行為と相当因果関係を有する弁護士費用"
                "として、前記（１）及び（２）の既確定額合計（金２,６１７,"
                "７００円）の１０％相当額である金２６１,７７０円を請求"
                "いたします。",
            )
            break

    # ---- 7. 連続する空行を1つに圧縮（表削除後の連続空行を解消） ----
    body = doc.element.body
    prev_was_empty_p = False
    to_remove = []
    for el in list(body):
        if el.tag != qn("w:p"):
            prev_was_empty_p = False
            continue
        para = Paragraph(el, body)
        is_empty = (para.text.strip() == "")
        if is_empty and prev_was_empty_p:
            to_remove.append(el)
        prev_was_empty_p = is_empty
    for el in to_remove:
        el.getparent().remove(el)

    doc.save(DST)
    print("saved:", DST)


if __name__ == "__main__":
    main()
