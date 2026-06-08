"""v6/v2契約書の本文・スタイル・コメントをダンプ。"""
from pathlib import Path
import zipfile
import sys
from docx import Document

paths = {
    "v6": Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v6.docx"),
    "v2": Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v2.docx"),
}

label = sys.argv[1] if len(sys.argv) > 1 else "v6"
SRC = paths[label]

doc = Document(str(SRC))
print("=" * 70)
print(f"【{label} 本文（段落・スタイル・揃え）】 file={SRC.name}")
print("=" * 70)
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ""
    align = p.alignment
    text = p.text
    # numbering 情報
    numId = None
    ilvl = None
    pPr = p._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if pPr is not None:
        numPr = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
        if numPr is not None:
            ni = numPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
            il = numPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
            numId = ni.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") if ni is not None else None
            ilvl = il.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") if il is not None else None
    print(f"[{i:03d}] style={style!r:25s} align={align} num=(id={numId},lv={ilvl}) text={text!r}")

print()
print("=" * 70)
print(f"【{label} 表】")
print("=" * 70)
for ti, table in enumerate(doc.tables):
    print(f"-- table {ti} rows={len(table.rows)} cols={len(table.columns)}")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            print(f"  r{ri}c{ci}: {cell.text!r}")

print()
print("=" * 70)
print(f"【{label} ZIP内エントリ & コメントXML】")
print("=" * 70)
with zipfile.ZipFile(SRC) as zf:
    names = zf.namelist()
    for n in names:
        if "comment" in n.lower() or n == "word/document.xml":
            print(f"  {n}")
    if "word/comments.xml" in names:
        with zf.open("word/comments.xml") as f:
            print("\n--- comments.xml ---")
            print(f.read().decode("utf-8"))
    else:
        print("(no word/comments.xml)")
