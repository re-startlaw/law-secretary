"""社宅管理規程制定のみの臨時社員総会議事録（第1号議案単独版）を作成。"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

DEST_DIR = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/"
    "14_弁護士法人/01_定款・規定/規定"
)
FNAME = "260506_社員総会議事録_社宅規程制定_2.docx"

HOJIN_NAME = "弁護士法人Re-Start法律事務所"
DAIHYO = "米谷 尚起"
SOKAI_BI = "令和8年5月7日"
SEKO_BI = "令和8年5月7日"


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "游明朝"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "游明朝")
    rfonts.set(qn("w:ascii"), "游明朝")
    rfonts.set(qn("w:hAnsi"), "游明朝")


def _title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def _right(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(text)


def _para(doc: Document, text: str = "") -> None:
    doc.add_paragraph(text)


def _bold_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def main() -> None:
    doc = Document()
    _set_default_font(doc)

    _title(doc, "臨時社員総会議事録")
    _para(doc)

    _para(
        doc,
        f"　{SOKAI_BI}午前10時00分、{HOJIN_NAME}（以下「法人」という。）"
        "の臨時社員総会を法人本店において開催した。",
    )
    _para(doc)
    _para(doc, "社員の総数　　　　　　　　1名")
    _para(doc, "出資金の総額　　　　　　　金10,000円")
    _para(doc, "出席社員数　　　　　　　　1名")
    _para(doc, "出席社員の出資金の総額　　金10,000円")
    _para(doc)
    _para(
        doc,
        f"　以上のとおり社員全員出席のうえ、満場一致をもって議長に代表社員{DAIHYO}を選出し、"
        "同人は議長席に着き、開会を宣し直ちに議事に入った。",
    )
    _para(doc)

    _bold_line(doc, "議案　社宅管理規程制定の件")
    _para(
        doc,
        "　議長は、法人の社員に対する社宅貸与の取扱いを明確にするため、"
        "別紙のとおり社宅管理規程を制定したい旨を述べ、その理由及び規程の内容を詳細に説明し、"
        "その可否を諮ったところ、出席社員全員の賛成をもって原案のとおり可決承認された。",
    )
    _para(doc, f"　なお、本規程の施行日は{SEKO_BI}とする。")
    _para(doc)

    _para(
        doc,
        f"　以上をもって本臨時社員総会の議事を終了したので、議長は{SOKAI_BI}午前10時30分閉会を宣した。",
    )
    _para(doc)
    _para(
        doc,
        "　上記の議事の経過の要領及びその結果を明確にするため、本議事録を作成し、"
        "出席代表社員において次のとおり記名押印する。",
    )
    _para(doc)
    _right(doc, SOKAI_BI)
    _para(doc)
    _right(doc, f"{HOJIN_NAME}　臨時社員総会")
    _para(doc)
    _right(doc, f"議長兼議事録作成者　代表社員　{DAIHYO}　　印")

    path = os.path.join(DEST_DIR, FNAME)
    doc.save(path)
    print(f"作成: {path}")


if __name__ == "__main__":
    main()
