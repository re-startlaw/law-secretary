#!/usr/bin/env python3
"""田村正宣案件 上申書(接見許可)を聴取報告書の内容を取り込んでv2作成

ルール:
- 元ファイルと同じフォルダに _2 を作成（既存なら _3, _4 ...）
- 元ファイル(260520_上申書.docx)は触らない
- 段落揃え: 日付RIGHT, タイトルCENTER, 以上RIGHT, 他LEFT
- 字下げは既存上申書の体裁(各段落フラッシュ)を踏襲
"""
import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

BASE = Path(
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/"
    "01_事件記録/た_田村正宣/09_ワードファイル"
)
SRC = BASE / "260520_上申書.docx"


def next_versioned_path(src: Path) -> Path:
    stem = src.stem  # 260520_上申書
    suffix = src.suffix  # .docx
    n = 2
    while True:
        cand = src.with_name(f"{stem}_{n}{suffix}")
        if not cand.exists():
            return cand
        n += 1


def clear_paragraphs(doc: Document) -> None:
    """body直下のp要素を全て削除"""
    body = doc.element.body
    # body末尾には sectPr があるので、それ以外の p を消す
    for p in list(body.iter(qn("w:p"))):
        # iterで深い階層まで拾うのは避け、直下だけにする
        if p.getparent() is body:
            body.remove(p)


def qn(tag: str) -> str:
    from docx.oxml.ns import qn as _qn
    return _qn(tag)


def add_para(doc, text="", align=None, bold=False, size=None, indent_first=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if indent_first:
        pf.first_line_indent = Pt(12)  # 1文字分相当
    if text:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if size is not None:
            run.font.size = size
    return p


def main():
    assert SRC.exists(), f"元ファイルが見つかりません: {SRC}"
    dst = next_versioned_path(SRC)
    shutil.copy2(SRC, dst)
    print(f"copied: {SRC.name} -> {dst.name}")

    doc = Document(str(dst))
    clear_paragraphs(doc)

    # ---- 表題 ----
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("上申書")
    r.bold = True
    r.font.size = Pt(18)

    add_para(doc)  # 空行

    # ---- 日付 ----
    add_para(doc, "令和８年５月２０日", align=WD_PARAGRAPH_ALIGNMENT.RIGHT)

    # ---- 宛先 ----
    add_para(doc, "東京地方裁判所　刑事部　御中", align=WD_PARAGRAPH_ALIGNMENT.LEFT)

    add_para(doc)

    # ---- 上申人 ----
    add_para(doc, "上申人（被疑者内縁の妻）", align=WD_PARAGRAPH_ALIGNMENT.RIGHT)
    add_para(doc, "氏名　　大宮　瑞希　　印", align=WD_PARAGRAPH_ALIGNMENT.RIGHT)

    add_para(doc)

    # ---- 導入 ----
    add_para(
        doc,
        "私は、本件被疑者である田村正宣の内縁の妻です。",
    )
    add_para(
        doc,
        "主人とは令和元年４月に出会って以来、約７年にわたって夫婦同然の生活を送ってまいりました。"
        "事情により婚姻届は提出しておりませんが、主人は私との間にもうけた長女・向日凜（ひまり。令和３年９月２９日生）、"
        "長男・優篤（ゆうま。令和５年１月５日生）、次男・護（まもる。令和８年１月２２日生）の３人の子をすべて認知し、"
        "責任をもって育ててくれており、文字どおり一家の大黒柱です。"
        "主人は仕事で出張することもありましたが、ほぼ日帰りで、これまで長期にわたって自宅を不在にしたことは一度もございませんでした。",
    )
    add_para(
        doc,
        "その主人が突然逮捕され、その後の予期せぬ事態に直面し、私自身の精神状態も限界に近く、"
        "また３人の子供たちの生活を守るために一刻の猶予もない状況にあります。"
        "主人との接見を認めていただきたく、重ねてお願い申し上げます。",
    )

    add_para(doc)

    # ---- 1. 自宅特定と避難先での限界 ----
    add_para(doc, "１　報道による自宅の特定と、避難先での限界について", bold=True)
    add_para(
        doc,
        "主人の逮捕後、実名や顔写真、さらには自宅住所の町名までが詳細に報道されました。"
        "主人は子供たちの幼稚園の送り迎えを毎日欠かさず行っており、地域の方々や園の保護者の方にも顔が知られていたため、"
        "一気に噂が広まり、志木市の自宅には身の危険を感じていられなくなりました。",
    )
    add_para(
        doc,
        "急遽、茨城県の実家へ子供３人を連れて避難いたしましたが、私の両親は花屋を営む自営業者であり、"
        "共働きで朝から晩まで店に出て働いております。そのため日中の育児サポートはまったく得られず、"
        "私一人が実家という慣れない環境で、生後３か月の乳児を含む３人の子供の世話に追われております。"
        "両親の生活リズムを壊している申し訳なさもあり、いつまでも身を寄せていられる状況ではありません。",
    )

    add_para(doc)

    # ---- 2. 精神的苦痛 ----
    add_para(doc, "２　精神的な苦痛と切迫感について", bold=True)
    add_para(
        doc,
        "主人が突然いなくなったショックに加え、連日のように流れる報道やネット上の反応、"
        "そして住み慣れた自宅を追われるという異常な事態に、私は現在、夜も全く眠れないほど精神的に追い詰められています。",
    )
    add_para(
        doc,
        "３人の子供たちの前では何とか気を張っていますが、今後の生活や住居、お金のこと、そして主人の安否を考えると、"
        "不安で押しつぶされそうになり、涙が止まらない時間もあります。"
        "特に生後３か月の次男を抱えながら、自分一人では何も決められない、どこへ行けばいいのかもわからないという状態です。",
    )

    add_para(doc)

    # ---- 3. 協議の必要性 ----
    add_para(doc, "３　主人との協議の必要性", bold=True)
    add_para(
        doc,
        "報道により元の自宅へは戻れないため、早急に新しい引越し先を見つける必要があります。"
        "しかし、引越し先をどの地域にするか、それに伴う幼稚園の転園手続き、そして一番下の子供の保育園入園の相談など、"
        "主人の同意なしに進めることは不可能です。",
    )
    add_para(
        doc,
        "また、家計のシミュレーション（別紙）のとおり、子供３人を育てるには多額の生活費がかかりますが、"
        "私は３人の子供の子育てのため現在働いておらず、生活費はすべて主人から受け取って参りました。"
        "手元の資金で何とか凌げるのは長くても１～２か月程度で、それ以上主人の勾留が続けば、"
        "３人の子供たちは路頭に迷ってしまいます。"
        "引越し費用も必要となる中で、主人の周囲の方々からのサポートをどう仰げばよいのか、"
        "今後の生活をどう立て直すべきか、主人の口から直接指示を受け、安心させてほしいと切に願っています。",
    )

    add_para(doc)

    # ---- 4. 罪証隠滅のおそれがないこと ----
    add_para(doc, "４　罪証隠滅のおそれがないこと", bold=True)
    add_para(
        doc,
        "今回主人と共に逮捕されている方々とは、いずれも主人を介して顔を合わせたことがあり、"
        "お子さんがいらっしゃる方も多いことから、家族ぐるみでお会いしたこともございます。"
        "もっとも、私はこれらの方々と連絡先を交換しておらず、主人を介さずに連絡を取る手段は全くありません。",
    )
    add_para(
        doc,
        "また、これらの方々と私から連絡を取ることは、証拠の隠滅を疑われる行為として絶対に行ってはならないと、"
        "弁護人である米谷弁護士からも繰り返し説明を受けて十分に理解しております。"
        "主人に対しても、決してそのような働きかけをしないよう強く言い聞かせる所存です。",
    )
    add_para(
        doc,
        "そもそも、生後３か月の乳児を含む３人の子供を抱える今の私には、子供たちの生活を守ることで精一杯であり、"
        "証拠隠滅に関与する理由も余裕もございません。",
    )

    add_para(doc)

    # ---- 5. 結び ----
    add_para(doc, "５　結びに", bold=True)
    add_para(
        doc,
        "面会には警察官の方が立ち会われると伺っており、事件の話をするつもりは一切ございません。"
        "これからの引越し先や、子供たちの幼稚園・保育園のこと、そして当面の生活費をどうやって工面するのかという、"
        "生きていくための切実な相談をさせてください。",
    )
    add_para(
        doc,
        "どうか、私と主人との面会をお許しくださいますよう、心よりお願い申し上げます。",
    )

    add_para(doc)
    add_para(doc, "以上", align=WD_PARAGRAPH_ALIGNMENT.RIGHT)

    doc.save(str(dst))
    print(f"saved: {dst}")
    return dst


if __name__ == "__main__":
    p = main()
    print(p)
