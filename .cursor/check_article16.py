"""第16条の項数を確認。"""
from pathlib import Path
from docx import Document
import re

v7_2_path = Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v7_2.docx")

doc = Document(str(v7_2_path))

print("=" * 100)
print("【第16条全文と項数を確認】")
print("=" * 100)

start_idx = None
for i, p in enumerate(doc.paragraphs):
    if re.match(r'^第16条', p.text.strip()):
        start_idx = i
        break

if start_idx:
    print(f"第16条開始位置: 段落{start_idx}\n")
    for i in range(start_idx, min(start_idx + 30, len(doc.paragraphs))):
        text = doc.paragraphs[i].text
        print(f"[{i:03d}] {text}")
        if i > start_idx and re.match(r'^第17条', text):
            break

print("\n" + "=" * 100)
print("【第16条項数確認】")
print("=" * 100)

# 第16条内のすべての項を数える
for i in range(start_idx, min(start_idx + 30, len(doc.paragraphs))):
    text = doc.paragraphs[i].text.strip()
    if re.match(r'^第17条', text):
        break
    item_match = re.match(r'^（?[\d１２３４５６７８９０]+[.]', text)
    if item_match:
        print(f"  項: {text[:50]}")

