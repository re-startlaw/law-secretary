"""v7_2 の詳細レビュー - 第11条と全文で違和感表現を抽出。"""
from pathlib import Path
from docx import Document
import re

v3_path = Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v3.docx")
v7_2_path = Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v7_2.docx")

def load_doc(path):
    return Document(str(path))

def extract_articles(doc):
    """条文ごとに段落をグループ化。"""
    articles = {}
    current_article = None
    current_article_num = None
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        match = re.match(r'^第(\d+)条', text)
        
        if match:
            current_article_num = int(match.group(1))
            articles[current_article_num] = {
                'start_idx': i,
                'title': text,
                'paragraphs': []
            }
            current_article = current_article_num
        elif current_article is not None:
            articles[current_article]['paragraphs'].append({
                'idx': i,
                'text': p.text
            })
    
    return articles

def check_suspicious_v7_2():
    """v7_2全文を走査して違和感表現を検出。"""
    doc = load_doc(v7_2_path)
    articles = extract_articles(doc)
    
    # 違和感パターン集
    patterns = {
        '推奨表現（望ましい）': r'（望ましい|望ましい。）',
        '推奨表現（良い）': r'（良い|要するに）',
        '意見・説明（考えられる）': r'（と考えられ|と考える|考えた|思われ|思える）',
        '説明的表現': r'（実質的に|いわゆる|結果として|当該|該当）',
        '解説的フレーズ': r'（価値の最大化|紛争の抑制|塩漬け|なり得る）',
        '法令説明が紛れ込み': r'法令上、|法令が|第三者提供|個人情報保護法',
        'やや曖昧な「単なる」': r'（単に|単なる|ただし）',
    }
    
    results = {}
    for article_num in sorted(articles.keys()):
        article = articles[article_num]
        title = article['title']
        suspicious_in_article = []
        
        # タイトルも検査
        for p_text in [title]:
            for reason, pattern in patterns.items():
                if re.search(pattern, p_text):
                    suspicious_in_article.append({
                        'section': f'（タイトル）',
                        'text': p_text,
                        'reason': reason
                    })
        
        # 各項を検査
        for para in article['paragraphs']:
            p_text = para['text']
            para_idx = para['idx']
            
            # 項番号を抽出（「１．」「２．」など）
            item_match = re.match(r'^（?[\d１２３４５６７８９０]+[.]）？(.*)$', p_text)
            section_label = f"第{article_num}条"
            
            for reason, pattern in patterns.items():
                if re.search(pattern, p_text):
                    suspicious_in_article.append({
                        'para_idx': para_idx,
                        'section': section_label,
                        'text': p_text,
                        'reason': reason
                    })
        
        if suspicious_in_article:
            results[article_num] = suspicious_in_article
    
    return results, articles

# ====== main ======
print("=" * 100)
print("【v7_2 契約書 違和感表現の詳細抽出】")
print("=" * 100)

results, articles = check_suspicious_v7_2()

print("\n【第11条（知的財産権の処理）全文抽出】")
print("=" * 100)
if 11 in articles:
    article_11 = articles[11]
    print(f"\nタイトル: {article_11['title']}")
    for para in article_11['paragraphs']:
        print(f"\n[段落{para['idx']:03d}] {para['text']!r}")

print("\n\n【全条文で抽出した違和感表現（第11条以外も含む）】")
print("=" * 100)
for article_num in sorted(results.keys()):
    suspicious_list = results[article_num]
    if not suspicious_list:
        continue
    print(f"\n【第{article_num}条】")
    for item in suspicious_list:
        print(f"\n  理由: {item['reason']}")
        print(f"  位置: {item.get('para_idx', 'N/A')}")
        print(f"  文: {item['text']!r}")

print("\n\n【v3 vs v7_2 条文数・構成比較】")
print("=" * 100)
doc_v3 = load_doc(v3_path)
doc_v7_2 = load_doc(v7_2_path)
articles_v3 = extract_articles(doc_v3)
articles_v7_2 = extract_articles(doc_v7_2)

print(f"v3: {len(doc_v3.paragraphs)} 段落、{len(articles_v3)} 条文")
print(f"v7_2: {len(doc_v7_2.paragraphs)} 段落、{len(articles_v7_2)} 条文")

v3_nums = set(articles_v3.keys())
v7_2_nums = set(articles_v7_2.keys())
common = v3_nums & v7_2_nums
v3_only = v3_nums - v7_2_nums
v7_2_only = v7_2_nums - v3_nums

print(f"\n共通の条文: {len(common)} 条")
print(f"v3のみ: {sorted(v3_only)}")
print(f"v7_2のみ: {sorted(v7_2_only)}")

