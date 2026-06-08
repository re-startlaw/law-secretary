"""学習誓約書（見学者用）.docx を作成する。"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = "/Users/kometaninaoki/law-secretary/obsidian-vault/50_プロンプト/学習誓約書_2.docx"

# 10.5pt の全角1文字幅 ≒ 0.37cm
CHAR1 = 0.37  # 1字
CHAR2 = 0.74  # 2字（「1　」「2　」相当）
CHAR_PAREN = 0.95  # 「(1)　」相当

FIRM_NAME = "弁護士法人Re-Start法律事務所"
LAWYER_NAME = "米谷尚起"
DATE_STR = "令和８年５月２０日"


def set_font(run, size=10.5, bold=False):
    run.font.name = "游明朝"
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.find(qn("w:rPr"))
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "游明朝")
    rFonts.set(qn("w:ascii"), "游明朝")
    rFonts.set(qn("w:hAnsi"), "游明朝")


def add_para(doc, text, *, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=10.5,
             space_after=None, first_line_indent=None, left_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p


def add_article(doc, heading, items=None, body=None):
    add_para(doc, heading, bold=True, size=11, space_after=2)
    if body is not None:
        add_para(doc, body, first_line_indent=CHAR1)
    if items is not None:
        for i, item in enumerate(items, start=1):
            add_para(doc, f"{i}　{item}", left_indent=CHAR2, first_line_indent=-CHAR2)


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "游明朝"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")

# タイトル
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(18)
run = title_p.add_run("守秘義務に関する誓約書（法律事務所見学者用）")
set_font(run, size=14, bold=True)

# 宛先
add_para(doc, f"{FIRM_NAME}", size=11)
add_para(doc, f"代表社員弁護士　{LAWYER_NAME}　御中", size=11, space_after=12)

# 前文
preamble = (
    "私は、貴職の法律事務所（以下「貴事務所」という。）を訪問し、また学習の目的で、"
    "貴事務所が処理中または過去に処理した事件等に関する資料（当事者の氏名等が秘匿される"
    "ようマスキング等の適切な匿名化措置が施されたもの。以下「本件教材用資料」という。）"
    "の閲覧・説明を受けるにあたり、以下の事項を遵守することを誓約いたします。"
)
add_para(doc, preamble, first_line_indent=CHAR1, space_after=12)

# 第１条
add_para(doc, "第１条（秘密情報の定義）", bold=True, size=11, space_after=2)
add_para(
    doc,
    "本誓約書において「秘密情報」とは、私が貴事務所への訪問、本件教材用資料の閲覧、"
    "および貴事務所の所属員との対話等の過程において知得した、以下の情報をいう。",
    first_line_indent=CHAR1,
)
add_para(
    doc,
    "(1)　貴事務所の顧客、依頼者、相談者、および事件関係者に関する一切の情報"
    "（氏名、住所、連絡先等の個人情報のほか、相談内容、事件内容、紛争の存在自体を含むが、"
    "これらに限られない。）",
    left_indent=CHAR_PAREN, first_line_indent=-CHAR_PAREN,
)
add_para(
    doc,
    "(2)　本件教材用資料に記載された、事件の背景、経緯、主張内容、ノウハウ等の情報"
    "（マスキングがなされているか否かを問わない。）",
    left_indent=CHAR_PAREN, first_line_indent=-CHAR_PAREN,
)
add_para(
    doc,
    "(3)　貴事務所の経営、業務運営、システム、および技術に関する一切の情報",
    left_indent=CHAR_PAREN, first_line_indent=-CHAR_PAREN,
)
add_para(
    doc,
    "(4)　その他、貴事務所が秘密である旨を指定した情報、または通常秘密として扱われる"
    "べき性質の情報",
    left_indent=CHAR_PAREN, first_line_indent=-CHAR_PAREN, space_after=10,
)

# 第２条
add_article(
    doc,
    "第２条（秘密保持義務）",
    items=[
        "私は、秘密情報を厳重に管理し、貴職の事前の書面による承諾がない限り、"
        "第三者（他の学生、友人、家族等を含むがこれらに限られない。）に対して一切開示、"
        "漏洩、または提供いたしません。",
        "私は、本件教材用資料のマスキングされた部分を推測、復元、または特定しようとする"
        "行為を行いません。",
    ],
)
add_para(doc, "", space_after=4)

# 第３条
add_article(
    doc,
    "第３条（目的外使用の禁止およびＳＮＳ等の発信禁止）",
    items=[
        "私は、秘密情報を私自身の司法試験等の勉強・学習の目的のみに使用し、"
        "他のいかなる目的（商業目的、第三者の利益、または自己の利益等）にも使用いたしません。",
        "私は、貴事務所での見学内容、本件教材用資料の内容、および知得した秘密情報"
        "（これらを推測させる内容を含む）について、ＳＮＳ（Ｘ、Ｉｎｓｔａｇｒａｍ、"
        "ＬＩＮＥ、ブログ等）やインターネット上の掲示板等への投稿・配信をはじめとする、"
        "不特定多数への発信行為を一切行いません。",
    ],
)
add_para(doc, "", space_after=4)

# 第４条
add_article(
    doc,
    "第４条（撮影・録音・複写等の禁止）",
    body="私は、貴事務所内において、また本件教材用資料について、貴職の許可なく写真撮影、"
         "録音、録画、スマートフォンのカメラ等によるスキャン、コピー、または外部への"
         "データ転送等を行いません。",
)
add_para(doc, "", space_after=4)

# 第５条
add_article(
    doc,
    "第５条（資料の返還・破棄）",
    body="私は、見学および閲覧が終了したときは、閲覧した資料等を直ちに貴事務所に返還します。"
         "また、万が一、私の私有端末等に秘密情報が含まれるデータが保存された場合は、"
         "貴職の指示に従い直ちに完全に消去・破棄します。",
)
add_para(doc, "", space_after=4)

# 第６条
add_article(
    doc,
    "第６条（損害賠償および法的責任）",
    body="私が本誓約書に違反し、貴事務所、その顧客、または第三者に損害を与えた場合、"
         "私はこれにより生じた一切の損害（合理的な弁護士費用を含む）を賠償する責任を負う"
         "ことを承諾します。また、民事上の責任に留まらず、弁護士法その他の法令に基づく"
         "社会的・法的責任が問われ得ることを理解しています。",
)
add_para(doc, "", space_after=4)

# 第７条
add_article(
    doc,
    "第７条（有効期間）",
    body="本誓約書に基づく私の守秘義務は、貴事務所への訪問および資料の閲覧が終了した後も、"
         "期間の定めなく存続するものとします。",
)

# 以上
add_para(doc, "以上", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

# 日付
add_para(doc, DATE_STR, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

# 署名欄
add_para(doc, "（誓約者）", space_after=10)
add_para(doc, "氏名（自署）：　　　　　　　　　　　　　　　　　　　　　　　　　　印", space_after=12)
add_para(doc, "所属（大学・学年等）：　　　　　　　　　　　　　　　　　　　　　　", space_after=12)
add_para(doc, "現住所：　　　　　　　　　　　　　　　　　　　　　　　　　　　　　", space_after=12)
add_para(doc, "連絡先（電話番号）：　　　　　　　　　　　　　　　　　　　　　　　", space_after=12)

doc.save(OUT)
print(f"Saved: {OUT}")
