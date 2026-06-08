#!/usr/bin/env python3
"""毛さんVSヒメラ 準備書面（８）_2 編集スクリプト

編集内容：
- 認否 第3（売上データ提出について）をユーザー指示に沿って書き直し
- 第2 第2（決算書等の提出義務がないこと）を準備書面（7）第2の3 引用に圧縮
- 第2 第3（売上データ提出義務がないこと）を同様に圧縮
"""
from docx import Document
from copy import deepcopy

PATH = '/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/も_毛さんVSヒメラ合同会社/訴訟/00主張/260514_準備書面（８）_2.docx'


def set_paragraph_text(para, text):
    """段落のテキストを置換。最初のrunに新テキストを入れ、他のrunは削除。"""
    if not para.runs:
        para.add_run(text)
        return
    first = para.runs[0]
    first.text = text
    # 2番目以降のrunを削除
    for r in para.runs[1:]:
        r._element.getparent().remove(r._element)


def replace_block(doc, start_idx, end_idx, new_texts):
    """段落 start_idx から end_idx（含む）を new_texts で置換。

    既存段落のスタイル（最初の段落のスタイル）を維持。
    new_texts が既存より少ない場合は余った段落を削除。
    多い場合は最後の段落の後に挿入。
    """
    paras = doc.paragraphs
    target = paras[start_idx:end_idx + 1]
    style = target[0].style
    # 既存段落を順に置換
    for i, text in enumerate(new_texts):
        if i < len(target):
            set_paragraph_text(target[i], text)
        else:
            # 足りない場合は前の段落の後ろに挿入
            new_para = target[-1]._element.addnext(deepcopy(target[-1]._element))
            from docx.text.paragraph import Paragraph
            p = Paragraph(target[-1]._element.getnext(), target[-1]._parent)
            set_paragraph_text(p, text)
    # 余った段落を削除
    if len(target) > len(new_texts):
        for p in target[len(new_texts):]:
            p._element.getparent().remove(p._element)


def main():
    doc = Document(PATH)

    # === 認否 第3（段落 24-26 のうち 24, 25 を書き換え、26 は維持） ===
    ninhi_3_p1 = (
        '第１段落のうち、被告が令和３年４月に乙２７をもって売上に係るデータの提出を求めた事実、'
        '被告が原告に対し令和３年７月２日及び同月１２日に提出を求めた事実、'
        '並びに原告が同月２７日に令和２年４月１日から令和３年４月３０日までの売上に係るデータを提出した事実（乙３１）は認め、'
        '被告が同年６月１２日及び同月１４日にも提出を求めたとの点は、提出された証拠からは確認することができないので不知、'
        'その余は争う。'
    )
    ninhi_3_p2 = (
        '第２段落のうち、原告が令和３年５月末日以降の売上に係るデータを別途提出しなかったことは認め、その余は争う。'
    )
    set_paragraph_text(doc.paragraphs[24], ninhi_3_p1)
    set_paragraph_text(doc.paragraphs[25], ninhi_3_p2)
    # 26 「争う部分の理由は、いずれも後述のとおりである。」は維持

    # === 第2 第2「決算書等の提出義務がないこと」段落 57-59 ===
    kessanshobun_p1 = (
        '被告書面第２（決算書等の提出について）の主張は、本件契約第７条を根拠として、'
        '原告に対し決算書、仕入台帳、財務諸表、税務申告書等の広範な書類の提出義務を課そうとするものである。'
    )
    kessanshobun_p2 = (
        'しかしながら、本件契約第７条が加盟店の義務として明示的に定めているのは、'
        '「前年の店舗の運営に関する収入と支出の明細書」の提出のみであり、'
        'これに続く「本店が合理的に必要とするフォーム、レポート、記録、明細書、財務諸表、およびその他の情報」の提出義務は、'
        'その文言上「本店が合理的に必要とする」場合という限定が付されている。'
        'また、被告書面が下線を付した「査看」（可以查看）の権限も、'
        'その文言上、合理的な事前通知を行った上で領帳、記録、スケジュール及び納税申告書等を閲覧することを許容するに過ぎず、'
        'かつ、その性格は契約期間中の監督権限にとどまる。'
    )
    kessanshobun_p3 = (
        'これらの点に加え、定額ロイヤリティ制の下では被告に損害が生じる関係にないこと（必要性の欠如）、'
        '本件契約終了後の元加盟店に対して遡及的に包括的な開示義務を課す根拠が存在しないこと、'
        '及び被告が求める情報が原告の営業秘密に属することは、'
        'いずれも令和８年２月１９日付準備書面（７）第２の３「文書提出義務の不存在」において詳述したとおりであり、'
        'これを援用する。'
    )
    kessanshobun_p4 = (
        'したがって、被告書面が求める決算書、仕入台帳等の仕入れ及び売上に関する資料、'
        '並びに税務申告書等の情報について、原告が被告に対しこれを開示すべき契約上の義務は存しない。'
    )
    # 段落 57-59 (3つ) を 4つに置換 → 1つ追加挿入が必要
    # 既存: 57, 58, 59 → 4つに拡張
    set_paragraph_text(doc.paragraphs[57], kessanshobun_p1)
    set_paragraph_text(doc.paragraphs[58], kessanshobun_p2)
    set_paragraph_text(doc.paragraphs[59], kessanshobun_p3)
    # 59の後に新しい段落を挿入（59のスタイルをコピー）
    p59 = doc.paragraphs[59]
    new_p_elem = deepcopy(p59._element)
    p59._element.addnext(new_p_elem)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p_elem, p59._parent)
    set_paragraph_text(new_para, kessanshobun_p4)

    # 段落番号がずれるため、再取得して 売上データ部分を編集
    doc2_paras = doc.paragraphs
    # 元 61, 62, 63 → 1つ追加分ずれて 62, 63, 64 になっている
    # 元: 60「ランク２　売上に係るデータの提出義務がないこと」見出し
    # 元: 61, 62, 63 本文 → 編集対象
    # 1つ挿入後: 見出し 61, 本文 62, 63, 64
    # ヘッダ「売上に係るデータの提出義務がないこと」を探す
    target_heading_idx = None
    for i, p in enumerate(doc2_paras):
        if '売上に係るデータの提出義務がないこと' in p.text:
            target_heading_idx = i
            break
    assert target_heading_idx is not None, '売上データ提出義務がない見出しが見つかりません'
    # 本文は次の3段落
    uriage_p1 = (
        '原告は、被告の求めに応じ、令和３年７月２７日に、'
        '令和２年４月１日から令和３年４月３０日までの売上に係るデータを提出している（乙３１）。'
    )
    uriage_p2 = (
        '被告は、令和３年５月末日以降の進行中の月分の売上データをも逐次提出すべきであった旨主張するが、'
        '本件契約第７条が加盟店の義務として明示的に定めているのは「前年の店舗の運営に関する収入と支出の明細書」の提出に限られ、'
        '進行中の月分の売上データの逐次提出を義務付けるものではない。'
        'また、本件契約終了後における詳細な経営データの開示義務が存在しないこと、'
        '及び被告が求めるデータが原告の営業秘密に属することは、'
        '前項「決算書等の提出義務がないこと」で援用した令和８年２月１９日付準備書面（７）第２の３「文書提出義務の不存在」において主張したとおりである。'
    )
    uriage_p3 = (
        'したがって、原告に売上に係るデータの提出義務違反は存しない。'
    )
    set_paragraph_text(doc.paragraphs[target_heading_idx + 1], uriage_p1)
    set_paragraph_text(doc.paragraphs[target_heading_idx + 2], uriage_p2)
    set_paragraph_text(doc.paragraphs[target_heading_idx + 3], uriage_p3)

    doc.save(PATH)
    print('saved:', PATH)
    # 確認用にダンプ
    d = Document(PATH)
    for i, p in enumerate(d.paragraphs):
        if p.text.strip():
            print(f'{i}\t[{p.style.name}]\t{p.text[:150]}')


if __name__ == '__main__':
    main()
