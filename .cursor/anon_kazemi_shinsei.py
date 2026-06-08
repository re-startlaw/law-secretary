# -*- coding: utf-8 -*-
import shutil
from docx import Document

src = "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/か_KAZEMI・HOSSEIN/提出書類/妻/申請理由書（サンプル）.docx"
dst = "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/か_KAZEMI・HOSSEIN/提出書類/妻/申請理由書（匿名化テンプレート）.docx"

shutil.copy2(src, dst)

# 置換は「特定性の高い・長い文字列」から先に適用する（順序重要）
replacements = [
    # 氏名（テンプレート化）
    ("KAZEMI HOSSEIN", "【配偶者氏名】"),
    ("DELBARI TAHEREH", "【申請人氏名】"),
    # 金額・通貨（特定情報）
    ("毎月850,000,000イランレアル（ＩＲＲ）", "毎月○○○○（現地通貨）"),
    ("約300万円", "約○○万円"),
    # 大使館・チャットツール（特定情報）
    ("在日本イラン大使館", "在日本○○大使館"),
    ("トルコのチャットツール「imo」", "○○のチャットツール「○○」"),
    # 戦争従軍歴（特定情報）
    ("イラン・イラク戦争に徴兵され２年従軍しました", "○○での戦争に徴兵され○年従軍しました"),
    # 不動産の個数（特定情報）
    ("不動産を３つ所有しており", "不動産を複数所有しており"),
    # 日付（特定情報）— 完全一致の長い順に
    ("2019年の11月1日", "○年○月○日"),
    ("2020年11月11日", "○年○月○日"),
    ("2022年10月31日", "○年○月○日"),
    ("２０１１年１２月１２日", "○年○月○日"),
    ("２０１４年１０月２８日", "○年○月○日"),
    ("2004年", "○年"),
    # 地名（特定情報）
    ("カラジ", "○○"),
    # 理由書日付欄（[37]）— 個別トークン。上で長い日付は処理済み
    ("2023年", "○年"),
    ("11月", "○月"),
    ("１日", "○日"),
    # 残存する国名（最後にまとめて）
    ("イラン", "○○"),
]

doc = Document(dst)

def apply_para(p):
    text = p.text
    new = text
    for old, rep in replacements:
        new = new.replace(old, rep)
    if new != text:
        # 段落先頭runの書式を維持して全文を差し替え
        if p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new)

for p in doc.paragraphs:
    apply_para(p)

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                apply_para(p)

doc.save(dst)
print("SAVED:", dst)

# 検証出力
doc2 = Document(dst)
for i, p in enumerate(doc2.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text}")
