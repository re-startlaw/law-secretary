"""馬さん案件・受任通知 中国語版_3 を作成（修正履歴付き）。

仕様:
- ベース: 既存 260512電子内容証明案_中国語版.docx を _中国語版_3 にコピー済み
- 最終形は _中国語版_2 と同じ（_4 反映済み）
- 旧中国語版からの差分を Word 変更履歴形式（w:ins/w:del）で表示
- 第４損害項目の番号繰上げは番号のみ差分、(7)律师费→(9)末尾移動は段落削除＋新規挿入で表現
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
NSMAP = {"w": W_NS}

AUTHOR = "米谷尚起"
DATE = "2026-05-12T00:00:00Z"

INPUT = Path(
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/"
    "共有用/01_事件記録/ま_馬さん/Wordファイル/260512電子内容証明案_中国語版_3.docx"
)


class IdGen:
    def __init__(self) -> None:
        self.n = 3000

    def next(self) -> str:
        self.n += 1
        return str(self.n)


IDS = IdGen()


def make_run(text: str, base_rpr: etree._Element | None = None) -> etree._Element:
    r = etree.Element(f"{W}r")
    if base_rpr is not None:
        r.append(copy.deepcopy(base_rpr))
    t = etree.SubElement(r, f"{W}t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def make_del_run(text: str, base_rpr: etree._Element | None = None) -> etree._Element:
    r = etree.Element(f"{W}r")
    if base_rpr is not None:
        r.append(copy.deepcopy(base_rpr))
    t = etree.SubElement(r, f"{W}delText")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def wrap_ins(runs: list[etree._Element]) -> etree._Element:
    ins = etree.Element(f"{W}ins")
    ins.set(f"{W}id", IDS.next())
    ins.set(f"{W}author", AUTHOR)
    ins.set(f"{W}date", DATE)
    for r in runs:
        ins.append(r)
    return ins


def wrap_del(runs: list[etree._Element]) -> etree._Element:
    d = etree.Element(f"{W}del")
    d.set(f"{W}id", IDS.next())
    d.set(f"{W}author", AUTHOR)
    d.set(f"{W}date", DATE)
    for r in runs:
        d.append(r)
    return d


def extract_first_rpr(p_elem: etree._Element) -> etree._Element | None:
    return p_elem.find(f".//{W}r/{W}rPr", namespaces=NSMAP)


def replace_paragraph_content(p_elem: etree._Element, new_text: str) -> None:
    base_rpr = extract_first_rpr(p_elem)
    old_text_parts: list[str] = []
    children_to_remove: list[etree._Element] = []
    for child in list(p_elem):
        if child.tag == f"{W}pPr":
            continue
        if child.tag == f"{W}r":
            for t in child.findall(f"{W}t", namespaces=NSMAP):
                old_text_parts.append(t.text or "")
        elif child.tag == f"{W}ins":
            for r in child.findall(f"{W}r", namespaces=NSMAP):
                for t in r.findall(f"{W}t", namespaces=NSMAP):
                    old_text_parts.append(t.text or "")
        children_to_remove.append(child)
    for c in children_to_remove:
        p_elem.remove(c)
    old_text = "".join(old_text_parts)
    if old_text:
        p_elem.append(wrap_del([make_del_run(old_text, base_rpr)]))
    if new_text:
        p_elem.append(wrap_ins([make_run(new_text, base_rpr)]))


def delete_paragraph_as_tracked(p_elem: etree._Element) -> None:
    base_rpr = extract_first_rpr(p_elem)
    old_text_parts: list[str] = []
    children_to_remove: list[etree._Element] = []
    pPr = p_elem.find(f"{W}pPr", namespaces=NSMAP)
    for child in list(p_elem):
        if child.tag == f"{W}pPr":
            continue
        if child.tag == f"{W}r":
            for t in child.findall(f"{W}t", namespaces=NSMAP):
                old_text_parts.append(t.text or "")
        children_to_remove.append(child)
    for c in children_to_remove:
        p_elem.remove(c)
    old_text = "".join(old_text_parts)
    if old_text:
        p_elem.append(wrap_del([make_del_run(old_text, base_rpr)]))
    if pPr is None:
        pPr = etree.SubElement(p_elem, f"{W}pPr")
        p_elem.insert(0, pPr)
    rPr = pPr.find(f"{W}rPr", namespaces=NSMAP)
    if rPr is None:
        rPr = etree.SubElement(pPr, f"{W}rPr")
    del_marker = etree.Element(f"{W}del")
    del_marker.set(f"{W}id", IDS.next())
    del_marker.set(f"{W}author", AUTHOR)
    del_marker.set(f"{W}date", DATE)
    rPr.insert(0, del_marker)


def insert_paragraph_after_as_tracked(
    after_p_elem: etree._Element, text: str
) -> etree._Element:
    base_pPr = after_p_elem.find(f"{W}pPr", namespaces=NSMAP)
    base_rpr = extract_first_rpr(after_p_elem)
    new_p = etree.Element(f"{W}p")
    if base_pPr is not None:
        new_pPr = copy.deepcopy(base_pPr)
        new_p.append(new_pPr)
        rPr = new_pPr.find(f"{W}rPr", namespaces=NSMAP)
        if rPr is None:
            rPr = etree.SubElement(new_pPr, f"{W}rPr")
        ins_marker = etree.Element(f"{W}ins")
        ins_marker.set(f"{W}id", IDS.next())
        ins_marker.set(f"{W}author", AUTHOR)
        ins_marker.set(f"{W}date", DATE)
        rPr.insert(0, ins_marker)
    if text:
        new_p.append(wrap_ins([make_run(text, base_rpr)]))
    after_p_elem.addnext(new_p)
    return new_p


def renumber_paragraph(p_elem: etree._Element, old_num: str, new_num: str) -> None:
    for child in list(p_elem):
        if child.tag != f"{W}r":
            continue
        for t in child.findall(f"{W}t", namespaces=NSMAP):
            text = t.text or ""
            if old_num in text:
                idx = text.index(old_num)
                prefix = text[:idx]
                suffix = text[idx + len(old_num):]
                t.text = prefix
                t.set(
                    "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                )
                run_rpr = child.find(f"{W}rPr", namespaces=NSMAP)
                del_elem = wrap_del([make_del_run(old_num, run_rpr)])
                ins_elem = wrap_ins([make_run(new_num, run_rpr)])
                suffix_r = etree.Element(f"{W}r")
                if run_rpr is not None:
                    suffix_r.append(copy.deepcopy(run_rpr))
                st = etree.SubElement(suffix_r, f"{W}t")
                st.text = suffix
                st.set(
                    "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                )
                child.addnext(suffix_r)
                child.addnext(ins_elem)
                child.addnext(del_elem)
                if prefix == "":
                    p_elem.remove(child)
                return
    raise ValueError(f"renumber: '{old_num}' not found")


def main() -> None:
    doc = Document(str(INPUT))
    paragraphs = doc.paragraphs

    # ---- 第２(1) 見出し ----
    p24 = paragraphs[24]
    assert "抓拽Angelina女士衣领将其拖移之行为" in (p24.text or "")
    replace_paragraph_content(
        p24._element,
        "（１）　小牧女士对Angelina女士衣领附近施以身体接触并使其移动、于跑道旁将其停止之行为",
    )

    # ---- 第２(1) 本文：時間・暴力性 ----
    p25 = paragraphs[25]
    assert p25.text.startswith("　2026年（令和8年）4月13日上午9时18分左右")
    replace_paragraph_content(
        p25._element,
        (
            "　2026年（令和8年）4月13日上午9时20分左右至9时30分左右之间，于贵校体育馆内，"
            "小牧女士从背后对Angelina女士衣领附近施以身体接触并将其抓住，使Angelina女士之"
            "身体被带动移动，于跑道旁将其停止，同时大声叱责。该事实，根据Angelina女士同日"
            "书面陈述、本方所持之事案当日影像等，足以认定。"
        ),
    )

    # ---- 第３ 数学授業 ----
    p34 = paragraphs[34]
    assert p34.text.startswith("　Angelina女士承认于2026年4月13日第一节算数课时间内")
    replace_paragraph_content(
        p34._element,
        (
            "　Angelina女士承认于2026年4月13日第一节数学课时间（上午8时40分至9时35分）内，"
            "曾前往体育馆。此乃Angelina女士在较早完成自身课题之后，为观看长男George先生"
            "约4分钟之越野跑（cross country）比赛，而前往体育馆。该时间，前半为教师讲课时间，"
            "后半为各学生进行课题之自习时间；且于该自习时间内学生暂时离开教室，亦系现场教师"
            "所默认。"
        ),
    )

    # ---- 第４(2) 精神状態悪化 ----
    p40 = paragraphs[40]
    assert "Angelina女士抑郁症恶化" in (p40.text or "")
    replace_paragraph_content(
        p40._element,
        (
            "（２）　因受贵校不当应对、无法到校、及违背事实地散布「Angelina女士在进行虚假"
            "陈述」之言论，导致Angelina女士精神状态恶化、需进行药物调整及继续接受心理咨询"
            "之状态，由此产生之精神科诊疗费及心理咨询费"
        ),
    )

    # ---- 第４(5) 妻関連 → 削除 ----
    p43 = paragraphs[43]
    assert "通知人之妻每日为George先生前往贵校" in (p43.text or "")
    delete_paragraph_as_tracked(p43._element)

    # ---- 第４(6) 通学定期券 → (5) に番号繰上げ（番号のみ差分） ----
    p_kippu = next(
        p for p in doc.paragraphs if "通学定期券之损失" in (p.text or "")
    )
    renumber_paragraph(p_kippu._element, "（６）", "（５）")

    # ---- 第４(7) 律师费 → 段落削除（その後、末尾(9)として新規挿入する） ----
    p_fee = next(
        p
        for p in doc.paragraphs
        if "通知人为本件法律应对所需之律师费" in (p.text or "")
    )
    delete_paragraph_as_tracked(p_fee._element)

    # ---- 新規 (6)(7)(8)(9) を通学定期券(5)の直後に挿入 ----
    # 通学定期券 段落は削除されておらず残っている
    anchor = p_kippu._element
    new_items = [
        "（６）　因Angelina女士长期处于学籍中断状态，导致Angelina女士学业连续性受到毁损所产生之损害",
        "（７）　对Angelina女士考试成绩、GPA及Transcript（成绩单）造成不利影响所产生之损害",
        "（８）　对Angelina女士升学路径造成现实性影响，以及教育机会丧失及未来升学风险所产生之损害",
        "（９）　通知人为本件法律应对所需之律师费及其他费用",
    ]
    for txt in new_items:
        anchor = insert_paragraph_after_as_tracked(anchor, txt)

    # ---- 第５(2) ----
    p_fukugaku = next(
        p
        for p in doc.paragraphs
        if (p.text or "").strip() == "（２）　立即允许Angelina女士复学"
    )
    replace_paragraph_content(
        p_fukugaku._element,
        "（２）　在确保对Angelina女士公平且安全之教育环境之基础上，实现Angelina女士之复学",
    )

    # ---- 第６(3) 学費相当額 → 1,388,000日元 ----
    p_gakuhi = next(
        p for p in doc.paragraphs if "学费相当额（金○○万日元）" in (p.text or "")
    )
    replace_paragraph_content(
        p_gakuhi._element,
        (
            "（３）　然而，Angelina女士及通知人强烈希望继续与贵校之在学合同。本方将学费"
            "相当额（金1,388,000日元。相当于Angelina女士G10年级第一期学费。）保全于"
            "本律师之律师预收金账户，待贵校恢复提供教育役务之同时，立即支付该金额。"
        ),
    )

    doc.save(str(INPUT))
    print(f"saved: {INPUT}")


if __name__ == "__main__":
    main()
