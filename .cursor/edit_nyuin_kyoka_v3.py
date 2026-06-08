# -*- coding: utf-8 -*-
"""
260511入院許可申請_2.docx をベースに _3 を作成し、次の修正を加える。
  (1) 第１の「大阪以下薬科大学病院」→「大阪医科薬科大学病院」
  (2) 第２の各項冒頭の「１　」「２　」… を削除（番号なしの段落に）
  (3) 末尾の「別紙」段落を削除
レイアウトは維持する。
"""
import shutil
import re
from pathlib import Path
from docx import Document

BASE = Path(
    "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/"
    ".shortcut-targets-by-id/1pEJgbHuaj4L4ovFEISOuq-0LeHGInxYl/"
    "大中 忠生 : 詐欺未遂/ワードファイル/260511入院許可申請_2.docx"
)
DST = BASE.with_name("260511入院許可申請_3.docx")

# 段落番号プレフィックス（全角数字 + 全角スペース）
NUM_PREFIX = re.compile(r"^[０-９一二三四五六七八九十]+[　\s]+")


def main():
    target = DST
    n = 3
    while target.exists():
        n += 1
        target = BASE.with_name(f"260511入院許可申請_{n}.docx")
    shutil.copy2(BASE, target)
    print(f"コピー作成: {target}")

    doc = Document(target)

    # 第２ と 附属書類 の範囲を特定
    paras = doc.paragraphs
    idx_dai2 = next(i for i, p in enumerate(paras) if p.text.strip().startswith("第２"))
    idx_appendix = next(i for i, p in enumerate(paras) if p.text.strip() == "附属書類")

    # (2) 段落番号削除
    for i in range(idx_dai2 + 1, idx_appendix):
        p = paras[i]
        txt = p.text
        new_txt = NUM_PREFIX.sub("", txt)
        if new_txt != txt:
            # 既存 run を全削除し、1 つの run にまとめて書き直す
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            for r in list(p._p.findall(f'{ns}r')):
                p._p.remove(r)
            run = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = new_txt
            t.set(qn('xml:space'), 'preserve')
            run.append(t)
            p._p.append(run)
            print(f"  番号除去 [{i}]: {new_txt[:30]}…")

    # (1) 第１ の入院先表記の誤字訂正
    for i, p in enumerate(doc.paragraphs):
        if "大阪以下薬科大学病院" in p.text:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            new_txt = p.text.replace("大阪以下薬科大学病院", "大阪医科薬科大学病院")
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            for r in list(p._p.findall(f'{ns}r')):
                p._p.remove(r)
            run = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = new_txt
            t.set(qn('xml:space'), 'preserve')
            run.append(t)
            p._p.append(run)
            print(f"  誤字訂正 [{i}]: {new_txt}")

    # (3) 末尾の「別紙」段落を削除（後続の空段落も含めて整理）
    # 「以上」の後に出てくる「別紙」とその後ろの段落をすべて削除する。
    paras = doc.paragraphs
    idx_ijou = next(i for i, p in enumerate(paras) if p.text.strip() == "以上")
    # 「別紙」段落を探索
    bessi_indices = [i for i, p in enumerate(paras) if i > idx_ijou and p.text.strip() == "別紙"]
    if bessi_indices:
        idx_bessi = bessi_indices[0]
        # 別紙 およびそれ以降のすべての段落を削除
        to_remove = paras[idx_bessi:]
        for p in to_remove:
            p._p.getparent().remove(p._p)
        print(f"  別紙以降 {len(to_remove)} 段落を削除")

    doc.save(target)
    print(f"保存完了: {target}")
    print(f"FINAL_PATH={target}")


if __name__ == "__main__":
    main()
