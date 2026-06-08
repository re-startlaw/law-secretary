"""馬強事件 損害賠償請求書 docx v3 生成。

修正指示（260522_馬氏4.md）：
- タイトル「損害賠償請求書」を太字16pt
- 差出人ブロックは左寄せで右端ギリギリ（既にユーザー編集で first_line_indent=3600450 適用済）
- 「ア　令和８年〜」のような箇条書き段落の2行目以降を1行目本文位置に揃える（hanging indent）
- タイトル以外の段落・表全体を1字下げ（left_indent=1字）
- 振込先口座を実口座に変更
  法人第二営業部 普通 2469137 弁護士法人Re-Start法律事務所 預かり口

学習点（_2 からの差分）：
- ユーザーは差出人を alignment=LEFT のまま first_line_indent=3600450 EMU
  で右寄りに配置していた → これを保持
- タイトル run.font.size を既に 16pt(203200 EMU)に設定 → bold=True を追加
"""

import os
import shutil
import re
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260522_損害賠償請求書（馬様）_2.docx"
)
DST = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260522_損害賠償請求書（馬様）_3.docx"
)

CHAR_EMU = 133350  # ≒1文字（10.5pt）


def rep(para, text):
    if para.runs:
        first = para.runs[0]
        first.text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def is_bullet_a_kou(text):
    """ア・イ・ウ・エ・オ・カで始まる箇条書き行か判定。"""
    t = text.strip()
    if not t:
        return False
    if t[0] in "アイウエオカキクケコ" and (len(t) >= 2 and t[1] in "　 "):
        return True
    return False


def set_table_indent(table, twips):
    """テーブル全体の左インデントを設定。"""
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblInd"))
    if existing is not None:
        tblPr.remove(existing)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(twips))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    paras = list(doc.paragraphs)

    # ---- 1. タイトル（Para 14: 損害賠償請求書）を太字16pt ----
    title = paras[14]
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(16)

    # ---- 2. 全段落の left_indent を 1字（133350 EMU）に設定 ----
    # 除外：ヘッダー要素（日付・宛先・差出人・タイトル・空行・以上）
    HEADER_INDICES = set(range(0, 16))  # 0:日付〜15:タイトル後の空行
    # 「以上」段落と末尾空行も除外
    for i, p in enumerate(paras):
        if i in HEADER_INDICES:
            continue
        text = p.text
        if text.strip() == "以上":
            continue
        # 末尾の空行（以上の後の段落）は left_indent なし
        # 通常本文・見出し・箇条書きは left_indent = 1字
        p.paragraph_format.left_indent = Emu(CHAR_EMU)

    # ---- 3. 箇条書きア〜カの hanging indent ----
    # 「ア　令和８年〜」など、left_indent = 3字、first_line_indent = -2字
    # 結果：1行目は1字（=本文と同じ）、2行目以降は3字位置で「ア　」の後ろに揃う
    for p in doc.paragraphs:
        if is_bullet_a_kou(p.text):
            p.paragraph_format.left_indent = Emu(CHAR_EMU * 3)
            p.paragraph_format.first_line_indent = Emu(-CHAR_EMU * 2)

    # ---- 4. 表の left_indent を 1字に ----
    # twips = EMU / 635 (1 twip = 1/20 pt, 1 pt = 12700 EMU)
    # 1 char ≈ 10.5pt = 210 twips, but we want consistent 133350 EMU
    # 133350 EMU / 12700 = 10.5 pt, in twips = 10.5 * 20 = 210 twips
    twips_1char = 210
    for t in doc.tables:
        set_table_indent(t, twips_1char)

    # ---- 5. 振込先情報を実口座に更新 ----
    new_bank = [
        "　金融機関：ＧＭＯあおぞらネット銀行（１０２）法人第二営業部支店",
        "　口座種別・口座番号：普通　２４６９１３７",
        "　口座名義：弁護士法人Ｒｅ－Ｓｔａｒｔ法律事務所　預かり口",
    ]
    bank_iter = iter(new_bank)
    for p in doc.paragraphs:
        if p.text.startswith("　　金融機関："):
            try:
                rep(p, next(bank_iter))
            except StopIteration:
                pass
        elif p.text.startswith("　　口座種別・口座番号："):
            try:
                rep(p, next(bank_iter))
            except StopIteration:
                pass
        elif p.text.startswith("　　口座名義："):
            try:
                rep(p, next(bank_iter))
            except StopIteration:
                pass

    doc.save(DST)
    print("saved:", DST)


if __name__ == "__main__":
    main()
