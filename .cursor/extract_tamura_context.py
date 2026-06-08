#!/usr/bin/env python3
"""田村案件の事件情報を抽出（PDFとdocx）"""
import os
import sys
import subprocess

with open(".cursor/shell_path_utf8.txt", encoding="utf-8") as f:
    paths = [ln.strip() for ln in f if ln.strip()]

for p in paths:
    print(f"\n========== {os.path.basename(p)} ==========")
    if p.endswith(".pdf"):
        try:
            res = subprocess.run(
                ["venv/bin/python", "-c",
                 "import sys; from pypdf import PdfReader;"
                 " r=PdfReader(sys.argv[1]);"
                 " [print('---p',i+1,'---') or print(pg.extract_text() or '') for i,pg in enumerate(r.pages)]",
                 p],
                capture_output=True, text=True, timeout=60,
            )
            print(res.stdout)
            if res.returncode != 0:
                print("STDERR:", res.stderr, file=sys.stderr)
        except Exception as e:
            print("ERR pdf:", e)
    elif p.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(p)
            for i, par in enumerate(doc.paragraphs):
                print(par.text)
            for ti, t in enumerate(doc.tables):
                print(f"--Table {ti}--")
                for r in t.rows:
                    print(" | ".join(c.text for c in r.cells))
        except Exception as e:
            print("ERR docx:", e)
