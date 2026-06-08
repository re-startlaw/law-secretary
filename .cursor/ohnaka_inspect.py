# -*- coding: utf-8 -*-
import os
from docx import Document

FOLDER = '/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/.shortcut-targets-by-id/1pEJgbHuaj4L4ovFEISOuq-0LeHGInxYl/大中 忠生 : 詐欺未遂/資料/★大中さんと共有/証拠アップロード用フォルダ/京阪事務機器/提出正式版'
REPORT = '/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/.shortcut-targets-by-id/1pEJgbHuaj4L4ovFEISOuq-0LeHGInxYl/大中 忠生 : 詐欺未遂/ワードファイル/260527_報告書(3)（案）  （京阪事務機器株式会社に関する資料）  .docx'

print('=== FOLDER EXISTS:', os.path.isdir(FOLDER), '===')
if os.path.isdir(FOLDER):
    for root, dirs, files in os.walk(FOLDER):
        rel = os.path.relpath(root, FOLDER)
        print('[DIR]', rel)
        for f in sorted(files):
            full = os.path.join(root, f)
            sz = os.path.getsize(full)
            print('   ', repr(f), sz)

print()
print('=== REPORT EXISTS:', os.path.isfile(REPORT), '===')
if os.path.isfile(REPORT):
    doc = Document(REPORT)
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            print(i, t)
    # tables
    for ti, tbl in enumerate(doc.tables):
        print('--- TABLE', ti, '---')
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            print(' | '.join(cells))
