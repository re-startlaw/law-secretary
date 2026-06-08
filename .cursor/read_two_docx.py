#!/usr/bin/env python3
"""Read both docx files and dump their content."""
from docx import Document
from pathlib import Path

BASE = Path("/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/た_田村正宣/09_ワードファイル")

files = [
    BASE / "聴取報告書.docx",
    BASE / "260512上申書.docx",
]

for f in files:
    print("=" * 80)
    print(f"FILE: {f.name}")
    print("=" * 80)
    if not f.exists():
        print("(file not found)")
        continue
    doc = Document(str(f))
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ""
        align = p.alignment
        print(f"[{i:03d}][{style}][align={align}] {p.text}")
    print()
    for ti, t in enumerate(doc.tables):
        print(f"--- TABLE {ti} ---")
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                print(f"  R{ri}C{ci}: {cell.text}")
    print()
