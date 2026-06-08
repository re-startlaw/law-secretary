"""段依辰示談書（案）_2 作成スクリプト
変更点:
1. 第７（接触禁止）に本示談書に基づく履行に必要な連絡を除外する例外を追加
2. 末尾（清算条項の次）に「許し・刑事処罰を求めない」条項を追加
"""
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SRC = Path("/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/ダ_ダンイシン/260519_示談書（案）.docx")
DST = SRC.with_name("260519_示談書（案）_2.docx")

shutil.copy(SRC, DST)
doc = Document(DST)

# 1) 接触禁止条項の差し替え（例外追加）
new_contact_text = (
    "　乙は、本示談書に基づく履行に必要な連絡を除き、今後一切、甲に対し、直接又は第三者を介して接触"
    "（面会、電話、ＳＭＳ、電子メール、ＷｅＣｈａｔを含むメッセンジャーアプリ、ＳＮＳその他一切の手段による連絡を含む。）"
    "してはならない。乙がこれに違反した場合、乙は、違反一回につき違約金として金〇円を甲に対し支払う。"
)

contact_par = None
seisan_par = None
for p in doc.paragraphs:
    if p.style.name == "ランク２" and p.text.startswith("　乙は、今後一切、甲に対し、直接又は第三者を介して接触"):
        contact_par = p
    if p.style.name == "ランク２" and p.text.startswith("　甲及び乙は、本示談書に定めるもののほか"):
        seisan_par = p

assert contact_par is not None, "接触禁止段落が見つかりません"
assert seisan_par is not None, "清算条項段落が見つかりません"

# Replace text in contact paragraph: clear existing runs, add a new run with same style
# Preserve the paragraph's pPr (style/numbering)
for run in list(contact_par.runs):
    run._r.getparent().remove(run._r)
contact_par.add_run(new_contact_text)

# 2) 「許し・刑事処罰を求めない」条項を清算条項の直後に挿入
forgive_text = (
    "　甲は、乙の第１項から第８項までの乙の対応に鑑み、本件について乙を許し、乙の刑事処罰を求めない。"
)

new_p = deepcopy(seisan_par._p)
# Clear all child elements except pPr (which carries the ランク２ style + numbering)
for child in list(new_p):
    if child.tag != qn("w:pPr"):
        new_p.remove(child)
# Insert new paragraph right after seisan_par
seisan_par._p.addnext(new_p)

# Add a run with the new text to new_p
from docx.oxml import OxmlElement
r = OxmlElement("w:r")
t = OxmlElement("w:t")
t.text = forgive_text
t.set(qn("xml:space"), "preserve")
r.append(t)
new_p.append(r)

doc.save(DST)
print(f"saved: {DST}")
