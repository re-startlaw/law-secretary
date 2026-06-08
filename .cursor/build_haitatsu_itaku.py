"""書類配達アルバイト向け業務委託契約書（A4 2枚以内）を作成する。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = Path(
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/"
    "業務委託契約書_書類配達_アルバイト.docx"
)


def set_font(run, size=10.5, bold=False):
    run.font.name = "ＭＳ 明朝"
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    rFonts.set(qn("w:ascii"), "ＭＳ 明朝")
    rFonts.set(qn("w:hAnsi"), "ＭＳ 明朝")


def add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=10.5,
             space_after=2, indent_cm=None, first_line_indent_cm=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    if indent_cm is not None:
        pf.left_indent = Cm(indent_cm)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p


def add_article(doc, heading, body_lines):
    add_para(doc, heading, bold=True, space_after=2)
    for line in body_lines:
        add_para(doc, line, first_line_indent_cm=0.0, space_after=2)


def add_numbered_items(doc, items):
    """項番号付きの箇条書き（１．２．… 形式、手書きで安定運用）。"""
    nums = ["１", "２", "３", "４", "５", "６", "７", "８", "９"]
    for i, text in enumerate(items):
        add_para(
            doc,
            f"{nums[i]}．{text}",
            indent_cm=0.7,
            first_line_indent_cm=-0.7,
            space_after=2,
        )


def build():
    doc = Document()

    # ページ余白を詰めて2枚に収める
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # 標準スタイル
    style = doc.styles["Normal"]
    style.font.name = "ＭＳ 明朝"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "ＭＳ 明朝")

    # タイトル
    add_para(doc, "業務委託契約書", align=WD_ALIGN_PARAGRAPH.CENTER,
             bold=True, size=14, space_after=8)

    # 前文
    add_para(
        doc,
        "弁護士法人Re-Start法律事務所（以下「甲」という。）と"
        "　　　　　　　　　（以下「乙」という。）は、"
        "甲の依頼する法律書面の配達業務に関し、次のとおり業務委託契約（以下「本契約」という。）"
        "を締結する。",
        first_line_indent_cm=0.0,
        space_after=4,
    )

    # 第1条
    add_article(
        doc,
        "第１条（業務内容）",
        ["　甲は、乙に対し、甲が指定する法律書面（訴状、準備書面、申立書、証拠書類等）"
         "を、甲が指定する裁判所、検察庁その他の機関に持参・提出する業務"
         "（以下「本業務」という。）を委託し、乙はこれを受託する。"],
    )

    # 第2条
    add_article(
        doc,
        "第２条（業務の遂行）",
        [],
    )
    add_numbered_items(
        doc,
        [
            "乙は、善良な管理者の注意をもって、甲の指示に従い本業務を遂行する。",
            "乙は、書面の受領後、甲が指定する期限までに指定の届出先へ確実に届けるものとし、"
            "届出完了後速やかに甲に報告する。",
        ],
    )

    # 第3条
    add_article(
        doc,
        "第３条（報酬及び交通費）",
        [],
    )
    add_numbered_items(
        doc,
        [
            "本業務の報酬は、１件（届出先１か所）あたり金２，５００円（消費税込み）とする。"
            "交通費その他本業務の遂行に通常要する実費は、当該報酬に含むものとする。",
            "甲は、毎月末日締めで前項の報酬を集計し、翌月末日までに、"
            "乙の指定する銀行口座への振込みにより支払う。振込手数料は甲の負担とする。",
        ],
    )

    # 第4条（守秘義務・書類の取扱い）
    add_article(
        doc,
        "第４条（守秘義務及び書類の取扱い）",
        [],
    )
    add_numbered_items(
        doc,
        [
            "乙は、本業務に関連して甲から預かり、又は知り得た書類、データ、依頼者・関係者の情報"
            "その他一切の情報（以下「秘密情報」という。）を、本業務遂行の目的以外に使用してはならず、"
            "甲の事前の書面による承諾なく、第三者に開示・漏えい・複製してはならない。",
            "乙は、預かった書類を移動中及び待機中を含め第三者の目に触れないよう、封筒、書類袋"
            "又はファイル等で常時遮蔽するとともに、車内、公共の場、飲食店その他の場所に放置してはならない。",
            "乙は、預かった書類を紛失、破損、盗難、誤配その他の事故が生じ、又はそのおそれが生じたときは、"
            "直ちに甲に報告し、甲の指示に従う。",
            "本条の義務は、本契約終了後も存続する。",
        ],
    )

    # 第5条（再委託の禁止）
    add_article(
        doc,
        "第５条（再委託の禁止）",
        ["　乙は、甲の事前の書面による承諾なく、本業務の全部又は一部を第三者に再委託してはならない。"],
    )

    # 第6条（契約期間）
    add_article(
        doc,
        "第６条（契約期間）",
        ["　本契約の有効期間は、契約締結日から１年間とする。期間満了の１か月前までに、"
         "甲又は乙のいずれからも書面による別段の意思表示がないときは、本契約は同一条件でさらに１年間"
         "更新されるものとし、以後も同様とする。"],
    )

    # 第7条（解除）
    add_article(
        doc,
        "第７条（解除）",
        ["　甲又は乙は、相手方が本契約に違反したとき、又は信頼関係を著しく損なう事由が生じたときは、"
         "何らの催告を要せず、直ちに本契約を解除することができる。"],
    )

    # 第8条（損害賠償）
    add_article(
        doc,
        "第８条（損害賠償）",
        ["　乙は、故意又は過失により、甲、甲の依頼者又は第三者に損害を与えたときは、"
         "その損害（弁護士費用を含む。）を賠償する責を負う。"],
    )

    # 第9条（協議）
    add_article(
        doc,
        "第９条（協議）",
        ["　本契約に定めのない事項又は本契約の解釈について疑義が生じた事項については、"
         "甲乙誠実に協議のうえ解決する。"],
    )

    add_para(doc, "", space_after=4)
    add_para(
        doc,
        "本契約締結の証として、本書２通を作成し、甲乙署名又は記名押印のうえ、各１通を保有する。",
        first_line_indent_cm=0.0,
        space_after=6,
    )

    # 日付（右寄せ）
    add_para(doc, "２０２６年　　月　　日", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)

    # 署名欄
    add_para(doc, "甲：東京都豊島区東池袋三丁目１番１号　サンシャイン６０　１２階",
             first_line_indent_cm=0.0, space_after=2)
    add_para(doc, "　　　弁護士法人Re-Start法律事務所", first_line_indent_cm=0.0, space_after=2)
    add_para(doc, "　　　代表社員弁護士　米　谷　尚　起　　　　　　　　　　　　　　印",
             first_line_indent_cm=0.0, space_after=8)

    add_para(doc, "乙：住所　　　　　　　　　　　　　　　　　　　　　　　　",
             first_line_indent_cm=0.0, space_after=2)
    add_para(doc, "　　　氏名　　　　　　　　　　　　　　　　　　　　　　　　印",
             first_line_indent_cm=0.0, space_after=2)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
