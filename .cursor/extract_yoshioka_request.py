#!/usr/bin/env python3
"""吉岡氏の証拠開示請求書docxを読み取り、テキスト出力"""
from docx import Document

src = "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/ん_終了/よ_吉岡和伸/wordファイル/260114証拠開示請求（提出）.docx"

doc = Document(src)

print("=== Paragraphs ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "?"
    align = p.alignment
    print(f"[{i:03d}|{style}|align={align}] {p.text}")

print()
print("=== Tables ===")
for ti, table in enumerate(doc.tables):
    print(f"--- Table {ti} ---")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            print(f"  [r{ri}c{ci}] {cell.text}")

print()
print("=== Sections ===")
for si, sec in enumerate(doc.sections):
    print(f"sec{si}: page={sec.page_width}x{sec.page_height} margins L{sec.left_margin} R{sec.right_margin} T{sec.top_margin} B{sec.bottom_margin}")
