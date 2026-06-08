"""社宅導入関連書類一式を作成する。

出力先：
  /Users/kometaninaoki/Library/CloudStorage/GoogleDrive-.../14_弁護士法人/01_定款・規定/規定/

作成ファイル：
  - 260506_社宅管理規程.docx
  - 260506_社員総会議事録_社宅規程制定.docx
  - 260506_社宅使用契約書.docx
  - 260506_賃貸料相当額計算書.xlsx
  - 260506_社宅導入オペレーション手順書.docx
"""
from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

# ---------------------------------------------------------------------------
# 定数（全書類で共有）
# ---------------------------------------------------------------------------
DEST_DIR = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/"
    "14_弁護士法人/01_定款・規定/規定"
)
PREFIX = "260506_"

# 法人情報
HOJIN_NAME = "弁護士法人Re-Start法律事務所"
HONTEN = "東京都豊島区東池袋三丁目1番1号 サンシャイン60 12階"
DAIHYO = "米谷 尚起"
SETSURITSU = "令和8年1月23日"
SHUSHIKIN = "10,000円"
JIGYO_NENDO = "毎年1月1日から12月31日まで"

# 物件情報
BUKKEN_MEI = "プリア常盤台パークフロント"
BUKKEN_HEYA = "101号室"
BUKKEN_JUSHO = "東京都板橋区常盤台四丁目8番10号"
BUKKEN_KOZO = "鉄筋コンクリート造（5階建）"
BUKKEN_MENSEKI = "61.36㎡"
BUKKEN_MADORI = "3LDK"
BUKKEN_TAIYO = 47  # 法定耐用年数（年）
BUKKEN_CHIKUNEN = "2006年7月（築20年）"

# 金額
GETSUGAKU_YACHIN = 195000
GETSUGAKU_YACHIN_KANJI = "金195,000円"
KOEKI = "なし"
SHIYORYO = 29250  # 家賃の15%（仮徴収額）
SHIYORYO_KANJI = "金29,250円"
SHIKIKIN = "賃料の1か月分"
REIKIN = "賃料の1か月分"

# 日付
SAKUSEI_BI = "令和8年5月6日"
SOKAI_BI = "令和8年5月7日"
SEKO_BI = "令和8年5月7日"
KEIYAKU_KAISHI = "令和8年6月1日"

# 計算用
KEISAN_NEMSU = 47  # 耐用年数
KOMOMSUMA_KIJUN = 99  # 耐用年数30年超→99㎡以下


# ---------------------------------------------------------------------------
# 共通ヘルパ
# ---------------------------------------------------------------------------
def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "游明朝"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    # CJK font
    from docx.oxml.ns import qn

    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "游明朝")
    rfonts.set(qn("w:ascii"), "游明朝")
    rfonts.set(qn("w:hAnsi"), "游明朝")


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def _add_centered(doc: Document, text: str, *, bold: bool = False, size: int = 10.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _add_right(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(text)


def _add_para(doc: Document, text: str = "") -> None:
    doc.add_paragraph(text)


def _add_jou(doc: Document, midashi: str, body_lines: list[str]) -> None:
    """条見出し＋本文（複数段落）。"""
    p = doc.add_paragraph()
    run = p.add_run(midashi)
    run.bold = True
    for line in body_lines:
        doc.add_paragraph(line)


def _save(doc: Document, fname: str) -> str:
    path = os.path.join(DEST_DIR, fname)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# 1. 社宅管理規程
# ---------------------------------------------------------------------------
def make_kitei() -> str:
    doc = Document()
    _set_default_font(doc)

    _add_title(doc, "社宅管理規程")
    _add_para(doc)

    _add_jou(
        doc,
        "第1条（目的）",
        [
            f"　本規程は、{HOJIN_NAME}（以下「法人」という。）が"
            "その社員（弁護士法人の業務を執行する社員をいう。以下同じ。）"
            "の住居の安定を図り、もって法人業務の円滑な遂行に資することを目的として、"
            "法人が借り上げ又は所有する住宅（以下「社宅」という。）を社員に貸与する場合の"
            "取扱いを定めるものである。",
        ],
    )

    _add_jou(
        doc,
        "第2条（定義）",
        [
            "　本規程において、次の各号に掲げる用語の意義は、当該各号に定めるところによる。",
            "(1)　社宅　法人が第三者から賃借し、又は所有する住宅であって、社員に居住の用に"
            "供するため貸与するものをいう。",
            "(2)　借上社宅　法人が第三者から賃借する社宅をいう。",
            "(3)　賃貸料相当額　国税庁タックスアンサーNo.2600（役員に社宅などを貸したとき）"
            "に定める計算式により算出した小規模住宅の賃貸料相当額をいう。",
            "(4)　小規模住宅　その建物の法定耐用年数が30年以下のものについては床面積132㎡以下、"
            "30年を超えるものについては床面積99㎡以下の住宅をいう。",
        ],
    )

    _add_jou(
        doc,
        "第3条（対象者）",
        [
            "　社宅貸与の対象者は、法人の社員のうち、法人が業務上必要と認めた者とする。",
        ],
    )

    _add_jou(
        doc,
        "第4条（物件選定基準）",
        [
            "　社宅として貸与する物件は、原則として小規模住宅の要件を満たすものとする。",
            "2　前項の小規模住宅の要件は次の各号のとおりとする。",
            "(1)　建物の法定耐用年数が30年以下のもの　床面積132㎡以下",
            "(2)　建物の法定耐用年数が30年を超えるもの　床面積99㎡以下",
            "3　社宅として貸与する物件は、社員の通勤の便、業務遂行上の必要性その他の事情を"
            "総合的に勘案して法人が選定する。",
        ],
    )

    _add_jou(
        doc,
        "第5条（賃貸借契約）",
        [
            "　借上社宅に係る賃貸借契約は、法人の名義をもって締結する。",
            "2　社員は、法人の事前承認なくして賃貸借契約の解約、変更、転貸又は又貸しをすることが"
            "できない。",
        ],
    )

    _add_jou(
        doc,
        "第6条（費用負担区分）",
        [
            "　借上社宅に係る費用は、次の各号の区分により法人又は社員がそれぞれ負担する。",
            "(1)　法人負担　敷金、礼金、仲介手数料、家賃、共益費、管理費、更新料、"
            "建物に係る火災保険料",
            "(2)　社員負担　光熱水費、通信費、駐車場使用料、引越費用、町内会費、"
            "家財に係る火災保険料、社員の故意又は過失による原状回復費用",
            "2　前項に定めのない費用については、その性質に応じ法人と社員が協議して定める。",
        ],
    )

    _add_jou(
        doc,
        "第7条（社宅使用料）",
        [
            "　法人は、社員に対し社宅を貸与する場合、当該社員から毎月の社宅使用料を徴収する。",
            "2　社宅使用料の額は、次に掲げる金額の合計額（以下「賃貸料相当額」という。）の"
            "100％以上の金額とする。",
            "(1)　その年度の建物の固定資産税の課税標準額×0.2％",
            "(2)　12円×（建物の総床面積（㎡）÷3.3）",
            "(3)　その年度の敷地の固定資産税の課税標準額×0.22％",
            "3　前項にかかわらず、固定資産税の課税標準額が判明するまでの間は、法人と社員の"
            "協議により仮の社宅使用料を定めることができる。この場合、課税標準額が判明し次第、"
            "速やかに前項の計算式により算定した額に改定するものとする。",
            "4　社宅使用料は、毎月の役員報酬から控除する方法により徴収する。",
        ],
    )

    _add_jou(
        doc,
        "第8条（入居手続）",
        [
            "　社宅の貸与を受けようとする社員は、入居予定日までに次の書類を法人に提出するものとする。",
            "(1)　社宅入居届",
            "(2)　社宅使用契約書",
            "(3)　その他法人が必要と認める書類",
        ],
    )

    _add_jou(
        doc,
        "第9条（退去）",
        [
            "　社員は、次の各号のいずれかに該当するときは、社宅を退去しなければならない。",
            "(1)　社員の地位を喪失したとき",
            "(2)　法人と社員との間で社宅使用契約を合意解除したとき",
            "(3)　社員から退去の申出があり、法人がこれを承認したとき",
            "(4)　第10条に定める管理上の義務に重大な違反があったとき",
            "2　社員は、退去にあたり、社員の費用負担により原状回復義務を履行しなければならない。",
            "ただし、通常損耗及び経年劣化に係るものはこの限りでない。",
        ],
    )

    _add_jou(
        doc,
        "第10条（管理上の義務）",
        [
            "　社員は、社宅を善良な管理者の注意をもって使用し、次の事項を遵守しなければならない。",
            "(1)　社宅を居住以外の用途に使用しないこと",
            "(2)　法人の事前承認なくして同居人を入居させないこと",
            "(3)　法人の事前承認なくして模様替え、増改築又は造作の付加をしないこと",
            "(4)　近隣住民との良好な関係を保ち、迷惑行為を行わないこと",
            "(5)　その他賃貸借契約の各条項を遵守すること",
        ],
    )

    _add_jou(
        doc,
        "第11条（改廃）",
        [
            "　本規程の改廃は、社員総会の決議により行う。",
        ],
    )

    _add_jou(
        doc,
        "第12条（適用上の疑義）",
        [
            "　本規程に定めのない事項又は本規程の解釈について疑義が生じたときは、"
            "代表社員がその都度決定する。",
        ],
    )

    _add_para(doc)
    p = doc.add_paragraph()
    run = p.add_run("附　則")
    run.bold = True
    _add_para(doc, f"　本規程は、{SEKO_BI}から施行する。")

    return _save(doc, f"{PREFIX}社宅管理規程.docx")


# ---------------------------------------------------------------------------
# 2. 臨時社員総会議事録
# ---------------------------------------------------------------------------
def make_giji() -> str:
    doc = Document()
    _set_default_font(doc)

    _add_title(doc, "臨時社員総会議事録")
    _add_para(doc)

    _add_para(
        doc,
        f"　{SOKAI_BI}午前10時00分、{HOJIN_NAME}（以下「法人」という。）"
        "の臨時社員総会を法人本店において開催した。",
    )
    _add_para(doc)
    _add_para(doc, "社員の総数　　　　　　　　1名")
    _add_para(doc, "出資金の総額　　　　　　　金10,000円")
    _add_para(doc, "出席社員数　　　　　　　　1名")
    _add_para(doc, "出席社員の出資金の総額　　金10,000円")
    _add_para(doc)
    _add_para(
        doc,
        f"　以上のとおり社員全員出席のうえ、満場一致をもって議長に代表社員{DAIHYO}を選出し、"
        "同人は議長席に着き、開会を宣し直ちに議事に入った。",
    )
    _add_para(doc)

    # 第1号議案
    p = doc.add_paragraph()
    run = p.add_run("第1号議案　社宅管理規程制定の件")
    run.bold = True
    _add_para(
        doc,
        "　議長は、法人の社員に対する社宅貸与の取扱いを明確にするため、"
        "別紙のとおり社宅管理規程を制定したい旨を述べ、その理由及び規程の内容を詳細に説明し、"
        "その可否を諮ったところ、出席社員全員の賛成をもって原案のとおり可決承認された。",
    )
    _add_para(doc, f"　なお、本規程の施行日は{SEKO_BI}とする。")
    _add_para(doc)

    # 第2号議案
    p = doc.add_paragraph()
    run = p.add_run("第2号議案　代表社員に対する社宅貸与の承認の件")
    run.bold = True
    _add_para(
        doc,
        "　議長は、前号議案により制定された社宅管理規程に基づき、"
        f"法人と代表社員{DAIHYO}との間で次のとおり社宅貸与を行うこととしたい旨を述べた。"
        "なお、本件は法人と代表社員との間の取引に該当するため、利益相反取引に準じて"
        "社員総会の承認を求める旨を併せて説明し、その可否を諮ったところ、"
        "出席社員全員の賛成をもって原案のとおり可決承認された。",
    )
    _add_para(doc)
    _add_para(doc, "記")
    _add_para(doc, f"1　貸与物件　{BUKKEN_MEI}{BUKKEN_HEYA}")
    _add_para(doc, f"　　所在地　 {BUKKEN_JUSHO}")
    _add_para(doc, f"　　構　造　 {BUKKEN_KOZO}（法定耐用年数{BUKKEN_TAIYO}年）")
    _add_para(doc, f"　　専有面積　{BUKKEN_MENSEKI}（{BUKKEN_MADORI}）")
    _add_para(doc, f"2　貸与開始日　{KEIYAKU_KAISHI}")
    _add_para(doc, f"3　月額家賃　 {GETSUGAKU_YACHIN_KANJI}（共益費・管理費{KOEKI}）")
    _add_para(
        doc,
        f"4　月額社宅使用料　{SHIYORYO_KANJI}（仮）",
    )
    _add_para(
        doc,
        "　　ただし、固定資産税の課税標準額が判明し次第、社宅管理規程第7条第2項所定の"
        "賃貸料相当額の計算式により正式金額を算定し、必要に応じ別途社員総会の決議を経て改定する。",
    )
    _add_para(doc, "5　法人負担　敷金、礼金、仲介手数料、家賃、共益費、管理費、更新料、建物火災保険料")
    _add_para(doc, "6　社員負担　光熱水費、通信費、駐車場使用料、引越費用、町内会費、家財火災保険料、故意過失の原状回復費用")
    _add_para(doc, "7　徴収方法　毎月の役員報酬から控除")
    _add_para(doc)

    _add_para(
        doc,
        f"　以上をもって本臨時社員総会の議事を終了したので、議長は{SOKAI_BI}午前10時30分閉会を宣した。",
    )
    _add_para(doc)
    _add_para(
        doc,
        "　上記の議事の経過の要領及びその結果を明確にするため、本議事録を作成し、"
        "出席代表社員において次のとおり記名押印する。",
    )
    _add_para(doc)
    _add_right(doc, SOKAI_BI)
    _add_para(doc)
    _add_right(doc, f"{HOJIN_NAME}　臨時社員総会")
    _add_para(doc)
    _add_right(doc, f"議長兼議事録作成者　代表社員　{DAIHYO}　　印")

    return _save(doc, f"{PREFIX}社員総会議事録_社宅規程制定.docx")


# ---------------------------------------------------------------------------
# 3. 社宅使用契約書
# ---------------------------------------------------------------------------
def make_keiyaku() -> str:
    doc = Document()
    _set_default_font(doc)

    _add_title(doc, "社宅使用契約書")
    _add_para(doc)

    _add_para(
        doc,
        f"　{HOJIN_NAME}（以下「法人」という。）と{DAIHYO}（以下「社員」という。）とは、"
        "法人が社員に対し社宅を貸与することについて、次のとおり契約（以下「本契約」という。）を締結する。",
    )
    _add_para(doc)

    _add_jou(
        doc,
        "第1条（目的）",
        [
            "　法人は、社宅管理規程の定めに従い、社員に対し次条に定める物件を社宅として貸与し、"
            "社員はこれを賃借する。",
        ],
    )

    _add_jou(
        doc,
        "第2条（貸与物件の表示）",
        [
            "　貸与物件の表示は、次のとおりとする。",
            f"(1)　建物の名称　{BUKKEN_MEI}",
            f"(2)　部屋番号　　{BUKKEN_HEYA}",
            f"(3)　所在地　　　{BUKKEN_JUSHO}",
            f"(4)　構　造　　　{BUKKEN_KOZO}（法定耐用年数{BUKKEN_TAIYO}年）",
            f"(5)　専有面積　　{BUKKEN_MENSEKI}（{BUKKEN_MADORI}）",
        ],
    )

    _add_jou(
        doc,
        "第3条（使用料）",
        [
            f"　社員は、法人に対し、本社宅の使用料として月額{SHIYORYO_KANJI}を支払う。",
            "2　前項の使用料は、当面の仮金額であり、固定資産税の課税標準額が判明し次第、"
            "社宅管理規程第7条第2項所定の賃貸料相当額の計算式により算定した金額に改定する。",
            "3　使用料は、毎月の役員報酬から控除する方法により徴収する。",
            "4　使用料の改定は、法人と社員の合意により書面で行う。",
        ],
    )

    _add_jou(
        doc,
        "第4条（費用負担区分）",
        [
            "　本社宅に係る費用負担は、次の各号のとおりとする。",
            "(1)　法人負担　敷金、礼金、仲介手数料、家賃、共益費、管理費、更新料、"
            "建物に係る火災保険料",
            "(2)　社員負担　光熱水費、通信費、駐車場使用料、引越費用、町内会費、"
            "家財に係る火災保険料、社員の故意又は過失による原状回復費用",
        ],
    )

    _add_jou(
        doc,
        "第5条（使用上の遵守事項）",
        [
            "　社員は、本社宅を善良な管理者の注意をもって使用し、次の事項を遵守する。",
            "(1)　本社宅を居住以外の用途に使用しないこと",
            "(2)　法人の事前承認なくして同居人を入居させないこと",
            "(3)　法人の事前承認なくして模様替え、増改築又は造作の付加をしないこと",
            "(4)　法人と賃貸人との賃貸借契約の各条項を遵守すること",
            "(5)　近隣住民との良好な関係を保ち、迷惑行為を行わないこと",
        ],
    )

    _add_jou(
        doc,
        "第6条（契約期間）",
        [
            f"　本契約の期間は、{KEIYAKU_KAISHI}から、社員が法人の社員の地位を喪失する日まで又は"
            "本契約が解除される日までとする。",
            "2　法人と賃貸人との賃貸借契約が終了したときは、本契約も同時に終了する。",
        ],
    )

    _add_jou(
        doc,
        "第7条（解除事由）",
        [
            "　法人は、社員が次の各号のいずれかに該当したときは、何らの催告を要することなく"
            "本契約を解除することができる。",
            "(1)　社員が法人の社員の地位を喪失したとき",
            "(2)　社員が本契約又は社宅管理規程に重大な違反をしたとき",
            "(3)　社員が法人と賃貸人との賃貸借契約に違反する行為をしたとき",
            "(4)　その他本契約を継続し難い事由が発生したとき",
            "2　社員が前項各号に該当した場合、社員は法人の指定する日までに本社宅を退去し、"
            "次条所定の原状回復義務を履行する。",
        ],
    )

    _add_jou(
        doc,
        "第8条（原状回復）",
        [
            "　社員は、退去にあたり社員の費用負担により本社宅を原状に回復する。"
            "ただし、通常損耗及び経年劣化に係るものはこの限りでない。",
        ],
    )

    _add_jou(
        doc,
        "第9条（協議事項）",
        [
            "　本契約に定めのない事項又は本契約の解釈について疑義が生じたときは、"
            "法人と社員が誠実に協議して解決する。",
        ],
    )

    _add_para(doc)
    _add_para(
        doc,
        "　本契約締結の証として、本書2通を作成し、法人及び社員が記名押印のうえ、各自1通を保有する。",
    )
    _add_para(doc)
    _add_right(doc, KEIYAKU_KAISHI)
    _add_para(doc)

    _add_para(doc, "（法　人）")
    _add_para(doc, f"　　所在地　{HONTEN}")
    _add_para(doc, f"　　名　称　{HOJIN_NAME}")
    _add_para(doc, f"　　代表者　代表社員　{DAIHYO}　　　印")
    _add_para(doc)
    _add_para(doc, "（社　員）")
    _add_para(doc, f"　　住　所　{HONTEN}")
    _add_para(doc, f"　　氏　名　{DAIHYO}　　　印")

    return _save(doc, f"{PREFIX}社宅使用契約書.docx")


# ---------------------------------------------------------------------------
# 4. 賃貸料相当額計算書（Excel）
# ---------------------------------------------------------------------------
def make_keisansho() -> str:
    wb = Workbook()
    ws = wb.active
    ws.title = "賃貸料相当額計算"

    blue = Font(color="0070C0", name="游明朝", size=11, bold=True)
    black = Font(color="000000", name="游明朝", size=11)
    bold_black = Font(color="000000", name="游明朝", size=11, bold=True)
    title_font = Font(color="000000", name="游明朝", size=14, bold=True)
    note_font = Font(color="595959", name="游明朝", size=9, italic=True)
    header_fill = PatternFill("solid", fgColor="D9E1F2")
    result_fill = PatternFill("solid", fgColor="FFF2CC")
    thin = Side(border_style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center")
    right = Alignment(horizontal="right", vertical="center")
    left_wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # 列幅
    widths = {"A": 4, "B": 38, "C": 18, "D": 6, "E": 38}
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    # タイトル
    ws.merge_cells("B2:E2")
    ws["B2"] = "賃貸料相当額計算書（役員社宅／小規模住宅）"
    ws["B2"].font = title_font
    ws["B2"].alignment = center

    ws.merge_cells("B3:E3")
    ws["B3"] = (
        f"対象物件：{BUKKEN_MEI}{BUKKEN_HEYA}　／　{BUKKEN_JUSHO}　／　"
        f"{BUKKEN_KOZO}　／　専有面積{BUKKEN_MENSEKI}"
    )
    ws["B3"].font = note_font
    ws["B3"].alignment = center

    # 入力欄ヘッダ
    ws["B5"] = "【入力欄】数値を青字で入力"
    ws["B5"].font = bold_black
    ws["B5"].fill = header_fill
    ws.merge_cells("B5:C5")

    inputs = [
        ("建物の固定資産税課税標準額（円）", "", "未取得の場合は空欄でも可"),
        ("敷地の固定資産税課税標準額（円）", "", "未取得の場合は空欄でも可"),
        ("建物の総床面積（㎡）", BUKKEN_MENSEKI.replace("㎡", ""), "登記簿面積。本物件は専有61.36㎡"),
        ("月額家賃（円）", GETSUGAKU_YACHIN, "本物件は195,000円"),
        ("法定耐用年数（年）", BUKKEN_TAIYO, "RC造=47年、木造=22年"),
    ]
    for i, (label, value, memo) in enumerate(inputs):
        row = 6 + i
        ws.cell(row=row, column=2, value=label).font = black
        ws.cell(row=row, column=2).alignment = left_wrap
        ws.cell(row=row, column=2).border = border
        c = ws.cell(row=row, column=3, value=value if value != "" else None)
        c.font = blue
        c.alignment = right
        c.border = border
        c.number_format = "#,##0"
        ws.cell(row=row, column=5, value=memo).font = note_font
        ws.cell(row=row, column=5).alignment = left_wrap

    # B = 建物課税標準額 → C6
    # D = 敷地課税標準額 → C7
    # 床面積 → C8
    # 家賃 → C9
    # 耐用年数 → C10

    # 計算欄ヘッダ
    ws["B12"] = "【計算欄】国税庁No.2600（小規模住宅）の計算式"
    ws["B12"].font = bold_black
    ws["B12"].fill = header_fill
    ws.merge_cells("B12:C12")

    calcs = [
        ("(1) 建物課税標準額 × 0.2％", "=IFERROR(C6*0.002,0)"),
        ("(2) 12円 × （総床面積 ÷ 3.3）", "=IFERROR(12*(C8/3.3),0)"),
        ("(3) 敷地課税標準額 × 0.22％", "=IFERROR(C7*0.0022,0)"),
    ]
    for i, (label, formula) in enumerate(calcs):
        row = 13 + i
        ws.cell(row=row, column=2, value=label).font = black
        ws.cell(row=row, column=2).alignment = left_wrap
        ws.cell(row=row, column=2).border = border
        c = ws.cell(row=row, column=3, value=formula)
        c.font = black
        c.alignment = right
        c.border = border
        c.number_format = "#,##0"

    # 合計
    ws.cell(row=16, column=2, value="月額賃貸料相当額（合計）").font = bold_black
    ws.cell(row=16, column=2).alignment = left_wrap
    ws.cell(row=16, column=2).border = border
    total = ws.cell(row=16, column=3, value="=ROUNDDOWN(C13+C14+C15,0)")
    total.font = bold_black
    total.fill = result_fill
    total.alignment = right
    total.border = border
    total.number_format = "#,##0"

    # 参考：家賃の50％
    ws.cell(row=18, column=2, value="【参考】月額家賃の50％").font = black
    ws.cell(row=18, column=2).alignment = left_wrap
    c = ws.cell(row=18, column=3, value="=IFERROR(C9*0.5,0)")
    c.font = black
    c.alignment = right
    c.number_format = "#,##0"

    # 仮徴収額
    ws.cell(row=19, column=2, value="【参考】月額家賃の15％（本件仮徴収額）").font = black
    ws.cell(row=19, column=2).alignment = left_wrap
    c = ws.cell(row=19, column=3, value="=IFERROR(C9*0.15,0)")
    c.font = black
    c.alignment = right
    c.number_format = "#,##0"

    # 小規模住宅判定
    ws["B21"] = "【判定】小規模住宅該当性"
    ws["B21"].font = bold_black
    ws["B21"].fill = header_fill
    ws.merge_cells("B21:C21")

    ws.cell(row=22, column=2, value="床面積基準（耐用年数30年以下=132㎡／30年超=99㎡）").font = black
    ws.cell(row=22, column=2).alignment = left_wrap
    c = ws.cell(
        row=22,
        column=3,
        value='=IF(C10<=30,132,99)',
    )
    c.font = black
    c.alignment = right
    c.number_format = "#,##0.00"

    ws.cell(row=23, column=2, value="判定結果").font = bold_black
    ws.cell(row=23, column=2).alignment = left_wrap
    c = ws.cell(
        row=23,
        column=3,
        value='=IF(C8<=C22,"該当（小規模住宅）","非該当")',
    )
    c.font = bold_black
    c.fill = result_fill
    c.alignment = center

    # 出典・注記
    ws.merge_cells("B25:E25")
    ws["B25"] = (
        "【出典】国税庁タックスアンサー No.2600 役員に社宅などを貸したとき"
        "（https://www.nta.go.jp/taxes/shiraberu/taxanswer/gensen/2600.htm）"
    )
    ws["B25"].font = note_font
    ws["B25"].alignment = left_wrap

    ws.merge_cells("B26:E27")
    ws["B26"] = (
        "【注意】小規模住宅に該当する場合、賃貸料相当額の100％以上を役員から徴収すれば"
        "給与課税は生じません。役員社宅で家賃の50％基準は使用人社宅の取扱いであり、役員には"
        "適用されない点に留意してください。固定資産税課税標準額は固定資産税納税通知書"
        "（賃貸人保有）または固定資産課税台帳の閲覧（区役所）で取得できます。"
    )
    ws["B26"].font = note_font
    ws["B26"].alignment = left_wrap

    # 保存
    path = os.path.join(DEST_DIR, f"{PREFIX}賃貸料相当額計算書.xlsx")
    wb.save(path)
    return path


# ---------------------------------------------------------------------------
# 5. 社宅導入オペレーション手順書
# ---------------------------------------------------------------------------
def make_tejunsyo() -> str:
    doc = Document()
    _set_default_font(doc)

    _add_title(doc, "社宅導入オペレーション手順書")
    _add_centered(doc, f"{HOJIN_NAME}", size=11)
    _add_centered(doc, f"作成日：{SAKUSEI_BI}", size=10)
    _add_para(doc)

    p = doc.add_paragraph()
    run = p.add_run("1.　目的")
    run.bold = True
    _add_para(
        doc,
        "　本手順書は、当法人が役員社宅制度を導入・運用するにあたり、"
        "物件契約から運用開始、税務調査対応までの実務手順を担当者向けに定めるものである。",
    )
    _add_para(doc)

    p = doc.add_paragraph()
    run = p.add_run("2.　導入時のステップ")
    run.bold = True

    steps = [
        (
            "STEP1　物件選定（小規模住宅要件チェック）",
            [
                "・床面積：耐用年数30年以下=132㎡以下／30年超=99㎡以下",
                "・本件物件：61.36㎡・RC造（耐用年数47年）→ 99㎡以下基準で該当",
                "・チェック書類：物件概要書、登記簿（建物表題部）",
            ],
        ),
        (
            "STEP2　社員総会・規程整備",
            [
                "・社宅管理規程の制定（社員総会決議）",
                "・社宅貸与の承認決議（利益相反取引に準じた承認）",
                "・必要書類：社宅管理規程、社員総会議事録",
            ],
        ),
        (
            "STEP3　賃貸借契約の締結（法人名義）",
            [
                "・契約名義は必ず法人とすること（個人名義は不可）",
                "・敷金・礼金・仲介手数料の領収書は法人名で受領",
                "・必要書類：賃貸借契約書、重要事項説明書、家賃保証契約書、火災保険証券（建物）",
            ],
        ),
        (
            "STEP4　社宅使用契約の締結（法人↔社員）",
            [
                "・本人覚書として社宅使用契約書を作成・記名押印",
                "・必要書類：社宅使用契約書",
            ],
        ),
        (
            "STEP5　仮社宅使用料の決定",
            [
                "・固定資産税課税標準額が未取得の段階では仮金額で運用開始",
                "・本件：月額家賃195,000円の15％＝29,250円を仮徴収",
                "・必要書類：社員総会議事録（仮金額の根拠記録）",
            ],
        ),
        (
            "STEP6　固定資産税課税標準額の取得",
            [
                "・賃貸人から固定資産税納税通知書の提供を依頼（建物・敷地の課税標準額）",
                "・取得困難な場合は、物件所在地の区役所で固定資産課税台帳を閲覧（借家人として可）",
                "・必要書類：固定資産税納税通知書（写）または固定資産評価証明書",
            ],
        ),
        (
            "STEP7　賃貸料相当額の正式算定",
            [
                "・賃貸料相当額計算書（Excel）に課税標準額を入力し算出",
                "・算定式（小規模住宅）：(1)建物課税標準額×0.2%＋(2)12円×(総床面積÷3.3)＋(3)敷地課税標準額×0.22%",
                "・必要書類：賃貸料相当額計算書（記入済）",
            ],
        ),
        (
            "STEP8　正式社宅使用料への改定",
            [
                "・賃貸料相当額の100％以上の金額を正式徴収額として決定",
                "・社員総会で改定決議（または法人と社員間の合意書）",
                "・必要書類：改定決議の社員総会議事録、社宅使用契約変更覚書",
            ],
        ),
        (
            "STEP9　毎月の運用",
            [
                "・社宅使用料を毎月の役員報酬から控除（給与計算ソフトで設定）",
                "・家賃は法人口座から賃貸人へ振込",
                "・freee：家賃は地代家賃／役員からの徴収は雑収入または地代家賃の貸方（経理処理は法人税基本通達9-2-9に基づき総額表示）",
            ],
        ),
    ]

    for title, lines in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        for line in lines:
            _add_para(doc, line)
        _add_para(doc)

    p = doc.add_paragraph()
    run = p.add_run("3.　税務調査対応のための保管書類リスト")
    run.bold = True
    _add_para(doc, "　以下の書類を一式まとめて、最低7年間保管する。")

    docs_to_keep = [
        "(1) 社宅管理規程（施行版）",
        "(2) 社員総会議事録（規程制定・社宅貸与承認）",
        "(3) 社宅使用契約書（法人↔社員）",
        "(4) 賃貸借契約書（法人↔賃貸人）／重要事項説明書／更新契約書",
        "(5) 賃貸料相当額計算書（記入済）",
        "(6) 固定資産税納税通知書または固定資産評価証明書（建物・敷地）",
        "(7) 月次の家賃支払証憑（振込明細）",
        "(8) 役員報酬控除明細（給与計算結果）",
        "(9) 火災保険証券（建物：法人負担／家財：個人負担）",
        "(10) 物件概要書、登記簿（小規模住宅要件確認用）",
    ]
    for line in docs_to_keep:
        _add_para(doc, line)
    _add_para(doc)

    p = doc.add_paragraph()
    run = p.add_run("4.　ありがちな指摘事項と対策")
    run.bold = True
    risks = [
        ("賃貸借契約が個人名義", "→ 必ず法人名義で締結。個人名義は社宅と認められず家賃全額が役員給与扱い。"),
        (
            "役員から徴収する賃料が賃貸料相当額未満",
            "→ 不足額が役員給与（定期同額給与の枠外）として源泉徴収・損金不算入リスク。"
            "課税標準額判明後すみやかに改定。",
        ),
        ("小規模住宅要件を超える物件", "→ 通常の役員社宅扱いとなり計算式が異なる（家賃の50％基準等）。"),
        ("社員総会の承認決議がない", "→ 利益相反取引の承認漏れ。後日の追認決議が必要。"),
        ("光熱水費等の個人負担分を法人が立替", "→ 役員給与認定リスク。費用負担区分を厳守。"),
    ]
    for title, body in risks:
        p = doc.add_paragraph()
        run = p.add_run(f"・{title}")
        run.bold = True
        _add_para(doc, f"　{body}")

    return _save(doc, f"{PREFIX}社宅導入オペレーション手順書.docx")


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------
def main() -> None:
    if not os.path.isdir(DEST_DIR):
        raise SystemExit(f"保存先フォルダが見つかりません: {DEST_DIR}")

    paths = [
        make_kitei(),
        make_giji(),
        make_keiyaku(),
        make_keisansho(),
        make_tejunsyo(),
    ]
    print("作成完了：")
    for p in paths:
        print(f"  - {p}")


if __name__ == "__main__":
    main()
