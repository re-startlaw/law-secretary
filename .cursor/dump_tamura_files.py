#!/usr/bin/env python3
"""聴取報告書.docx と 260520_上申書.docx の内容をダンプ"""
import sys
from pathlib import Path
from docx import Document

BASE = Path("/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/た_田村正宣/09_ワードファイル")

files = [
    BASE / "聴取報告書.docx",
    BASE / "260520_上申書.docx",
]

for f in files:
    print("=" * 80)
    print(f"FILE: {f.name}")
    print("=" * 80)
    if not f.exists():
        print("(not found)")
        continue
    doc = Document(str(f))
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "?"
        align = p.alignment
        text = p.text
        print(f"[{i:03d}] style={style} align={align}")
        print(f"      text={text!r}")
    print()
    print(f"-- TABLES: {len(doc.tables)} --")
    for ti, table in enumerate(doc.tables):
        print(f"  Table {ti}: rows={len(table.rows)} cols={len(table.columns)}")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                print(f"    ({ri},{ci}): {cell.text!r}")
