"""通常の連絡書面 雛形 docx を作成。

馬様事件 損害賠償請求書 v3 のレイアウト規律を踏襲し、依頼者固有情報を
プレースホルダー化して `02_ひな形/00_一般/連絡書面.docx` に保存する。

このひな型を `shutil.copy2` でコピーし、依頼者フォルダの Wordファイル
配下で本件用に書き換えるのが標準フロー。
"""

import os
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/02_ひな形/"
    "00_一般/連絡書面.docx"
)

CHAR_EMU = 133350  # ≒1文字


def add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  left_indent=None, first_line_indent=None,
                  bold=False, size_pt=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size_pt:
        run.font.size = Pt(size_pt)
    return p


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    # ---- ヘッダー ----
    add_paragraph(doc, "２０◯◯年◯月◯日", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    # 宛先
    add_paragraph(doc, "〒◯◯◯－◯◯◯◯")
    add_paragraph(doc, "◯◯都◯◯区◯◯町◯丁目◯番◯号")
    add_paragraph(doc, "◯◯◯◯法律事務所")
    add_paragraph(doc, "弁護士　◯◯　◯◯　先生")

    # 差出人（左寄せで右端寄り：first_line_indent でぐっと右に）
    sender_fi = Emu(3600450)
    add_paragraph(doc, "（差出人）", first_line_indent=sender_fi)
    add_paragraph(doc, "〒１７０-６０１２", first_line_indent=sender_fi)
    add_paragraph(doc, "東京都豊島区東池袋３丁目１－１",
                  first_line_indent=sender_fi)
    add_paragraph(doc, "サンシャイン６０　１２階", first_line_indent=sender_fi)
    add_paragraph(doc, "弁護士法人Ｒｅ－Ｓｔａｒｔ法律事務所",
                  first_line_indent=sender_fi)
    add_paragraph(doc, "弁護士　米谷尚起", first_line_indent=sender_fi)
    add_paragraph(doc, "TEL：０３－６８２０－３８１５",
                  first_line_indent=sender_fi)

    add_paragraph(doc, "")

    # タイトル（太字16pt・中央揃え）
    add_paragraph(doc, "◯◯◯◯について", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  bold=True, size_pt=16)
    add_paragraph(doc, "")

    # 本文骨子レイアウト：
    # - セクションタイトル（第○、（○）、【○】）: 左マージンから開始（インデント無し）
    # - 本文段落: left_indent=1字 + first_line_indent=1字
    #   → 1行目=2字下げ、2行目以降=1字下げ
    # - 表: tblInd=210twips (1字下げ)
    body_li = Emu(CHAR_EMU)
    body_fi = Emu(CHAR_EMU)

    # 前文（本文段落）
    add_paragraph(doc, "当職は、◯◯◯◯（以下「通知人」といいます。）の"
                       "代理人として、貴職に対し、以下のとおりご連絡申し"
                       "上げます。", left_indent=body_li, first_line_indent=body_fi)
    add_paragraph(doc, "")

    # 見出し例（インデント無し）
    add_paragraph(doc, "第１　◯◯◯◯")
    add_paragraph(doc, "（本文）", left_indent=body_li, first_line_indent=body_fi)
    add_paragraph(doc, "")

    add_paragraph(doc, "第２　◯◯◯◯")
    add_paragraph(doc, "（本文）", left_indent=body_li, first_line_indent=body_fi)
    add_paragraph(doc, "")

    # 振込先（実口座を雛形に固定）
    add_paragraph(doc, "【振込先】")
    add_paragraph(doc, "金融機関：ＧＭＯあおぞらネット銀行（１０２）"
                       "法人第二営業部支店",
                  left_indent=body_li, first_line_indent=body_fi)
    add_paragraph(doc, "口座種別・口座番号:普通　２４６９１３７",
                  left_indent=body_li, first_line_indent=body_fi)
    add_paragraph(doc, "口座名義:弁護士法人Ｒｅ－Ｓｔａｒｔ法律事務所　"
                       "預かり口",
                  left_indent=body_li, first_line_indent=body_fi)
    add_paragraph(doc, "")

    # 結語
    add_paragraph(doc, "本書面の内容をご確認いただき、ご不明な点がござ"
                       "いましたら、ご遠慮なくお尋ねくださいますようお願い"
                       "申し上げます。",
                  left_indent=body_li, first_line_indent=body_fi)
    add_paragraph(doc, "")

    # 末尾「以上」（RIGHT揃え、インデント無し）
    add_paragraph(doc, "以上", alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    main()
