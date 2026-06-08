#!/usr/bin/env python3
"""Extract text from 段衣信 通知書 and 鈴木七海 内容証明 (style template)."""
from docx import Document

paths = [
    "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/ダ_ダンイシン/260507_通知書（内容証明郵便案）.docx",
    "/Users/kometaninaoki/Library/CloudStorage/GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/す_鈴木七海/07_ワードファイル/260507電子内容証明-送付.docx",
]

for p in paths:
    print("=" * 80)
    print(p.split("/")[-1])
    print("=" * 80)
    doc = Document(p)
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else "Normal"
        text = para.text
        print(f"[{i:03d}|{style}] {text}")
    # also tables
    for ti, table in enumerate(doc.tables):
        print(f"--- Table {ti} ---")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                print(f"  [{ri},{ci}] {cell.text}")
    print()
