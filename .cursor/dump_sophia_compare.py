"""v3 と v7_2 の契約書本文・条文タイトルをダンプして差分を抽出。"""
from pathlib import Path
from docx import Document
import re
import sys

v3_path = Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v3.docx")
v7_2_path = Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v7_2.docx")

def load_doc(path):
    return Document(str(path))

def extract_article_titles(doc):
    """「第○条」で始まる段落を抽出。"""
    articles = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # 「第○条」で始まる行を抽出
        match = re.match(r'^第(\d+)条\s*(.*)$', text)
        if match:
            article_num = int(match.group(1))
            title = text
            articles[article_num] = {
                'idx': i,
                'title': title,
                'text': text
            }
    return articles

def dump_paragraphs(doc, label):
    """段落ごとにダンプ。"""
    output = []
    output.append("=" * 80)
    output.append(f"【{label} 本文（段落・スタイル・文字列）】")
    output.append("=" * 80)
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "（スタイルなし）"
        text = p.text
        output.append(f"[{i:03d}] style={style:30s} | {text!r}")
    return output

def extract_suspicious_expressions(doc):
    """契約条項として違和感のある表現を抽出。"""
    suspicious_patterns = [
        (r'（?～が望ましい）?', '推奨表現（望ましい）'),
        (r'～と考えられる', '意見・説明表現'),
        (r'～と思われる', '意見・説明表現'),
        (r'～となり得る', '可能性の説明（曖昧）'),
        (r'実質的に', 'editorial表現'),
        (r'いわゆる', 'editorial表現'),
        (r'結果として', '説明的フレーズ'),
        (r'当該|該当', 'potentially editorial'),
        (r'价値の最大化|紛争の抑制', '解説的フレーズ'),
        (r'塩漬けとなり得る', '解説的フレーズ'),
        (r'法令上、.*となる場合', '法令説明が条文に紛れ込み'),
        (r'ただし（ママ）', '単なる感想・推奨'),
    ]
    
    results = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        style = p.style.name if p.style else ""
        
        # 第○条の判定
        article_match = re.match(r'^第(\d+)条', text)
        article_num = None
        if article_match:
            article_num = int(article_match.group(1))
        
        for pattern, reason in suspicious_patterns:
            if re.search(pattern, text):
                results.append({
                    'para_idx': i,
                    'article': article_num,
                    'style': style,
                    'text': text,
                    'pattern': pattern,
                    'reason': reason
                })
    return results

# ====== main ======
print("=" * 80)
print("【v3 と v7_2 の契約書構成比較】")
print("=" * 80)

doc_v3 = load_doc(v3_path)
doc_v7_2 = load_doc(v7_2_path)

articles_v3 = extract_article_titles(doc_v3)
articles_v7_2 = extract_article_titles(doc_v7_2)

print(f"\nv3: {len(doc_v3.paragraphs)} 段落、{len(articles_v3)} 条文")
print(f"v7_2: {len(doc_v7_2.paragraphs)} 段落、{len(articles_v7_2)} 条文")

print("\n--- v3 条文タイトル一覧 ---")
for num in sorted(articles_v3.keys()):
    print(f"  {articles_v3[num]['title']}")

print("\n--- v7_2 条文タイトル一覧 ---")
for num in sorted(articles_v7_2.keys()):
    print(f"  {articles_v7_2[num]['title']}")

print("\n" + "=" * 80)
print("【v7_2 全段落ダンプ】")
print("=" * 80)
for line in dump_paragraphs(doc_v7_2, "v7_2"):
    print(line)

print("\n" + "=" * 80)
print("【v7_2 で抽出した違和感のある表現】")
print("=" * 80)
suspicious = extract_suspicious_expressions(doc_v7_2)
if suspicious:
    for item in suspicious:
        print(f"\n[段落{item['para_idx']:03d}] article={item.get('article')} | reason={item['reason']}")
        print(f"  text: {item['text']!r}")
else:
    print("(該当なし)")

