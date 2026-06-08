"""馬さん案件・受任通知_4 を作成。

仕様:
- 最終形は _3 と同じ（指示6 不適用、(5)削除に伴う番号繰上げあり、新規(7)(8)(9)追加）
- 違いは「Word 変更履歴付き」であること（元 .docx と _3 の差分が w:ins/w:del で見える）
- ベース: 元 260512電子内容証明案.docx を _4 にコピー済み
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
NSMAP = {"w": W_NS}

AUTHOR = "米谷尚起"
DATE = "2026-05-12T00:00:00Z"

INPUT = Path(
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/"
    "共有用/01_事件記録/ま_馬さん/Wordファイル/260512電子内容証明案_4.docx"
)


class IdGen:
    def __init__(self) -> None:
        self.n = 2000

    def next(self) -> str:
        self.n += 1
        return str(self.n)


IDS = IdGen()


def make_run(text: str, base_rpr: etree._Element | None = None) -> etree._Element:
    r = etree.Element(f"{W}r")
    if base_rpr is not None:
        r.append(copy.deepcopy(base_rpr))
    t = etree.SubElement(r, f"{W}t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def make_del_run(text: str, base_rpr: etree._Element | None = None) -> etree._Element:
    r = etree.Element(f"{W}r")
    if base_rpr is not None:
        r.append(copy.deepcopy(base_rpr))
    t = etree.SubElement(r, f"{W}delText")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def wrap_ins(runs: list[etree._Element]) -> etree._Element:
    ins = etree.Element(f"{W}ins")
    ins.set(f"{W}id", IDS.next())
    ins.set(f"{W}author", AUTHOR)
    ins.set(f"{W}date", DATE)
    for r in runs:
        ins.append(r)
    return ins


def wrap_del(runs: list[etree._Element]) -> etree._Element:
    d = etree.Element(f"{W}del")
    d.set(f"{W}id", IDS.next())
    d.set(f"{W}author", AUTHOR)
    d.set(f"{W}date", DATE)
    for r in runs:
        d.append(r)
    return d


def extract_first_rpr(p_elem: etree._Element) -> etree._Element | None:
    return p_elem.find(f".//{W}r/{W}rPr", namespaces=NSMAP)


def replace_paragraph_content(p_elem: etree._Element, new_text: str) -> None:
    base_rpr = extract_first_rpr(p_elem)
    old_text_parts: list[str] = []
    children_to_remove: list[etree._Element] = []
    for child in list(p_elem):
        if child.tag == f"{W}pPr":
            continue
        if child.tag == f"{W}r":
            for t in child.findall(f"{W}t", namespaces=NSMAP):
                old_text_parts.append(t.text or "")
        elif child.tag == f"{W}ins":
            for r in child.findall(f"{W}r", namespaces=NSMAP):
                for t in r.findall(f"{W}t", namespaces=NSMAP):
                    old_text_parts.append(t.text or "")
        children_to_remove.append(child)
    for c in children_to_remove:
        p_elem.remove(c)
    old_text = "".join(old_text_parts)
    if old_text:
        p_elem.append(wrap_del([make_del_run(old_text, base_rpr)]))
    if new_text:
        p_elem.append(wrap_ins([make_run(new_text, base_rpr)]))


def delete_paragraph_as_tracked(p_elem: etree._Element) -> None:
    base_rpr = extract_first_rpr(p_elem)
    old_text_parts: list[str] = []
    children_to_remove: list[etree._Element] = []
    pPr = p_elem.find(f"{W}pPr", namespaces=NSMAP)
    for child in list(p_elem):
        if child.tag == f"{W}pPr":
            continue
        if child.tag == f"{W}r":
            for t in child.findall(f"{W}t", namespaces=NSMAP):
                old_text_parts.append(t.text or "")
        children_to_remove.append(child)
    for c in children_to_remove:
        p_elem.remove(c)
    old_text = "".join(old_text_parts)
    if old_text:
        p_elem.append(wrap_del([make_del_run(old_text, base_rpr)]))
    if pPr is None:
        pPr = etree.SubElement(p_elem, f"{W}pPr")
        p_elem.insert(0, pPr)
    rPr = pPr.find(f"{W}rPr", namespaces=NSMAP)
    if rPr is None:
        rPr = etree.SubElement(pPr, f"{W}rPr")
    del_marker = etree.Element(f"{W}del")
    del_marker.set(f"{W}id", IDS.next())
    del_marker.set(f"{W}author", AUTHOR)
    del_marker.set(f"{W}date", DATE)
    rPr.insert(0, del_marker)


def insert_paragraph_after_as_tracked(
    after_p_elem: etree._Element, text: str
) -> etree._Element:
    base_pPr = after_p_elem.find(f"{W}pPr", namespaces=NSMAP)
    base_rpr = extract_first_rpr(after_p_elem)
    new_p = etree.Element(f"{W}p")
    if base_pPr is not None:
        new_pPr = copy.deepcopy(base_pPr)
        new_p.append(new_pPr)
        rPr = new_pPr.find(f"{W}rPr", namespaces=NSMAP)
        if rPr is None:
            rPr = etree.SubElement(new_pPr, f"{W}rPr")
        ins_marker = etree.Element(f"{W}ins")
        ins_marker.set(f"{W}id", IDS.next())
        ins_marker.set(f"{W}author", AUTHOR)
        ins_marker.set(f"{W}date", DATE)
        rPr.insert(0, ins_marker)
    if text:
        new_p.append(wrap_ins([make_run(text, base_rpr)]))
    after_p_elem.addnext(new_p)
    return new_p


def renumber_paragraph(p_elem: etree._Element, old_num: str, new_num: str) -> None:
    """段落本文の先頭にある番号 (例: 「（６）」) のみを del/ins で書換。
    本文はそのまま温存して、変更箇所が明確になるようにする。"""
    for child in list(p_elem):
        if child.tag != f"{W}r":
            continue
        for t in child.findall(f"{W}t", namespaces=NSMAP):
            text = t.text or ""
            if old_num in text:
                idx = text.index(old_num)
                prefix = text[:idx]
                suffix = text[idx + len(old_num):]
                # 現 run のテキストを prefix に縮める
                t.text = prefix
                t.set(
                    "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                )
                # 元 run の rPr を取得
                run_rpr = child.find(f"{W}rPr", namespaces=NSMAP)
                # del/ins/suffix run を作って child の後ろに順に挿入
                del_elem = wrap_del([make_del_run(old_num, run_rpr)])
                ins_elem = wrap_ins([make_run(new_num, run_rpr)])
                # suffix を持つ新 run
                suffix_r = etree.Element(f"{W}r")
                if run_rpr is not None:
                    suffix_r.append(copy.deepcopy(run_rpr))
                st = etree.SubElement(suffix_r, f"{W}t")
                st.text = suffix
                st.set(
                    "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                )
                child.addnext(suffix_r)
                child.addnext(ins_elem)
                child.addnext(del_elem)
                # prefix が空ならオリジナル run を削除（番号が run 先頭の場合）
                if prefix == "":
                    p_elem.remove(child)
                return
    raise ValueError(f"renumber: '{old_num}' not found in paragraph")


def main() -> None:
    doc = Document(str(INPUT))
    paragraphs = doc.paragraphs

    # ---- 修正1+2: Para 24 見出し ----
    p24 = paragraphs[24]
    assert p24.text.strip() == "（１）　小牧氏がAngelina氏の襟元を掴んで引き移動させた行為"
    replace_paragraph_content(
        p24._element,
        "（１）　小牧氏がAngelina氏の襟元付近に身体的接触を加えて移動させ、トラック脇で停止させた行為",
    )

    # ---- 修正1+2: Para 25 本文 ----
    p25 = paragraphs[25]
    assert p25.text.startswith("　令和８年４月１３日午前９時１８分頃")
    replace_paragraph_content(
        p25._element,
        (
            "　令和８年４月１３日午前９時２０分頃から同９時３０分頃までの間、貴校体育館において、"
            "小牧氏は、Angelina氏に対し、後方からその襟元付近に身体的接触を加えてこれを掴み、"
            "Angelina氏の身体を伴って移動させた上、トラック脇においてこれを停止させるとともに、"
            "大声で叱責しました。当該事実は、Angelina氏の同日付書面陳述、当方が保有する事案"
            "当日の映像等から、優に認められるものです。"
        ),
    )

    # ---- 修正3: Para 34 数学授業 ----
    p34 = paragraphs[34]
    assert p34.text.startswith("　Angelina氏は、令和８年４月１３日の１時間目の算数")
    replace_paragraph_content(
        p34._element,
        (
            "　Angelina氏は、令和８年４月１３日の１時間目の数学の時間（午前８時４０分から"
            "同９時３５分まで）中、約４分間体育館に立ち寄ったことを認めています。これは、"
            "Angelina氏が自身の課題を比較的早く終えた後、長男であるGeorge氏の約４分間の"
            "クロスカントリー競技を観覧するため、体育館に向かったものです。当該時間は、"
            "前半が教員による講義の時間、後半が各生徒が課題を行う自習の時間として運用"
            "されていたものであり、また、当該自習時間中に生徒が一時的に教室を離れること"
            "についても、現場の教員において従前から黙認されてきたものです。"
        ),
    )

    # ---- 修正4: Para 40 (2) 精神状態悪化 ----
    p40 = paragraphs[40]
    assert p40.text.startswith("（２）　貴校から不適切な対応を受けたこと")
    replace_paragraph_content(
        p40._element,
        (
            "（２）　貴校から不適切な対応を受けたこと、登校できないこと、及び事実に反して"
            "Angelina氏が虚偽の主張をしている旨を周囲に言いふらされたことによりAngelina氏"
            "の精神状態が悪化し、投薬調整及び継続的な心理カウンセリングを要する状態となった"
            "ことに伴う精神科通院費用及び心理カウンセリング費用"
        ),
    )

    # ---- 修正5: Para 43 (5) 妻関連損害 → 段落削除 ----
    p43 = paragraphs[43]
    assert p43.text.startswith("（５）　George氏の看護者の校内立入りが何らの根拠なく拒絶")
    delete_paragraph_as_tracked(p43._element)

    # ---- 番号繰上げ: 旧(6)→新(5), 旧(7)→新(6) ----
    p_kippu = next(
        p
        for p in doc.paragraphs
        if (p.text or "").startswith("（６）　既に購入済みのAngelina氏の通学定期券")
    )
    renumber_paragraph(p_kippu._element, "（６）", "（５）")

    p_fee = next(
        p
        for p in doc.paragraphs
        if (p.text or "").startswith("（７）　弁護士費用その他本件への法的対応")
    )
    renumber_paragraph(p_fee._element, "（７）", "（６）")

    # ---- 修正9: 弁護士費用の直後に新規 (7)(8)(9) を挿入 ----
    anchor = p_fee._element
    new_items = [
        "（７）　Angelina氏の長期にわたる学籍中断により、Angelina氏の学業継続性が毀損されていることに関する損害",
        "（８）　Angelina氏の試験成績、GPA及びTranscriptへの不利益な影響に関する損害",
        "（９）　Angelina氏の大学進学経路への現実的影響並びに教育機会の喪失及び将来的進学リスクに関する損害",
    ]
    for txt in new_items:
        anchor = insert_paragraph_after_as_tracked(anchor, txt)

    # ---- 修正8: 第５(2) 直ちに復学 → 公平・安全な教育環境 ----
    p_fukugaku = next(
        p
        for p in doc.paragraphs
        if (p.text or "").strip() == "（２）　Angelina氏を直ちに復学させていただくこと"
    )
    replace_paragraph_content(
        p_fukugaku._element,
        (
            "（２）　Angelina氏に対する公平かつ安全な教育環境を確保した上で、"
            "Angelina氏の復学を実現していただくこと（なお、当方として復学権自体を放棄"
            "するものではありません。）"
        ),
    )

    # ---- 修正7: 第６(3) 学費相当額 → 1,388,000円 ----
    p_gakuhi = next(
        p for p in doc.paragraphs if "学費相当額（金○○万円）" in (p.text or "")
    )
    replace_paragraph_content(
        p_gakuhi._element,
        (
            "（３）　もっとも、Angelina氏及び通知人は、貴校との在学契約の継続を強く"
            "希望しております。当方は、学費相当額（金１,３８８,０００円。Angelina氏の"
            "Ｇ１０学年第一期学費に相当します。）を当職の弁護士預り金口座において保全し、"
            "貴校による教育役務の提供再開と引換えに、直ちに同額をお支払いする用意がございます。"
        ),
    )

    # ---- 指示6 は適用しない ----

    doc.save(str(INPUT))
    print(f"saved: {INPUT}")


if __name__ == "__main__":
    main()
