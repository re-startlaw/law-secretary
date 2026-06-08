"""Create 受領証 for 田村正宣 case with filled-in items."""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

OUTPUT_PATH = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/"
    "01_事件記録/た_田村正宣/09_ワードファイル/260427_受領証.docx"
)

ITEMS = [
    ("免許証", "１枚"),
    ("マイナンバーカード", "１枚"),
    ("保険証", "１枚"),
    ("楽天クレジットカード", "１枚"),
    ("ポイントカード", "５枚"),
    ("Suica", "１枚"),
    ("PASMO", "１枚"),
    ("ゆうちょキャッシュカード", "１枚"),
    ("ゆうちょ通帳", "２冊"),
    ("SMBCキャッシュカード", "１枚"),
    ("クオカード", "１枚"),
    ("印鑑登録証", "１通"),
    ("巣鴨キャッシュカード", "１枚"),
    ("診察券", "８枚"),
    ("警察の名刺", "３枚"),
    ("ネックレス", "２点"),
    ("ブレスレット", "１点"),
    ("腕時計", "１点"),
]

JP_FONT = "ＭＳ 明朝"
ABSTRACT_NUM_ID = "99"
NUM_ID = "99"


def set_jp_font(run, size_pt=11, bold=False):
    run.font.name = JP_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), JP_FONT)
    rFonts.set(qn("w:ascii"), JP_FONT)
    rFonts.set(qn("w:hAnsi"), JP_FONT)


def add_full_width_numbering(doc):
    numbering = doc.part.numbering_part.element

    abstract_num_xml = (
        '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:abstractNumId="{ABSTRACT_NUM_ID}">'
        '<w:multiLevelType w:val="singleLevel"/>'
        '<w:lvl w:ilvl="0">'
        '<w:start w:val="1"/>'
        '<w:numFmt w:val="decimalFullWidth"/>'
        '<w:suff w:val="space"/>'
        '<w:lvlText w:val="%1"/>'
        '<w:lvlJc w:val="left"/>'
        '<w:pPr>'
        '<w:ind w:left="0" w:hanging="0"/>'
        '</w:pPr>'
        '</w:lvl>'
        '</w:abstractNum>'
    )
    abstract_num_elem = parse_xml(abstract_num_xml)
    nums = numbering.findall(qn("w:num"))
    if nums:
        numbering.insert(list(numbering).index(nums[0]), abstract_num_elem)
    else:
        numbering.append(abstract_num_elem)

    num_xml = (
        '<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:numId="{NUM_ID}">'
        f'<w:abstractNumId w:val="{ABSTRACT_NUM_ID}"/>'
        '</w:num>'
    )
    num_elem = parse_xml(num_xml)
    numbering.append(num_elem)


def _add_dotted_bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for tag in ("bottom", "between"):
        b = OxmlElement(f"w:{tag}")
        b.set(qn("w:val"), "dotted")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "auto")
        pBdr.append(b)
    pPr.append(pBdr)


def add_list_item(doc, name_text, count_text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), NUM_ID)
    numPr.append(ilvl)
    numPr.append(numId_el)
    pPr.append(numPr)

    p.paragraph_format.tab_stops.add_tab_stop(Mm(160), WD_TAB_ALIGNMENT.RIGHT)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:hanging"), "0")
    pPr.append(ind)

    _add_dotted_bottom_border(p)
    p.paragraph_format.space_after = Pt(6)

    set_jp_font(p.add_run(name_text + "\t" + count_text))
    return p


def add_para(doc, text, *, align=None, size=11, bold=False, left_indent_mm=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if left_indent_mm is not None:
        p.paragraph_format.left_indent = Mm(left_indent_mm)
    if text:
        set_jp_font(p.add_run(text), size_pt=size, bold=bold)
    return p


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)

    normal = doc.styles["Normal"]
    normal.font.name = JP_FONT
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), JP_FONT)
    rFonts.set(qn("w:ascii"), JP_FONT)
    rFonts.set(qn("w:hAnsi"), JP_FONT)

    add_full_width_numbering(doc)

    add_para(doc, "№　　　　　")
    add_para(doc, "受　領　証", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True)
    add_para(doc, "")
    add_para(doc, "令和　　年　　月　　日", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, "")
    add_para(doc, "弁護士　米谷尚起　　殿")
    add_para(doc, "")
    add_para(doc, "住　所　　　　　　　　　　　　　　　　", left_indent_mm=70)
    add_para(doc, "氏　名　　　　　　　　　　　　　　　印", left_indent_mm=70)
    add_para(doc, "")
    add_para(doc, "件　名　　　　　　　　　　　　　　　　　　　　　　　　　　事件")
    add_para(doc, "")
    add_para(doc, "標記事件について、次の書類の返還を受け、正に受領致しました。")
    add_para(doc, "")

    for name, count in ITEMS:
        add_list_item(doc, name, count)

    add_para(doc, "")
    add_para(doc, "以　上", align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUTPUT_PATH)
    print("saved:", OUTPUT_PATH)


if __name__ == "__main__":
    main()
