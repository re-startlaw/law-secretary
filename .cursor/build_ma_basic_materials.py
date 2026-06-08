"""馬強事件_基礎資料整理.docx を生成する。

続編内容証明作成の前提として、依頼者・学校・寺井弁護士とのメール内容、
案件フォルダ内の資料一覧、馬さんの要求項目・損害項目、当方が前回内容証明
で挙げた項目、5/15付寺井弁護士書面の内容などを網羅的に整理する。
"""

import json
import os
import re
from datetime import datetime, timezone, timedelta

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

JST = timezone(timedelta(hours=9))

OUT = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260522_馬強事件_基礎資料整理.docx"
)

PREV_NAIYO_DOCX = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260515電子内容証明案_提出版_名前訂正.docx"
)

JIANGAI_DOCX = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/事案の概要 (1).docx"
)


def fmt_jst(iso_utc):
    try:
        dt = datetime.fromisoformat(iso_utc.replace("Z", "+00:00"))
        return dt.astimezone(JST).strftime("%Y-%m-%d %H:%M")
    except Exception:
        return iso_utc


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt({1: 14, 2: 12, 3: 11}.get(level, 11))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, indent=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if italic:
        run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_quote(doc, text):
    """元メール本文を引用ブロックとして追加。"""
    if not text.strip():
        return
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run("│ " + line if line.strip() else "│")
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def dump_docx_paragraphs(path):
    try:
        d = Document(path)
        return [p.text for p in d.paragraphs if p.text.strip()]
    except Exception as e:
        return [f"[読込失敗: {e}]"]


def main():
    doc = Document()

    # ---- タイトル ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("馬強事件　基礎資料整理")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub.add_run("作成日：2026年5月22日").font.size = Pt(10)

    add_para(
        doc,
        "本書は、馬強氏（依頼者）からの相談・委任事項、KIST及び寺井勇人弁護士"
        "とのやり取り、案件記録フォルダ内の資料を整理し、次回内容証明（追加請求"
        "書面）作成の基礎資料とすることを目的に作成する。",
    )

    # ---- 第1部　事案の概要 ----
    add_heading(doc, "第１　事案の概要", 1)
    add_para(doc, "（出典：事案の概要 (1).docx／案件記録フォルダ）")
    for line in dump_docx_paragraphs(JIANGAI_DOCX):
        add_para(doc, line, indent=0.5)

    # ---- 第2部　関係者一覧 ----
    add_heading(doc, "第２　関係者一覧", 1)
    persons = [
        ("通知人（依頼者）", "馬強　Qiang Ma（martin.ma@letour.co.jp）"),
        ("通知人の配偶者", "Ni Liu"),
        ("次女", "Angelina Feili Ma（KIST G9A 在籍）"),
        ("長男", "George Man Tzar（KIST K2A 在籍）"),
        ("学校", "K. International School Tokyo（KIST）"),
        ("理事会議長", "Takako Komaki（小牧孝子）"),
        ("副理事長", "Kei Komaki（小牧氏ご子息）"),
        ("総校長", "Kevin Yoshihara（吉原）"),
        ("中学部校長", "Mark Cowe"),
        ("スチューデントケアコーディネーター", "Karen Donald Godfrey"),
        ("学校代理人・理事兼顧問弁護士", "寺井勇人弁護士（dzr03674@nifty.com）"),
        ("通知人代理人", "弁護士法人Re-Start法律事務所 弁護士 米谷尚起"),
    ]
    for role, name in persons:
        add_para(doc, f"・{role}：{name}", indent=0.5)

    # ---- 第3部　経緯（時系列） ----
    add_heading(doc, "第３　経緯（時系列）", 1)
    timeline = [
        ("2026-04-13 09:20頃-09:30頃",
         "KIST体育館にて、小牧氏がAngelinaの襟元に身体的接触を加えてトラック脇まで移動させ、停止のうえ叱責。"
         "Angelinaは事案当日に書面陳述を作成・提出。Mark Cowe校長より、協議完了までAngelina登校を認めない旨のメール送信。"),
        ("2026-04-14",
         "保護者がKISTを訪問し学校と面談（録音あり）。学校側は理事長の有形力行使を否定し、登校拒否の意思を継続。"),
        ("2026-04-14〜15",
         "学校とのメールやり取り継続。Angelinaの精神状態が悪化、精神科受診。"),
        ("2026-04-16",
         "同日体育館で撮影をしていた他の保護者から動画（約4分11秒・約4秒の2本）を受領。"),
        ("2026-04-17",
         "学校から、George氏看護者の校内立入りを4月20日以降禁止する旨通知。"),
        ("2026-04-19",
         "保護者が動画の存在を学校へ開示の上、学校見解の根拠を再度求めるも回答変わらず。"
         "対応窓口が学校法務顧問・寺井勇人弁護士へ切替え。"),
        ("2026-04-20",
         "寺井弁護士よりMa夫妻宛「Official Communication from KIST Board of Directors and Legal Counsel」メール（事案を刑事的構成・登校復帰拒否を明示）。"),
        ("2026-04-21",
         "馬さんから学校・寺井弁護士宛「今後は当方代理人を通じて連絡する」旨メール送付。"),
        ("2026-05-07",
         "馬さんが弁護士法人Re-Start法律事務所に来所、米谷弁護士と相談。委任契約書草案を発送。"),
        ("2026-05-08〜09",
         "成功報酬構成・経済的損害項目について馬さんと協議。修正合意。"),
        ("2026-05-10",
         "委任契約書 Adobe Sign 署名完了。"),
        ("2026-05-11",
         "馬さんから着手金33万円の入金確認。同日付で馬さんより「G10学費支払いを保留し、退学扱いを阻止する内容の通知書」作成要請。"),
        ("2026-05-12",
         "内容証明（受任通知及び御連絡）原案作成、馬さん修正コメント9項目を反映、馬さんから学費預り金1,388,000円の入金確認、当日内容証明郵便を速達発送。"),
        ("2026-05-14",
         "馬さんよりKIST Accounts宛の事前メール案（5/15期限前連絡）について確認依頼。"),
        ("2026-05-15",
         "内容証明、寺井弁護士事務所到着確認（高輪→赤坂転送）。"
         "寺井弁護士の指摘を受け、宛名表記訂正版PDFをメール送付。"
         "馬さんよりKIST Accounts宛メール送付。"
         "同日、寺井弁護士よりMa夫妻宛「Regarding Tuition Payment for G9A」メール送信"
         "（学費期限延長を承認しつつ、4月21日以降4週間以上代理人連絡がない点を指摘）。"),
        ("2026-05-18",
         "馬さんが寺井弁護士に対し、Reply All形式で事実関係を整理する英文返信送付（米谷弁護士BCC共有）。"),
        ("2026-05-19",
         "馬さんから追加の請求書面（金額請求一覧・損害項目）の準備要請。直近数週間の医療費を含める旨依頼。"),
        ("2026-05-22",
         "本書（基礎資料整理）作成。続編内容証明案の作成準備。"),
    ]
    for when, what in timeline:
        add_para(doc, f"■ {when}", indent=0.3)
        add_para(doc, what, indent=0.8)

    # ---- 第4部　馬さんの要求項目・成功報酬構成 ----
    add_heading(doc, "第４　馬さんの要求項目・成功報酬構成", 1)
    add_para(doc, "（出典：2026-05-08付martin.maメール、2026-05-09付当方回答メール）")
    add_heading(doc, "（１）馬さんが実現を希望している事項", 2)
    for item in [
        "本件における学校側の対応および処理過程に不適切な点があったことの確認",
        "Angelinaの名誉および人格的評価の回復",
        "Georgeの看護者による正常な校内立入り権限の回復",
        "本件によって生じた実際の経済的損害および精神的損害に対する適切な補償",
        "復学問題に関する検討（公平かつ安全な学習環境の確保が前提）",
    ]:
        add_para(doc, f"・{item}", indent=0.5)

    add_heading(doc, "（２）成功報酬構成（委任契約反映済み）", 2)
    for item in [
        "学校側が不適切性を認め、Angelinaに書面謝罪・名誉回復措置を行った場合 → 15万円",
        "Georgeの看護者による正常な校内立入り権限が回復された場合 → 15万円",
        "学校側が合理的かつ公平な復学方案を提示し、双方合意のうえ実際に正常な就学が再開された場合 → 15万円",
        "経済的損害について → 経済的利益の15％",
    ]:
        add_para(doc, f"・{item}", indent=0.5)

    # ---- 第5部　損害項目・想定金額 ----
    add_heading(doc, "第５　損害項目・想定金額（依頼者主張）", 1)
    add_para(doc, "（出典：2026-05-08付martin.maメール、5-12付martin.maメール、5-19付martin.maメール）")
    add_heading(doc, "（１）実費損害", 2)
    items_actual = [
        ("CGA（Crimson Global Academy）G9 Term 2学費（代替教育費）", "1,260,000円"),
        ("CGA G10学年オンラインスクール預託金（10％ Deposit）", "244,800円"),
        ("既に支払済みのKIST G9学費のうち、登校拒否期間相当額", "G9年額2,897,000円を按分（2026/4/14〜正常復学日まで）"),
        ("購入済み通学定期券損失", "2026/1/17購入・6か月35,980円のうち2026/4/14以降未利用相当分"),
        ("精神科診療費（2026/4/16）", "21,850円"),
        ("精神科診療費（2026/4/30）", "57,050円"),
        ("心理カウンセリング費（2026/5/1）", "17,000円"),
        ("心理カウンセリング費（2026/5/8）", "17,000円"),
        ("精神科往復交通費", "1回840円（2〜3週間に1回程度）"),
        ("心理カウンセリング往復交通費", "1回840円（週1回程度）"),
        ("2026-05-19以降に発生した医療費（領収書受領待ち）", "追って加算"),
        ("弁護士費用", "本件法的対応費用一式"),
    ]
    for name, amount in items_actual:
        add_para(doc, f"・{name}：{amount}", indent=0.5)

    add_heading(doc, "（２）精神的損害", 2)
    add_para(doc, "依頼者は暫定請求額として約7,000,000円を想定（2026-05-08付martin.maメール）。", indent=0.5)
    add_para(doc, "本件発生後の精神的損害の具体的内容：", indent=0.5)
    for item in [
        "本件および学校側のその後の対応によりAngelina本人が受けた精神的損害",
        "長期間正常に登校できず、約1か月にわたり授業を受けられなかったことによる不安・心理的負担",
        "学業中断・授業継続性の破壊による試験成績、GPA、Transcript完整性、将来の大学進学経路への現実的影響",
        "Angelina本人の名誉権・人格権・校内評価に対する侵害",
        "学校側の対応によって生じた教育機会損失および将来の進学上の不利益リスク",
    ]:
        add_para(doc, f"  ・{item}", indent=1.0)

    add_heading(doc, "（３）参考：4月30日打合せメモ上の想定", 2)
    add_para(doc, "事案の概要 (1).docx（4月30日時点の方針整理）では、慰謝料額は50万円程度を妥当と想定。"
                 "ただし、依頼者の意向を踏まえ最終調整するとされている。", indent=0.5)

    # ---- 第6部　当方が前回内容証明で主張した項目（5/12発送） ----
    add_heading(doc, "第６　前回内容証明（2026-05-12発送・5/15訂正版）の構成・主張項目", 1)
    add_para(doc, "（出典：260515電子内容証明案_提出版_名前訂正.docx）")
    for line in dump_docx_paragraphs(PREV_NAIYO_DOCX):
        add_para(doc, line, indent=0.5)

    # ---- 第7部　5/15付寺井弁護士書面の要旨と論点 ----
    add_heading(doc, "第７　2026-05-15付寺井勇人弁護士書面（KIST Accounts宛馬様メール返信）", 1)
    add_para(doc,
             "寺井弁護士からMa夫妻宛（学校管理層CC）に送信されたメールの要旨は以下のとおり。"
             "学校設立・KIST関係者（小牧理事長、Naito、Yoshihara、Komaki Kei、school info、Accounts）"
             "宛にもCCされており、校内関係者間で「家長側代理人から長期間連絡が行われていない」"
             "との印象を作出する内容となっている点に、依頼者は懸念を示している。")
    add_heading(doc, "（１）寺井弁護士書面の要点", 2)
    for item in [
        "2026〜2027年度学費納付期限の延長を承認する。",
        "Ma夫妻が4月21日に「代理人を通じて連絡する」と明言したにもかかわらず、その後4週間以上Ma側代理人から一切連絡がない点を強調。",
        "Angelinaに関する一切の照会・依頼・連絡は、今後すべて寺井弁護士宛に行うことを要求。",
        "Angelinaが書面陳述で主張する「dragged by the collar」は刑事事件相当の重大行為であり、学校としては事実調査義務を負う。",
        "学校・家族側で事案の事実について完全合意し、合意した措置を講ずるまでAngelinaの登校は認められない。",
        "Ma側の説明は「ぶれ」が大きく、調査を長期化させているのはMa側であると指摘。",
        "学校側に対し、当該動画データを直ちに学校へ提出すること、入手経路の明確化を要求。"
        "Ma家族の所持機器または公開情報以外からの入手は、4月以降の情報拡散禁止依頼に反する旨指摘。",
        "学校所属アドバイザーである元警視庁刑事 Akinari Hirano による現場検証の実施を要求し、Angelina及び保護者の来校日程提示を要求。",
    ]:
        add_para(doc, f"・{item}", indent=0.5)

    add_heading(doc, "（２）当方として整理すべき論点", 2)
    for item in [
        "動画は2026-04-19の段階で学校側へ存在を開示済み。提出義務の存否（民事上の任意提出か、刑事事件・民事訴訟手続による提出か）。",
        "動画の入手元は他の保護者であり、4月以降の情報拡散禁止依頼に反するものではない（家族の動画取得経緯はWeChat記録で確認可能）。",
        "現場検証を要求している Akinari Hirano 氏は学校アドバイザーであり、中立性（独立性）に重大な疑義あり。"
        "中立的検証であれば、双方合意の第三者または捜査機関を介すべき。",
        "学校による「完全合意までは登校不可」との取扱いは、在学契約上の教育役務提供義務との関係で法的根拠を欠く。",
        "馬さん主張のように4月21日以降の期間は証拠整理・GW期間を含むものであり、「4週間以上連絡なし」との校内CC送信は、"
        "対外的印象操作の一面があり、当方代理人を通じて事実関係の訂正を求めるべき。",
    ]:
        add_para(doc, f"・{item}", indent=0.5)

    # ---- 第8部　メール時系列（martin.ma & 寺井弁護士 thread） ----
    add_heading(doc, "第８　メール時系列（原文）", 1)
    add_para(doc, "依頼者martin.ma@letour.co.jpおよび学校側寺井勇人弁護士dzr03674@nifty.comとの送受信を時系列で示す。")

    data = json.load(open(".cursor/ma_thread_data.json", encoding="utf-8"))
    # Combine and sort
    all_msgs = []
    for m in data["thread_main"]:
        m["thread"] = "馬様(契約・受任通知・追加方針) thread"
        all_msgs.append(m)
    for m in data["thread_kist_forward"]:
        m["thread"] = "KIST学費・寺井弁護士 forward thread"
        all_msgs.append(m)
    all_msgs.sort(key=lambda x: x["date"] or "")

    for m in all_msgs:
        head = f"〔{fmt_jst(m['date'])}〕 {m['sender']} → {', '.join(m.get('to') or [])}"
        if m.get("cc"):
            head += f"（CC: {', '.join(m['cc'])}）"
        add_para(doc, head, indent=0.0)
        if m.get("subject"):
            add_para(doc, f"Subject: {m['subject']}", indent=0.3)
        add_para(doc, f"［{m['thread']}］", indent=0.3, italic=True)
        add_quote(doc, m.get("body") or "")
        doc.add_paragraph()  # spacer

    # ---- 第9部　案件記録フォルダ内資料一覧 ----
    add_heading(doc, "第９　案件記録フォルダ内資料一覧", 1)
    add_para(doc, "（パス：01_事件記録/ま_馬強/）")

    folder_struct = {
        "ルート": [
            "260507_ネット塾費用は5月1日に支払い済みです.jpg",
            "260511_案件整理票（馬様）.gsheet",
            "260512_web21_60 - 2026-05-12T180043.860.pdf（内容証明配達証拠）",
            "事案の概要 (1).docx",
            "整理表（馬強）.gsheet",
        ],
        "03連絡文書": [
            "260514_馬強事件_Angelina登校問題に関する通知書面.pdf",
        ],
        "06資料/260507受領": [
            "Angelina書面陳述（4月13日付）.pdf",
            "CGA入学プロセス資料.pdf",
            "WeChat記録（証拠動画取得経緯）.pdf",
            "事案整理書面（依頼者作成）.pdf",
            "事案発生現場の位置関係図.pdf",
            "医療費領収書（AmericanClinicTokyo_4月16日精神科）.pdf",
            "学校メール_KarenDonaldGodfreyスチューデントケアコーディネーター.pdf",
            "学校メール_KarenDonaldGodfreyスチューデントケアコーディネーター_2.pdf",
            "学校メール_MarkCowe（4月13日事案当日）.pdf",
            "学校法務顧問メール_寺井隼人（4月20日）.pdf",
            "案件概要.pdf",
            "相談カード.pdf",
            "証拠一覧と各証拠の目的説明.pdf",
            "証拠映像説明書面.pdf",
            "面談録音主要抜粋（4月14日学校面談）.pdf",
            "面談録音主要抜粋（4月14日学校面談）_2.pdf",
        ],
        "USB": [
            "注記：小牧家族に関するインターネット上の報道（http含むdocx）",
            "証拠1：2026年4月13日付　Angelinaによる手書きの書面陳述.pdf",
            "証拠2：2026年4月15日に取得した動画（約4分11秒・約4秒の2本のmp4）",
            "証拠3：事案発生現場の位置関係図（当方作成）.pdf",
            "証拠4：上記動画取得に関する微信（WeChat）のやり取り記録（jpg2点）",
            "証拠5：学校管理層とのメールのやり取り.pdf",
            "証拠6：学校の法務顧問とのメールのやり取り.pdf",
            "証拠7：2026年4月14日の面談録音.m4a",
            "証拠8：学校のスクールカウンセラーとのメールのやり取り.pdf",
            "証拠9：医療費に関する領収書（4/15・4/30・5/1心理カウンセリングinvoice）",
            "証拠10：オンライン高校（Crimson Global Academy）G9第2学期の学費（学費通知メール・銀行振込証明書）",
        ],
        "Wordファイル": [
            "260512電子内容証明案.docx（〜_3、中国語版、修正版、提出版、コピー、下書き）",
            "260515電子内容証明_提出版_名前訂正.pdf",
            "260515電子内容証明案_提出版_名前訂正.docx",
            "在学契約と同時履行に関する法律調査メモ.docx（_2あり）",
            "第１事案の概要.docx",
            "１まず謝罪させてほしい.docx",
        ],
        "委任関係": [
            "260507_委任契約書（馬様）.docx／.pdf／_2.docx",
            "260509_委任契約書（馬様）.pdf／署名済み.pdf",
        ],
        "経費": [
            "20260507_日本郵便株式会社_3043_領収書_通信費.pdf",
        ],
    }
    for folder, files in folder_struct.items():
        add_para(doc, f"■ {folder}", indent=0.3)
        for f in files:
            add_para(doc, f"・{f}", indent=0.8)

    # ---- 第10部　続編内容証明（追加請求書面）作成方針 ----
    add_heading(doc, "第１０　続編内容証明（追加請求書面）作成方針", 1)
    add_para(doc, "上記資料を踏まえた続編内容証明の構成案：")
    plan = [
        ("第１　前回内容証明（2026-05-12発送・5/15訂正）への対応状況",
         "送達確認、寺井弁護士による受領確認（PDF送付済）、現時点で具体的回答なし。"),
        ("第２　2026-05-15付寺井弁護士書面に対する当方見解",
         "①動画提出要求への対応方針（既に存在開示済み、強制力なし）、"
         "②現場検証要求への対応方針（Hirano氏は学校アドバイザーで中立性に疑義）、"
         "③学校CC送信による印象操作的記載への抗議、"
         "④Angelinaの登校拒絶解除が先決事項であることの再強調。"),
        ("第３　損害賠償の具体的請求",
         "①実費損害（CGA学費、KIST学費返金請求、定期券損失、医療費・交通費、弁護士費用）を金額付きで請求、"
         "②精神的損害として慰謝料7,000,000円を請求、"
         "③これらに対する支払期限を設定（例：通知到達後2週間）。"),
        ("第４　継続要求事項",
         "①Angelinaの登校拒絶解除、②George氏看護者の校内立入権回復、"
         "③Angelinaへの書面謝罪・名誉訂正声明発出、④学業遅延補填措置、"
         "⑤公平・安全な学習環境を保障する復学方案の提示、"
         "に対する具体的回答期限の設定。"),
        ("第５　応訴予告",
         "期限内に誠実な回答がなければ、生徒地位保全仮処分・本案訴訟提起の方針を予告。"),
    ]
    for h, b in plan:
        add_para(doc, f"■ {h}", indent=0.3)
        add_para(doc, b, indent=0.8)

    add_heading(doc, "確認事項（米谷弁護士）", 2)
    for item in [
        "慰謝料額：依頼者想定700万円を全額記載するか、当面300万円・残額留保とするか。"
        "（依頼者は5/22時点で『700万円で記載』を選好）",
        "動画提出要求：完全拒絶か、復学・看護者立入回復等を条件として応諾する余地ありとするか。",
        "現場検証要求：拒絶か、双方合意の中立第三者を提案するか。",
        "支払期限・回答期限：通知到達後2週間が標準だが、復学関連は学期末（2026-06-12）を考慮し短縮すべきか。",
        "送付方法：再び内容証明郵便（速達）か、寺井弁護士宛メール添付（PDF）を併用するか。",
        "言語：日本語のみか、中国語訳併用か（前回は日本語のみ）。",
    ]:
        add_para(doc, f"・{item}", indent=0.5)

    # save
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    main()
