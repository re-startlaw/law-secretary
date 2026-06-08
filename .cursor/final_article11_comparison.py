"""第11条の v3 vs v7_2 詳細比較。"""
from pathlib import Path
from docx import Document
import re

v3_path = Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v3.docx")
v7_2_path = Path("/Users/kometaninaoki/Downloads/合弁会社設立契約書_株式会社Sophia_v7_2.docx")

def get_article_text(doc, article_num):
    """条文を段落ごとに抽出。"""
    start_idx = None
    lines = []
    
    for i, p in enumerate(doc.paragraphs):
        if re.match(r'^第' + str(article_num) + r'条', p.text.strip()):
            start_idx = i
            break
    
    if start_idx is None:
        return None
    
    for i in range(start_idx, len(doc.paragraphs)):
        text = doc.paragraphs[i].text
        if i > start_idx and re.match(r'^第\d+条', text.strip()) and int(re.match(r'^第(\d+)条', text.strip()).group(1)) > article_num:
            break
        lines.append((i, text))
    
    return lines

doc_v3 = Document(str(v3_path))
doc_v7_2 = Document(str(v7_2_path))

article11_v3 = get_article_text(doc_v3, 11)
article11_v7_2 = get_article_text(doc_v7_2, 11)

print("=" * 120)
print("【第11条（知的財産権の処理）v3 vs v7_2 詳細比較】")
print("=" * 120)

print("\n【v3版】")
for idx, text in article11_v3:
    print(f"[{idx:03d}] {text}")

print("\n\n【v7_2版】")
for idx, text in article11_v7_2:
    print(f"[{idx:03d}] {text}")

print("\n\n" + "=" * 120)
print("【主な差分】")
print("=" * 120)

print("""
1. 第11条第2項の「なお」文が追加（v7_2のみ）
   v3: 「既存知的財産権は、引き続きそれぞれの当事者に帰属するものとする。甲又は乙が...」
   v7_2: （同上）＋ 「なお、顧客リスト等に個人情報が含まれる場合...」
   → 個人情報保護法に関する説明文が挿入

2. 第11条第3項全体が大幅改訂（v3 vs v7_2）
   v3: 簡潔に「本会社が解散・清算した場合、新規知的財産権の処理は甲乙協議で決定」
   v7_2: 冒頭に「将来、本会社を第三者に譲渡する場合に備え...」から始まり、
         解説的・推奨調の文章が多数含まれる（望ましい、価値の最大化と紛争の抑制、塩漬け、など）

3. 第11条第4項が新たに追加（v7_2のみ）
   「顧客リスト等に係る個人情報については、目的外利用...個人情報保護法等を遵守」

4. 第20条第5項で参照先が変更
   v3: 「第11条第3項に定めるところに従う」
   v7_2: 「第11条第3項及び第4項に定めるところに従う」
""")

