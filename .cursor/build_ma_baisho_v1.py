"""馬強事件 損害賠償請求書 docx 生成（普通Word版・v1）。

修正指示（260522_馬氏２.md）：
- 電子内容証明ではなく普通のword
- 精神的損害は100万円
- 金額は表で見やすく、各項目合計・全体合計を記載
- 第４は「前回の要求事項と併せて」とだけ書く（繰り返しなし）

雛形 02_ひな形/00_一般/内容証明.docx を shutil.copy2 でコピーし、
python-docx で run を温存しながら段落と表を組み立てる。
元ファイル 260522電子内容証明案_損害賠償請求.docx は保全のため残す。
"""

import os
import shutil
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/02_ひな形/"
    "00_一般/内容証明.docx"
)
DST = (
    "/Users/kometaninaoki/Library/CloudStorage/"
    "GoogleDrive-n.kometani@re-startlaw.com/マイドライブ/共有用/01_事件記録/"
    "ま_馬強/Wordファイル/260522_損害賠償請求書（馬様）.docx"
)


def rep(para, text):
    if para.runs:
        first = para.runs[0]
        first.text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def insert_before(anchor_para, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    new_p_el = deepcopy(anchor_para._element)
    for child in list(new_p_el):
        if child.tag.endswith("}r"):
            new_p_el.remove(child)
    anchor_para._element.addprevious(new_p_el)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p_el, anchor_para._parent)
    p.alignment = alignment
    p.add_run(text)
    return p


def insert_para_after(anchor_para, text):
    new_p = deepcopy(anchor_para._element)
    for child in list(new_p):
        if child.tag.endswith("}r"):
            new_p.remove(child)
    anchor_para._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, anchor_para._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)
    return p


def remove_para(para):
    para._element.getparent().remove(para._element)


def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    # 既存の tcBorders を上書き
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def add_table_before(doc, anchor_para, rows_data, col_widths=None):
    """anchor_para の直前に表を挿入。

    rows_data: List[List[str]]（最初の行はヘッダー想定）
    col_widths: List[Cm]（None なら自動）
    """
    n_rows = len(rows_data)
    n_cols = len(rows_data[0])
    # 末尾に一度作ってから XML を移動する
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = w
    for ri, row in enumerate(rows_data):
        for ci, text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            if ri == 0:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if ci == n_cols - 1:
                    # 金額列は右揃え
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_borders(cell)

    # 末尾に追加された表を anchor の前へ移動
    table_el = table._element
    table_el.getparent().remove(table_el)
    anchor_para._element.addprevious(table_el)
    return table


def main():
    os.makedirs(os.path.dirname(DST), exist_ok=True)
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    paras = list(doc.paragraphs)

    # ---- ヘッダー部分の書き換え ----
    rep(paras[0], "２０２６年５月◯日")
    rep(paras[1], "〒１０８－００７４")
    rep(paras[2], "東京都港区高輪四丁目２３番６号")
    rep(paras[3], "ハイホーム高輪７０８")
    rep(paras[4], "寺井勇人法律事務所")
    insert_para_after(paras[4], "弁護士　寺井　勇人　先生")

    # Para 5: （差出人）/ 6-10: 差出人 / 11: 空行 → TEL に差し替え
    rep(paras[11], "TEL：０３－６８２０－３８１５")

    # Para 12-16: 共同受任弁護士欄は不要なので削除
    for p in [paras[12], paras[13], paras[14], paras[15], paras[16]]:
        remove_para(p)

    # Para 17: 空行を残す
    # Para 18: タイトル (CENTER)
    rep(paras[18], "損害賠償請求書")

    # Para 19: 空行 / Para 20-34: 本文を全削除 / Para 35: 「草々」→「以上」
    for p in paras[20:35]:
        remove_para(p)
    rep(paras[35], "以上")

    # ---- 本文挿入 ----
    anchor = paras[35]  # 「以上」段落

    # 前文
    insert_before(anchor,
                  "　当職は、貴職が代理人として対応されているK. International "
                  "School Tokyo（以下「貴校」といいます。）にG9学年として在籍"
                  "するAngelina Feili Ma氏（以下「Angelina氏」といいます。）"
                  "の父であり、貴校との間で在学契約を締結している馬強氏（以下"
                  "「通知人」といいます。）の代理人として、貴職に対し、令和８年"
                  "５月１２日付「受任通知及び御連絡」（以下「前回通知」と"
                  "いいます。）に引き続き、以下のとおり通知人の貴校及び貴校"
                  "理事会議長小牧孝子氏（以下「小牧氏」といいます。）に対する"
                  "損害賠償請求について、御連絡申し上げます。")
    insert_before(anchor, "")

    # 第１ 前回通知の整理
    insert_before(anchor, "第１　前回通知の整理")
    insert_before(anchor,
                  "　前回通知においては、貴校及び小牧氏による不法行為の概要、"
                  "現に通知人及びその家族に生じている損害の項目並びに通知人の"
                  "貴校に対する要求事項を御連絡したほか、貴校による教育役務提供"
                  "義務の履行拒絶に対し、通知人が同時履行の抗弁権を行使し、"
                  "Angelina氏のG10学年第一期学費相当額（金１,３８８,０００円）"
                  "を当職弁護士預り金口座において保全している旨を御連絡いたし"
                  "ました。")
    insert_before(anchor,
                  "　本書面は、前回通知においては個別の金額を留保していた損害"
                  "賠償請求の具体的内容について、改めて御連絡するものです。")
    insert_before(anchor, "")

    # 第２ 損害賠償請求
    insert_before(anchor, "第２　損害賠償請求")
    insert_before(anchor,
                  "　前記前回通知第２記載の貴校及び小牧氏による不法行為（民法"
                  "第７０９条、同法第７１５条又は第７１９条）並びに貴校による"
                  "在学契約上の教育役務提供義務の履行拒絶（同法第４１５条）に"
                  "より、通知人及びその家族には、以下のとおり、金銭に換算可能"
                  "な損害が現に生じております。通知人は、貴校及び小牧氏に対し、"
                  "連帯して、下記合計額の支払を請求いたします。")
    insert_before(anchor, "")

    # （１）実費損害
    insert_before(anchor, "（１）　実費損害")
    actual_rows = [
        ["項目", "金額"],
        ["ア　Angelina氏の代替教育（Crimson Global Academy "
         "G9 Term 2）学費", "金１,２６０,０００円"],
        ["イ　Angelina氏のG10学年オンラインスクール預託金（１０％ Deposit）",
         "金２４４,８００円"],
        ["ウ　KIST G9学年学費のうち登校拒絶期間に対応する按分額\n"
         "（年額金２,８９７,０００円のうち令和８年４月１４日から正常復学日まで）",
         "復学日確定後に確定"],
        ["エ　通学定期券損失（令和８年１月１７日購入の６か月通学定期券\n"
         "金３５,９８０円のうち令和８年４月１４日以降の未利用期間相当額）",
         "未利用期間確定後に確定"],
        ["オ　令和８年４月１６日付精神科診療費", "金２１,８５０円"],
        ["カ　令和８年４月３０日付精神科診療費", "金５７,０５０円"],
        ["キ　令和８年５月１日付心理カウンセリング費用", "金１７,０００円"],
        ["ク　令和８年５月８日付心理カウンセリング費用", "金１７,０００円"],
        ["ケ　Angelina氏の精神科及び心理カウンセリング通院に伴う交通費\n"
         "（保護者１名同行を要し、精神科及び心理カウンセリング各回往復金"
         "８４０円）", "通院終了後に確定"],
        ["コ　以後の精神科及び心理カウンセリング治療費用",
         "治療終了後に確定"],
        ["実費損害　既確定額合計", "金１,６１７,７００円"],
    ]
    add_table_before(doc, anchor, actual_rows,
                     col_widths=[Cm(11.0), Cm(5.5)])
    insert_before(anchor, "")
    insert_before(anchor,
                  "　ウ、エ、ケ及びコの各項目については、現に継続して発生して"
                  "おり、最終的な金額は復学日及び治療終了日の確定をもって精算"
                  "するものといたします。")
    insert_before(anchor, "")

    # （２）精神的損害
    insert_before(anchor, "（２）　精神的損害（慰謝料）")
    insert_before(anchor,
                  "　貴校及び小牧氏の不法行為により、Angelina氏及び通知人らは、"
                  "以下に列挙する精神的苦痛を受けております。")
    insert_before(anchor,
                  "ア　令和８年４月１３日における小牧氏の有形力行使（襟元付近"
                  "の身体的接触を伴う移動及び大声による叱責）により、"
                  "Angelina氏が受けた精神的苦痛")
    insert_before(anchor,
                  "イ　約１か月以上にわたる登校拒絶措置の継続により、"
                  "Angelina氏が正常な学習機会及び学級内の人間関係を喪失した"
                  "ことに伴う精神的苦痛並びに不安、心理的圧迫")
    insert_before(anchor,
                  "ウ　貴校及び小牧氏が、Angelina氏の書面陳述記載の事実関係に"
                  "ついて、事実に反する主張を行っているかのような取扱いを継続"
                  "したことにより、Angelina氏が受けた名誉権、人格権及び校内"
                  "評価に対する侵害並びに精神的苦痛")
    insert_before(anchor,
                  "エ　Angelina氏の精神状態の悪化、投薬量の増量、継続的心理"
                  "カウンセリングを要する状態に至った結果、Angelina氏及び"
                  "通知人らが受けた精神的苦痛")
    insert_before(anchor,
                  "オ　Angelina氏の学業継続性、試験成績、GPA及びTranscriptへ"
                  "の不利益並びに将来の大学進学経路への現実的影響に伴う精神的"
                  "苦痛")
    insert_before(anchor,
                  "カ　George氏の看護者に対する校内立入禁止措置の継続により、"
                  "George氏の就学環境が害されたこと並びにこれに対応するために"
                  "通知人らが負担を強いられたことに伴う精神的苦痛")
    insert_before(anchor, "")
    mental_rows = [
        ["項目", "金額"],
        ["Angelina氏及び通知人らに対する慰謝料", "金１,０００,０００円"],
        ["精神的損害　合計", "金１,０００,０００円"],
    ]
    add_table_before(doc, anchor, mental_rows,
                     col_widths=[Cm(11.0), Cm(5.5)])
    insert_before(anchor, "")

    # （３）弁護士費用
    insert_before(anchor, "（３）　弁護士費用")
    insert_before(anchor,
                  "　通知人は、本件法的対応のために当職に委任せざるを得ず、"
                  "貴校及び小牧氏の不法行為と相当因果関係を有する弁護士費用と"
                  "して、前記（１）及び（２）の既確定額合計（金２,６１７,"
                  "７００円）の１０％相当額を請求いたします。")
    insert_before(anchor, "")
    fee_rows = [
        ["項目", "金額"],
        ["弁護士費用（実費損害＋精神的損害の１０％）", "金２６１,７７０円"],
    ]
    add_table_before(doc, anchor, fee_rows,
                     col_widths=[Cm(11.0), Cm(5.5)])
    insert_before(anchor, "")

    # （４）総合計
    insert_before(anchor, "（４）　総合計")
    total_rows = [
        ["項目", "金額"],
        ["実費損害（既確定額）", "金１,６１７,７００円"],
        ["精神的損害（慰謝料）", "金１,０００,０００円"],
        ["弁護士費用", "金２６１,７７０円"],
        ["総合計（既確定額）", "金２,８７９,４７０円"],
    ]
    add_table_before(doc, anchor, total_rows,
                     col_widths=[Cm(11.0), Cm(5.5)])
    insert_before(anchor,
                  "　なお、実費損害ウ、エ、ケ及びコの各項目に係る損害額は、"
                  "前記のとおり継続発生中であり、これらが確定した場合には、"
                  "総合計額に加算してご請求申し上げます。")
    insert_before(anchor, "")

    # 第３ 支払方法及び支払期限
    insert_before(anchor, "第３　支払方法及び支払期限")
    insert_before(anchor,
                  "　貴校及び小牧氏は、本書面到達後◯日以内に、前記第２（４）"
                  "記載の総合計額を、下記当職指定口座へお振込みの方法により"
                  "お支払いください。")
    insert_before(anchor, "　【振込先】")
    insert_before(anchor,
                  "　　金融機関：（追って当職よりご指定申し上げます。）")
    insert_before(anchor, "　　口座種別・口座番号：（同上）")
    insert_before(anchor,
                  "　　口座名義：弁護士法人Ｒｅ－Ｓｔａｒｔ法律事務所　預り口")
    insert_before(anchor, "")

    # 第４ 要求事項及び御回答期限
    insert_before(anchor, "第４　要求事項及び御回答期限")
    insert_before(anchor,
                  "　前回通知第５記載の各要求事項と併せて、本書面到達後◯日以内"
                  "に、貴校の対応方針について御回答くださいますよう求めます。")
    insert_before(anchor, "")

    # 第５ 法的措置の予告
    insert_before(anchor, "第５　法的措置の予告")
    insert_before(anchor,
                  "　万一、前記第３記載の支払期限内に貴校及び小牧氏から本件請求"
                  "金額の全部又は相当部分の支払がない場合、又は前記第４記載の"
                  "御回答期限内に貴校から誠実な御回答が得られない場合、当方と"
                  "しては、Angelina氏の生徒たる地位保全の仮処分命令申立て"
                  "（民事保全法第２３条第２項）並びに本案訴訟の提起その他の"
                  "法的措置を直ちに講じる所存です。")
    insert_before(anchor,
                  "　また、前記第３記載の支払期限を経過した本件請求金額に"
                  "ついては、支払期限の翌日から支払済みまで、年３％の割合に"
                  "よる遅延損害金を併せ請求いたします。")
    insert_before(anchor, "")

    # 結語
    insert_before(anchor,
                  "　本書面の内容を御確認いただき、御不明な点がございましたら、"
                  "ご遠慮なくお尋ねくださいますようお願い申し上げます。")
    insert_before(anchor, "")

    doc.save(DST)
    print("saved:", DST)


if __name__ == "__main__":
    main()
