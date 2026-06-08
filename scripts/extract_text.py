"""ファイル内容抽出ユーティリティ（PDF/Word/Excel/テキスト統合）。

secretary.py の LLM 分類フォールバックや手動分類タスクから呼ばれる。
1ファイル1抽出を強制するためにモジュールレベルのメモリキャッシュを持つ。
"""

from __future__ import annotations

import json
import os
import sys
from typing import Any

_CACHE: dict[str, dict[str, Any]] = {}


def _extract_pdf_via_ocr(path: str, max_chars: int) -> tuple[str, str]:
    """fitzでページをラスタ→tesseract OCR。tesseract本体と jpn traineddata が
    両方揃っているときのみ動く。揃っていなければ ("", "ocr-...") を返す。
    """
    try:
        import io

        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except ImportError:
        return "", "ocr-no-deps"

    try:
        langs = set(pytesseract.get_languages(config=""))
    except Exception as e:  # tesseract 本体未インストール等
        return "", f"ocr-tesseract-error:{e.__class__.__name__}"

    if "jpn" not in langs:
        return "", "ocr-no-jpn-traineddata"

    lang_codes = ["jpn"]
    if "jpn_vert" in langs:
        lang_codes.append("jpn_vert")
    if "eng" in langs:
        lang_codes.append("eng")
    lang = "+".join(lang_codes)

    try:
        doc = fitz.open(path)
    except Exception as e:
        return "", f"fitz-open-error:{e.__class__.__name__}"

    parts: list[str] = []
    try:
        for page in doc:
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(img, lang=lang)
            if text and text.strip():
                parts.append(text)
            if sum(len(p) for p in parts) >= max_chars:
                break
    except Exception as e:
        doc.close()
        return "", f"ocr-runtime-error:{e.__class__.__name__}"
    finally:
        try:
            doc.close()
        except Exception:
            pass

    text = "\n".join(parts).strip()
    if text:
        return text[:max_chars], f"tesseract:{lang}"
    return "", "ocr-no-text"


def _extract_pdf(path: str, max_chars: int) -> tuple[str, str]:
    """PDFテキスト抽出。fitz→pdfminer→OCR(tesseract) の順で試行し (text, extractor) を返す。"""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(path)
        parts: list[str] = []
        for page in doc:
            parts.append(page.get_text())
            if sum(len(p) for p in parts) >= max_chars:
                break
        doc.close()
        text = "".join(parts)
        if text.strip():
            return text[:max_chars], "fitz"
    except Exception:
        pass

    try:
        from pdfminer.high_level import extract_text as pdfminer_extract

        text = pdfminer_extract(path) or ""
        if text.strip():
            return text[:max_chars], "pdfminer"
    except Exception:
        pass

    text, kind = _extract_pdf_via_ocr(path, max_chars)
    if text.strip():
        return text, kind
    return "", kind or "pdf-failed"


def _extract_docx(path: str, max_chars: int) -> tuple[str, str]:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
        if sum(len(p) for p in parts) >= max_chars:
            break
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
            if sum(len(p) for p in parts) >= max_chars:
                break
    return "\n".join(parts)[:max_chars], "python-docx"


def _extract_xlsx(path: str, max_chars: int) -> tuple[str, str]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[シート: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append("\t".join(cells))
            if sum(len(p) for p in parts) >= max_chars:
                break
        if sum(len(p) for p in parts) >= max_chars:
            break
    wb.close()
    return "\n".join(parts)[:max_chars], "openpyxl"


def _extract_doc(path: str, max_chars: int) -> tuple[str, str]:
    """旧 .doc（Word97-2003）の本文抽出。

    優先順:
      1. macOS 標準の `textutil` で utf-8 txt に変換して読む（最も安定）
      2. `antiword` コマンドが PATH にあれば使う
    どれも失敗したら ("", "doc-failed")。
    """
    import shutil as _sh
    import subprocess
    import tempfile

    textutil = "/usr/bin/textutil"
    if os.path.exists(textutil):
        out_path = ""
        try:
            with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tf:
                out_path = tf.name
            subprocess.run(
                [
                    textutil, "-convert", "txt", "-encoding", "UTF-8",
                    "-output", out_path, path,
                ],
                check=True, capture_output=True, timeout=30,
            )
            with open(out_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read(max_chars)
            if text.strip():
                return text[:max_chars], "textutil"
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
            pass
        finally:
            if out_path:
                try:
                    os.unlink(out_path)
                except OSError:
                    pass

    antiword = _sh.which("antiword")
    if antiword:
        try:
            res = subprocess.run(
                [antiword, "-w", "0", path],
                check=True, capture_output=True, timeout=30,
            )
            text = res.stdout.decode("utf-8", errors="replace")
            if text.strip():
                return text[:max_chars], "antiword"
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
            pass

    return "", "doc-failed"


def _extract_plain(path: str, max_chars: int) -> tuple[str, str]:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read(max_chars), "plain"


def _prepare_image(path: str) -> tuple[str | None, str | None]:
    """画像ファイルを vision 送信用 JPG パスへ正規化する。

    .jpg/.jpeg/.png はそのまま、.heic/.heif は macOS の `sips` で JPG 変換。
    Returns: (image_path, error)。失敗時は (None, err文字列)。
    """
    import subprocess
    import tempfile

    ext = os.path.splitext(path)[1].lower()
    if ext in (".jpg", ".jpeg", ".png"):
        return path, None
    if ext in (".heic", ".heif"):
        sips = "/usr/bin/sips"
        if not os.path.exists(sips):
            return None, "sips-not-found"
        try:
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tf:
                out_path = tf.name
            subprocess.run(
                [sips, "-s", "format", "jpeg", path, "--out", out_path],
                check=True, capture_output=True, timeout=30,
            )
            if os.path.getsize(out_path) > 0:
                return out_path, None
            return None, "sips-empty-output"
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError) as e:
            return None, f"sips-error:{e.__class__.__name__}"
    return None, None


def extract_text(path: str, max_chars: int = 4000) -> dict[str, Any]:
    """ファイル内容を抽出する。

    Returns:
        {
            "text": str,              # 抽出本文（max_chars で切詰）
            "extracted_chars": int,
            "needs_ocr": bool,        # PDFで文字が取れなかった（画像のみ）場合 True
            "extractor": str,         # 使用した抽出器名
            "error": str | None,
            "image_path": str | None, # vision 送信用画像パス（画像ファイルのみ）
        }
    """
    cache_key = f"{path}:{max_chars}"
    if cache_key in _CACHE:
        return _CACHE[cache_key]

    if not os.path.isfile(path):
        result = {
            "text": "",
            "extracted_chars": 0,
            "needs_ocr": False,
            "extractor": "none",
            "error": f"file not found: {path}",
        }
        _CACHE[cache_key] = result
        return result

    ext = os.path.splitext(path)[1].lower()
    text = ""
    extractor = "unsupported"
    error: str | None = None
    image_path: str | None = None

    try:
        if ext == ".pdf":
            text, extractor = _extract_pdf(path, max_chars)
        elif ext == ".docx":
            text, extractor = _extract_docx(path, max_chars)
        elif ext == ".doc":
            text, extractor = _extract_doc(path, max_chars)
        elif ext == ".xlsx":
            text, extractor = _extract_xlsx(path, max_chars)
        elif ext in (".txt", ".md", ".csv"):
            text, extractor = _extract_plain(path, max_chars)
        elif ext in (".heic", ".heif", ".jpg", ".jpeg", ".png"):
            image_path, image_err = _prepare_image(path)
            if image_path:
                extractor = f"image:{ext.lstrip('.')}"
            else:
                error = image_err or f"image-prepare-failed: {ext}"
        else:
            error = f"unsupported extension: {ext}"
    except Exception as e:
        error = f"{type(e).__name__}: {e}"

    needs_ocr = ext == ".pdf" and not text.strip() and error is None

    result = {
        "text": text,
        "extracted_chars": len(text),
        "needs_ocr": needs_ocr,
        "extractor": extractor,
        "error": error,
        "image_path": image_path,
    }
    _CACHE[cache_key] = result
    return result


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: extract_text.py <path> [<path>...]", file=sys.stderr)
        return 2
    out = [extract_text(p) for p in sys.argv[1:]]
    print(json.dumps(out, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
